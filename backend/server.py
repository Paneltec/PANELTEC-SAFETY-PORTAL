from fastapi import FastAPI, APIRouter, HTTPException, Depends, Header, WebSocket, WebSocketDisconnect, Request
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import json
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Any, Dict
import uuid
import secrets
import base64
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

JWT_SECRET = os.environ.get('JWT_SECRET', 'change-me')
JWT_ALG = 'HS256'
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

app = FastAPI(title="Paneltec Safety Portal API")
api_router = APIRouter(prefix="/api")

# ============== UTILS ==============
def now_iso():
    return datetime.now(timezone.utc).isoformat()

def hash_pw(pw: str) -> str:
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()

def verify_pw(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode(), hashed.encode())
    except Exception:
        return False

def make_token(user_id: str, role: str) -> str:
    payload = {
        'sub': user_id,
        'role': role,
        'exp': datetime.now(timezone.utc) + timedelta(days=30)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

async def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith('Bearer '):
        raise HTTPException(401, 'Missing token')
    token = authorization.split(' ', 1)[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
    except Exception:
        raise HTTPException(401, 'Invalid token')
    user = await db.users.find_one({'id': payload['sub']}, {'_id': 0, 'password': 0})
    if not user:
        raise HTTPException(401, 'User not found')
    return user

async def require_admin(user=Depends(get_current_user)):
    if user.get('role') != 'admin':
        raise HTTPException(403, 'Admin only')
    return user

# ============== MODELS ==============
class LoginIn(BaseModel):
    email: str
    password: str

class RegisterIn(BaseModel):
    email: str
    password: str
    name: str
    role: str = 'worker'

class WorkerAvailabilityDay(BaseModel):
    enabled: bool = False
    start: str = "06:00"  # 24hr HH:MM
    end: str = "15:00"

class Worker(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str  # legacy combined name (kept for backwards compat)
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    role: str = 'worker'  # worker, supervisor, admin, manager
    trade: Optional[str] = None
    # Address
    country: Optional[str] = 'AUSTRALIA'
    state: Optional[str] = None
    street_address: Optional[str] = None
    suburb: Optional[str] = None
    postal_code: Optional[str] = None
    birth_date: Optional[str] = None  # YYYY-MM-DD
    additional_notes: Optional[str] = None
    # Assignment
    location_ids: List[str] = []
    client_ids: List[str] = []  # Simpro clients/companies
    skills: List[str] = []
    # Status
    is_manager: bool = False
    license_allocated: bool = False
    license_allocated_by: Optional[str] = None  # admin user name who granted access
    status: str = 'active'  # active, inactive
    # Availability (Mon..Sun)
    availability: Dict[str, Dict[str, Any]] = Field(default_factory=lambda: {
        d: {'enabled': False, 'start': '06:00', 'end': '15:00'}
        for d in ['monday','tuesday','wednesday','thursday','friday','saturday','sunday']
    })
    # Simpro
    simpro_id: Optional[str] = None
    simpro_company_id: Optional[str] = None
    simpro_company_name: Optional[str] = None
    source: str = 'manual'  # manual, simpro
    signature_b64: Optional[str] = None
    photo_b64: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

class Certification(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    worker_id: str
    name: str
    issuer: Optional[str] = None
    issued_date: Optional[str] = None
    expiry_date: Optional[str] = None
    image_b64: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

class Location(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    address: Optional[str] = None
    project_code: Optional[str] = None
    worker_ids: List[str] = []
    created_at: str = Field(default_factory=now_iso)

class FormField(BaseModel):
    id: str
    label: str
    type: str  # text, textarea, number, date, select, checkbox, radio, signature, photo, gps, repeatable
    required: bool = False
    options: List[str] = []  # for select/radio
    placeholder: Optional[str] = None
    is_critical: bool = False  # flag this form as critical (incident/near-miss)

class FormTemplate(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    category: str = 'general'  # general, incident, inspection, toolbox, near_miss
    is_private: bool = False
    fields: List[Dict[str, Any]] = []
    created_at: str = Field(default_factory=now_iso)

class FormSubmission(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    template_id: str
    template_name: str
    category: str = 'general'
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    worker_id: Optional[str] = None
    worker_name: Optional[str] = None
    answers: Dict[str, Any] = {}
    signature_b64: Optional[str] = None
    photos_b64: List[str] = []
    gps_lat: Optional[float] = None
    gps_lng: Optional[float] = None
    flagged: bool = False
    flag_note: Optional[str] = None
    ai_summary: Optional[str] = None
    submitted_at: str = Field(default_factory=now_iso)

# ============== AUTH ==============
@api_router.post('/auth/register')
async def register(body: RegisterIn):
    existing = await db.users.find_one({'email': body.email})
    if existing:
        raise HTTPException(400, 'Email already registered')
    uid = str(uuid.uuid4())
    doc = {
        'id': uid,
        'email': body.email,
        'password': hash_pw(body.password),
        'name': body.name,
        'role': body.role,
        'created_at': now_iso(),
    }
    await db.users.insert_one(doc)
    # Also create worker profile if worker
    if body.role == 'worker':
        w = Worker(name=body.name, email=body.email, role='worker').model_dump()
        await db.workers.insert_one(w)
    token = make_token(uid, body.role)
    return {'token': token, 'user': {'id': uid, 'email': body.email, 'name': body.name, 'role': body.role}}

@api_router.post('/auth/login')
async def login(body: LoginIn):
    user = await db.users.find_one({'email': body.email})
    if not user or not verify_pw(body.password, user['password']):
        raise HTTPException(401, 'Invalid credentials')
    token = make_token(user['id'], user['role'])
    return {'token': token, 'user': {'id': user['id'], 'email': user['email'], 'name': user['name'], 'role': user['role']}}

@api_router.get('/auth/me')
async def me(user=Depends(get_current_user)):
    return user

# ============== WORKERS ==============
@api_router.get('/workers')
async def list_workers(user=Depends(get_current_user)):
    items = await db.workers.find({}, {'_id': 0}).to_list(1000)
    return items

@api_router.post('/workers')
async def create_worker(w: Worker, user=Depends(require_admin)):
    doc = w.model_dump()
    await db.workers.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/workers/{wid}')
async def update_worker(wid: str, w: Worker, user=Depends(require_admin)):
    doc = w.model_dump()
    doc['id'] = wid
    await db.workers.update_one({'id': wid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/workers/{wid}')
async def delete_worker(wid: str, user=Depends(require_admin)):
    await db.workers.delete_one({'id': wid})
    await db.certifications.delete_many({'worker_id': wid})
    return {'ok': True}

# ============== CERTIFICATIONS ==============
@api_router.get('/certifications')
async def list_certs(user=Depends(get_current_user)):
    items = await db.certifications.find({}, {'_id': 0}).to_list(2000)
    return items

@api_router.post('/certifications')
async def create_cert(c: Certification, user=Depends(get_current_user)):
    doc = c.model_dump()
    await db.certifications.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/certifications/{cid}')
async def delete_cert(cid: str, user=Depends(get_current_user)):
    await db.certifications.delete_one({'id': cid})
    return {'ok': True}

# ============== LOCATIONS ==============
@api_router.get('/locations')
async def list_locations(user=Depends(get_current_user)):
    items = await db.locations.find({}, {'_id': 0}).to_list(1000)
    return items

@api_router.post('/locations')
async def create_location(l: Location, user=Depends(require_admin)):
    doc = l.model_dump()
    await db.locations.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/locations/{lid}')
async def update_location(lid: str, l: Location, user=Depends(require_admin)):
    doc = l.model_dump()
    doc['id'] = lid
    await db.locations.update_one({'id': lid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/locations/{lid}')
async def delete_location(lid: str, user=Depends(require_admin)):
    await db.locations.delete_one({'id': lid})
    return {'ok': True}

# ============== FORM TEMPLATES ==============
@api_router.get('/forms/templates')
async def list_templates(user=Depends(get_current_user)):
    items = await db.form_templates.find({}, {'_id': 0}).to_list(1000)
    return items

@api_router.get('/forms/templates/{tid}')
async def get_template(tid: str, user=Depends(get_current_user)):
    t = await db.form_templates.find_one({'id': tid}, {'_id': 0})
    if not t:
        raise HTTPException(404, 'Not found')
    return t

@api_router.post('/forms/templates')
async def create_template(t: FormTemplate, user=Depends(require_admin)):
    doc = t.model_dump()
    await db.form_templates.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/forms/templates/{tid}')
async def update_template(tid: str, t: FormTemplate, user=Depends(require_admin)):
    doc = t.model_dump()
    doc['id'] = tid
    await db.form_templates.update_one({'id': tid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/forms/templates/{tid}')
async def delete_template(tid: str, user=Depends(require_admin)):
    await db.form_templates.delete_one({'id': tid})
    return {'ok': True}

# ============== SUBMISSIONS ==============
@api_router.get('/submissions')
async def list_submissions(user=Depends(get_current_user)):
    items = await db.submissions.find({}, {'_id': 0}).sort('submitted_at', -1).to_list(2000)
    return items

@api_router.get('/submissions/{sid}')
async def get_submission(sid: str, user=Depends(get_current_user)):
    s = await db.submissions.find_one({'id': sid}, {'_id': 0})
    if not s:
        raise HTTPException(404, 'Not found')
    return s

@api_router.post('/submissions')
async def create_submission(s: FormSubmission, user=Depends(get_current_user)):
    doc = s.model_dump()
    # auto-flag critical categories
    if doc.get('category') in ('incident', 'near_miss'):
        doc['flagged'] = True
    await db.submissions.insert_one(doc)
    doc.pop('_id', None)
    # Trigger auto-share rules (mocked email log)
    try:
        await evaluate_share_rules(doc)
    except Exception as e:
        logger.warning(f"Share rule eval failed: {e}")
    return doc

@api_router.delete('/submissions/{sid}')
async def delete_submission(sid: str, user=Depends(require_admin)):
    await db.submissions.delete_one({'id': sid})
    return {'ok': True}

# ============== AI ==============
class AISummaryIn(BaseModel):
    submission_id: str

@api_router.post('/ai/summarize')
async def ai_summarize(body: AISummaryIn, user=Depends(get_current_user)):
    sub = await db.submissions.find_one({'id': body.submission_id}, {'_id': 0})
    if not sub:
        raise HTTPException(404, 'Submission not found')
    
    # Build context
    parts = [
        f"Form: {sub.get('template_name')}",
        f"Category: {sub.get('category')}",
        f"Location: {sub.get('location_name')}",
        f"Worker: {sub.get('worker_name')}",
        f"Submitted: {sub.get('submitted_at')}",
        "\nAnswers:",
    ]
    for k, v in (sub.get('answers') or {}).items():
        parts.append(f"- {k}: {v}")
    context = "\n".join(parts)

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"summary-{body.submission_id}",
            system_message=(
                "You are a safety compliance analyst at Paneltec, a civil contractor. "
                "Summarize the following safety form submission in 3-5 concise bullet points. "
                "Highlight key risks, root causes, and recommended corrective actions. "
                "Be specific, professional, and actionable."
            )
        ).with_model('openai', 'gpt-4o-mini')
        msg = UserMessage(text=context)
        result = await chat.send_message(msg)
        summary = str(result)
    except Exception as e:
        logger.error(f"AI error: {e}")
        summary = f"(AI summary unavailable: {str(e)[:120]})"

    await db.submissions.update_one({'id': body.submission_id}, {'$set': {'ai_summary': summary}})
    return {'summary': summary}

# ============== ANALYTICS ==============
@api_router.get('/analytics/overview')
async def analytics_overview(user=Depends(get_current_user)):
    subs = await db.submissions.find({}, {'_id': 0}).to_list(5000)
    workers = await db.workers.find({}, {'_id': 0}).to_list(2000)
    locations = await db.locations.find({}, {'_id': 0}).to_list(2000)
    certs = await db.certifications.find({}, {'_id': 0}).to_list(5000)

    incidents = [s for s in subs if s.get('category') == 'incident']
    near_misses = [s for s in subs if s.get('category') == 'near_miss']
    inspections = [s for s in subs if s.get('category') == 'inspection']
    toolbox = [s for s in subs if s.get('category') == 'toolbox']

    # By location
    by_loc = {}
    for s in subs:
        ln = s.get('location_name') or 'Unassigned'
        by_loc[ln] = by_loc.get(ln, 0) + 1
    by_location = [{'location': k, 'count': v} for k, v in sorted(by_loc.items(), key=lambda x: -x[1])][:10]

    # By worker (incident-prone)
    by_worker = {}
    for s in subs:
        if s.get('category') in ('incident', 'near_miss'):
            wn = s.get('worker_name') or 'Unknown'
            by_worker[wn] = by_worker.get(wn, 0) + 1
    accident_prone = [{'worker': k, 'count': v} for k, v in sorted(by_worker.items(), key=lambda x: -x[1])][:8]

    # Monthly trend
    monthly = {}
    for s in subs:
        try:
            dt = datetime.fromisoformat(s['submitted_at'].replace('Z', '+00:00'))
            key = dt.strftime('%Y-%m')
            monthly.setdefault(key, {'month': key, 'incidents': 0, 'near_misses': 0, 'inspections': 0, 'toolbox': 0})
            cat = s.get('category', 'general')
            if cat == 'incident':
                monthly[key]['incidents'] += 1
            elif cat == 'near_miss':
                monthly[key]['near_misses'] += 1
            elif cat == 'inspection':
                monthly[key]['inspections'] += 1
            elif cat == 'toolbox':
                monthly[key]['toolbox'] += 1
        except Exception:
            pass
    monthly_trend = sorted(monthly.values(), key=lambda x: x['month'])[-12:]

    # Expiring certs (30 days)
    soon = []
    today = datetime.now(timezone.utc).date()
    for c in certs:
        if c.get('expiry_date'):
            try:
                exp = datetime.fromisoformat(c['expiry_date']).date()
                days = (exp - today).days
                if days <= 60:
                    worker_name = next((w['name'] for w in workers if w['id'] == c.get('worker_id')), 'Unknown')
                    soon.append({
                        'cert_id': c['id'],
                        'name': c['name'],
                        'worker_name': worker_name,
                        'expiry_date': c['expiry_date'],
                        'days_remaining': days,
                        'status': 'expired' if days < 0 else ('critical' if days <= 14 else 'warning')
                    })
            except Exception:
                pass
    soon.sort(key=lambda x: x['days_remaining'])

    return {
        'kpis': {
            'total_submissions': len(subs),
            'incidents': len(incidents),
            'near_misses': len(near_misses),
            'inspections': len(inspections),
            'toolbox_talks': len(toolbox),
            'workers': len(workers),
            'locations': len(locations),
            'flagged': len([s for s in subs if s.get('flagged')]),
            'expiring_certs': len(soon),
        },
        'by_location': by_location,
        'accident_prone': accident_prone,
        'monthly_trend': monthly_trend,
        'expiring_certifications': soon[:20],
    }

# ============== SEED ==============
@api_router.post('/seed')
async def seed_demo():
    """Seed demo data for a civil contractor (Paneltec)."""
    # Clear
    await db.users.delete_many({})
    await db.workers.delete_many({})
    await db.locations.delete_many({})
    await db.certifications.delete_many({})
    await db.form_templates.delete_many({})
    await db.submissions.delete_many({})

    # Admin + worker
    admin_id = str(uuid.uuid4())
    await db.users.insert_one({
        'id': admin_id, 'email': 'admin@paneltec.com',
        'password': hash_pw('admin123'), 'name': 'Admin Paneltec',
        'role': 'admin', 'created_at': now_iso()
    })
    worker_uid = str(uuid.uuid4())
    await db.users.insert_one({
        'id': worker_uid, 'email': 'worker@paneltec.com',
        'password': hash_pw('worker123'), 'name': 'John Carpenter',
        'role': 'worker', 'created_at': now_iso()
    })

    # Locations - civil contractor jobsites
    locs = [
        {'name': 'Highway 401 Expansion - Section A', 'address': 'Toronto, ON', 'project_code': 'PT-HW401-A'},
        {'name': 'Downtown Bridge Reconstruction', 'address': 'Mississauga, ON', 'project_code': 'PT-BR-DT'},
        {'name': 'Industrial Park Foundation Site', 'address': 'Brampton, ON', 'project_code': 'PT-IP-FN'},
        {'name': 'Waterfront Drainage Project', 'address': 'Hamilton, ON', 'project_code': 'PT-WF-DR'},
    ]
    location_docs = []
    for l in locs:
        d = Location(**l).model_dump()
        location_docs.append(d)
    await db.locations.insert_many([dict(x) for x in location_docs])

    # Workers - civil contractor trades
    worker_data = [
        ('John Carpenter', 'Site Supervisor', 'supervisor'),
        ('Mike Rodriguez', 'Heavy Equipment Operator', 'worker'),
        ('Sarah Thompson', 'Concrete Finisher', 'worker'),
        ('David Chen', 'Surveyor', 'worker'),
        ('Emily Brooks', 'Safety Officer', 'supervisor'),
        ('Carlos Mendez', 'Excavator Operator', 'worker'),
        ('Liam OConnor', 'Foreman', 'supervisor'),
        ('Aisha Patel', 'Quality Inspector', 'worker'),
        ('Robert Kane', 'Pipelayer', 'worker'),
        ('Tom Walker', 'Crane Operator', 'worker'),
    ]
    worker_docs = []
    for name, trade, role in worker_data:
        w = Worker(name=name, trade=trade, role=role,
                   email=f"{name.split()[0].lower()}@paneltec.com",
                   location_ids=[location_docs[len(worker_docs) % len(location_docs)]['id']]).model_dump()
        worker_docs.append(w)
    await db.workers.insert_many([dict(x) for x in worker_docs])

    # Certifications with varied expiries
    today = datetime.now(timezone.utc).date()
    cert_data = [
        ('WHMIS 2015', 'Ontario Labour', 200),
        ('Working at Heights', 'MOL', 45),
        ('Confined Space Entry', 'IHSA', -10),
        ('First Aid / CPR', 'Red Cross', 90),
        ('Heavy Equipment Operator', 'IHSA', 12),
        ('Traffic Control Person', 'OTM Book 7', 365),
        ('Fall Protection', 'IHSA', 7),
        ('Excavation & Trenching', 'CSAO', 180),
    ]
    cert_docs = []
    for i, w in enumerate(worker_docs):
        # 2 certs per worker
        for j in range(2):
            cname, issuer, days = cert_data[(i + j) % len(cert_data)]
            exp = today + timedelta(days=days + (i * 5))
            issued = today - timedelta(days=365 - (i * 3))
            cert_docs.append(Certification(
                worker_id=w['id'], name=cname, issuer=issuer,
                issued_date=issued.isoformat(), expiry_date=exp.isoformat()
            ).model_dump())
    await db.certifications.insert_many([dict(x) for x in cert_docs])

    # Form Templates
    templates = [
        FormTemplate(
            name='Incident Report',
            category='incident',
            description='Report any workplace incident on a civil construction site',
            fields=[
                {'id': 'f1', 'label': 'Date & Time of Incident', 'type': 'date', 'required': True},
                {'id': 'f2', 'label': 'Type of Incident', 'type': 'select', 'required': True,
                 'options': ['Injury', 'Property Damage', 'Equipment Failure', 'Environmental Spill', 'Near Miss', 'Other']},
                {'id': 'f3', 'label': 'Description of Incident', 'type': 'textarea', 'required': True,
                 'placeholder': 'Describe what happened in detail...'},
                {'id': 'f4', 'label': 'Immediate Cause', 'type': 'textarea', 'required': False},
                {'id': 'f5', 'label': 'Body Part Affected (if injury)', 'type': 'select',
                 'options': ['N/A', 'Head', 'Eye', 'Hand', 'Arm', 'Back', 'Leg', 'Foot', 'Other']},
                {'id': 'f6', 'label': 'Witnesses', 'type': 'text'},
                {'id': 'f7', 'label': 'Corrective Actions Required', 'type': 'textarea'},
                {'id': 'f8', 'label': 'Photo Evidence', 'type': 'photo'},
                {'id': 'f9', 'label': 'Worker Signature', 'type': 'signature', 'required': True},
            ]
        ).model_dump(),
        FormTemplate(
            name='Daily Site Inspection',
            category='inspection',
            description='Pre-work safety inspection of the jobsite',
            fields=[
                {'id': 'f1', 'label': 'Inspection Date', 'type': 'date', 'required': True},
                {'id': 'f2', 'label': 'Weather Conditions', 'type': 'select',
                 'options': ['Clear', 'Cloudy', 'Rain', 'Snow', 'Windy', 'Fog']},
                {'id': 'f3', 'label': 'PPE Available & In Use?', 'type': 'radio', 'options': ['Yes', 'No', 'Partial']},
                {'id': 'f4', 'label': 'Excavations Properly Shored?', 'type': 'radio', 'options': ['Yes', 'No', 'N/A']},
                {'id': 'f5', 'label': 'Traffic Control in Place?', 'type': 'radio', 'options': ['Yes', 'No', 'N/A']},
                {'id': 'f6', 'label': 'Equipment Inspected?', 'type': 'radio', 'options': ['Yes', 'No']},
                {'id': 'f7', 'label': 'Hazards Identified', 'type': 'textarea'},
                {'id': 'f8', 'label': 'Site Photo', 'type': 'photo'},
                {'id': 'f9', 'label': 'Inspector Signature', 'type': 'signature', 'required': True},
            ]
        ).model_dump(),
        FormTemplate(
            name='Toolbox Talk',
            category='toolbox',
            description='Daily safety briefing with the crew',
            fields=[
                {'id': 'f1', 'label': 'Date', 'type': 'date', 'required': True},
                {'id': 'f2', 'label': 'Topic Discussed', 'type': 'text', 'required': True},
                {'id': 'f3', 'label': 'Key Points', 'type': 'textarea', 'required': True},
                {'id': 'f4', 'label': 'Number of Attendees', 'type': 'number'},
                {'id': 'f5', 'label': 'Questions / Concerns Raised', 'type': 'textarea'},
                {'id': 'f6', 'label': 'Foreman Signature', 'type': 'signature', 'required': True},
            ]
        ).model_dump(),
        FormTemplate(
            name='Near Miss Report',
            category='near_miss',
            description='Report a near miss to prevent future incidents',
            fields=[
                {'id': 'f1', 'label': 'Date & Time', 'type': 'date', 'required': True},
                {'id': 'f2', 'label': 'What Happened?', 'type': 'textarea', 'required': True},
                {'id': 'f3', 'label': 'Potential Severity', 'type': 'select',
                 'options': ['Low', 'Medium', 'High', 'Critical']},
                {'id': 'f4', 'label': 'Contributing Factors', 'type': 'textarea'},
                {'id': 'f5', 'label': 'Recommended Actions', 'type': 'textarea'},
                {'id': 'f6', 'label': 'Photo', 'type': 'photo'},
                {'id': 'f7', 'label': 'Signature', 'type': 'signature', 'required': True},
            ]
        ).model_dump(),
        FormTemplate(
            name='Equipment Pre-Use Checklist',
            category='inspection',
            description='Inspect heavy equipment before use',
            fields=[
                {'id': 'f1', 'label': 'Equipment Type', 'type': 'select',
                 'options': ['Excavator', 'Bulldozer', 'Crane', 'Loader', 'Dump Truck', 'Grader', 'Compactor']},
                {'id': 'f2', 'label': 'Equipment ID / Serial #', 'type': 'text', 'required': True},
                {'id': 'f3', 'label': 'Fluid Levels OK?', 'type': 'radio', 'options': ['Yes', 'No']},
                {'id': 'f4', 'label': 'Tires/Tracks Condition', 'type': 'radio', 'options': ['Good', 'Fair', 'Poor']},
                {'id': 'f5', 'label': 'Lights & Alarms Working?', 'type': 'radio', 'options': ['Yes', 'No']},
                {'id': 'f6', 'label': 'Defects Found', 'type': 'textarea'},
                {'id': 'f7', 'label': 'Operator Signature', 'type': 'signature', 'required': True},
            ]
        ).model_dump(),
    ]
    await db.form_templates.insert_many([dict(x) for x in templates])

    # Generate realistic submissions (~50)
    import random
    random.seed(42)
    incident_descriptions = [
        ('Worker slipped on muddy ground near excavation, sustained minor ankle sprain. Wet conditions from morning rain not properly cordoned off.', 'Injury'),
        ('Excavator bucket struck overhead utility line. Power outage to adjacent building. No injuries but significant property damage.', 'Property Damage'),
        ('Operator failed to perform pre-use brake check. Truck rolled 3 feet before being stopped. No damage.', 'Equipment Failure'),
        ('Hydraulic fluid leak from grader contaminated drainage ditch. Environmental cleanup required.', 'Environmental Spill'),
        ('Worker not wearing fall protection while installing rebar on elevated form. Caught by supervisor and corrected.', 'Near Miss'),
        ('Concrete splash to eye despite safety glasses (worn improperly). Flushed at site, sent for medical check.', 'Injury'),
    ]
    near_miss_descs = [
        'Steel beam swung dangerously close to ground crew during crane lift due to gust of wind.',
        'Worker stepped backward without looking, came within 2 feet of operating excavator swing radius.',
        'Trench collapse on unshored section just minutes after crew exited for break.',
        'Vehicle entered closed traffic zone after barricades were knocked over by wind overnight.',
    ]
    
    submissions = []
    for i in range(55):
        days_ago = random.randint(0, 180)
        sub_date = (datetime.now(timezone.utc) - timedelta(days=days_ago)).isoformat()
        cat_choice = random.choices(
            ['incident', 'near_miss', 'inspection', 'toolbox'],
            weights=[0.18, 0.20, 0.35, 0.27]
        )[0]
        tpl = next((t for t in templates if t['category'] == cat_choice), templates[0])
        loc = random.choice(location_docs)
        worker = random.choice(worker_docs)

        answers = {}
        if cat_choice == 'incident':
            desc, itype = random.choice(incident_descriptions)
            answers = {
                'Date & Time of Incident': sub_date[:10],
                'Type of Incident': itype,
                'Description of Incident': desc,
                'Immediate Cause': random.choice([
                    'Inadequate hazard assessment', 'Rushing to meet deadline',
                    'Equipment not properly maintained', 'Lack of training on new procedure'
                ]),
                'Body Part Affected (if injury)': random.choice(['N/A', 'Hand', 'Eye', 'Back', 'Foot']),
                'Corrective Actions Required': 'Retrain crew on procedure. Update site hazard assessment. Toolbox talk scheduled.',
            }
        elif cat_choice == 'near_miss':
            answers = {
                'Date & Time': sub_date[:10],
                'What Happened?': random.choice(near_miss_descs),
                'Potential Severity': random.choice(['Medium', 'High', 'Critical']),
                'Contributing Factors': 'Weather conditions; communication breakdown between operator and ground crew.',
                'Recommended Actions': 'Implement stop-work protocol for high winds. Add spotter requirement.'
            }
        elif cat_choice == 'inspection':
            answers = {
                'Inspection Date': sub_date[:10],
                'Weather Conditions': random.choice(['Clear', 'Cloudy', 'Rain']),
                'PPE Available & In Use?': random.choice(['Yes', 'Partial']),
                'Excavations Properly Shored?': random.choice(['Yes', 'N/A']),
                'Traffic Control in Place?': 'Yes',
                'Equipment Inspected?': 'Yes',
                'Hazards Identified': random.choice(['None', 'Loose debris near walkway', 'Standing water in trench'])
            }
        else:  # toolbox
            topics = ['Fall Protection', 'Trenching Safety', 'Crane Lift Plan', 'Heat Stress', 'PPE Compliance', 'Backup Alarms']
            answers = {
                'Date': sub_date[:10],
                'Topic Discussed': random.choice(topics),
                'Key Points': 'Reviewed proper use of equipment, hazard identification, emergency procedures.',
                'Number of Attendees': random.randint(5, 18),
            }
        sub = FormSubmission(
            template_id=tpl['id'],
            template_name=tpl['name'],
            category=cat_choice,
            location_id=loc['id'],
            location_name=loc['name'],
            worker_id=worker['id'],
            worker_name=worker['name'],
            answers=answers,
            gps_lat=43.6532 + random.uniform(-0.5, 0.5),
            gps_lng=-79.3832 + random.uniform(-0.5, 0.5),
            flagged=(cat_choice in ('incident', 'near_miss')),
            submitted_at=sub_date,
        ).model_dump()
        submissions.append(sub)
    await db.submissions.insert_many([dict(x) for x in submissions])

    return {'ok': True, 'message': 'Demo data seeded',
            'counts': {'workers': len(worker_docs), 'locations': len(location_docs),
                       'templates': len(templates), 'submissions': len(submissions),
                       'certifications': len(cert_docs)}}

# ============== AUTO-SHARE RULES ==============
class AutoShareRule(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    location_id: Optional[str] = None
    category: Optional[str] = None  # incident, near_miss, inspection, toolbox, general, or None=all
    emails: List[str] = []
    enabled: bool = True
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/share-rules')
async def list_share_rules(user=Depends(get_current_user)):
    items = await db.share_rules.find({}, {'_id': 0}).to_list(500)
    return items

@api_router.post('/share-rules')
async def create_share_rule(r: AutoShareRule, user=Depends(require_admin)):
    doc = r.model_dump()
    await db.share_rules.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/share-rules/{rid}')
async def delete_share_rule(rid: str, user=Depends(require_admin)):
    await db.share_rules.delete_one({'id': rid})
    return {'ok': True}

@api_router.get('/share-log')
async def list_share_log(user=Depends(get_current_user)):
    items = await db.share_log.find({}, {'_id': 0}).sort('sent_at', -1).to_list(500)
    return items

async def evaluate_share_rules(submission: dict):
    """Evaluate auto-share rules. Send real email via M365 if configured, else MOCK to log."""
    rules = await db.share_rules.find({'enabled': True}, {'_id': 0}).to_list(500)
    m365_cfg = await db.integrations.find_one({'id': 'm365'}, {'_id': 0})
    m365_ready = bool(m365_cfg and m365_cfg.get('verified'))

    for r in rules:
        loc_match = (not r.get('location_id')) or r.get('location_id') == submission.get('location_id')
        cat_match = (not r.get('category')) or r.get('category') == submission.get('category')
        if loc_match and cat_match and r.get('emails'):
            subject = f"[Paneltec Safety] {submission.get('template_name', 'Form')} — {submission.get('location_name', '')}"
            html = f"""
<h2 style="color:#0B0B0F">Paneltec Safety Portal</h2>
<p>A new <b>{submission.get('template_name')}</b> ({submission.get('category')}) has been submitted at <b>{submission.get('location_name')}</b> by {submission.get('worker_name')}.</p>
<table style="border-collapse:collapse;width:100%;margin-top:10px">
  {''.join(f'<tr><td style="border:1px solid #E5E7EB;padding:6px;background:#F9FAFB"><b>{k}</b></td><td style="border:1px solid #E5E7EB;padding:6px">{v}</td></tr>' for k, v in (submission.get('answers') or {}).items())}
</table>
<p style="margin-top:14px;color:#6B7280;font-size:12px">Submitted: {submission.get('submitted_at')}<br>Submission ID: {submission.get('id')}</p>
"""
            for email in r['emails']:
                status = 'mocked'
                response_detail = None
                if m365_ready:
                    try:
                        result = await _send_m365_email(SendEmailIn(to=[email], subject=subject, body=html))
                        status = 'sent' if result.get('ok') else 'failed'
                        response_detail = str(result.get('response'))[:300] if not result.get('ok') else None
                    except Exception as e:
                        status = 'failed'
                        response_detail = str(e)[:300]
                log = {
                    'id': str(uuid.uuid4()),
                    'rule_id': r['id'],
                    'submission_id': submission['id'],
                    'template_name': submission.get('template_name'),
                    'location_name': submission.get('location_name'),
                    'category': submission.get('category'),
                    'recipient': email,
                    'status': status,
                    'response': response_detail,
                    'sent_at': now_iso(),
                }
                await db.share_log.insert_one(log)

# ============== API TOKENS (Public API) ==============
class ApiToken(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    token: str = Field(default_factory=lambda: 'ptk_' + secrets.token_urlsafe(24))
    scopes: List[str] = ['read']  # read, write
    last_used: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/api-tokens')
async def list_tokens(user=Depends(require_admin)):
    items = await db.api_tokens.find({}, {'_id': 0}).to_list(200)
    return items

@api_router.post('/api-tokens')
async def create_token(t: ApiToken, user=Depends(require_admin)):
    doc = t.model_dump()
    await db.api_tokens.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/api-tokens/{tid}')
async def delete_token(tid: str, user=Depends(require_admin)):
    await db.api_tokens.delete_one({'id': tid})
    return {'ok': True}

async def verify_api_key(x_api_key: Optional[str] = Header(None)):
    if not x_api_key:
        raise HTTPException(401, 'Missing X-API-Key header')
    tok = await db.api_tokens.find_one({'token': x_api_key}, {'_id': 0})
    if not tok:
        raise HTTPException(401, 'Invalid API key')
    await db.api_tokens.update_one({'id': tok['id']}, {'$set': {'last_used': now_iso()}})
    return tok

# Public API endpoints (consumed by Zapier-like integrations)
@api_router.get('/public/submissions')
async def public_list_submissions(tok=Depends(verify_api_key)):
    items = await db.submissions.find({}, {'_id': 0}).sort('submitted_at', -1).to_list(500)
    return items

@api_router.get('/public/workers')
async def public_list_workers(tok=Depends(verify_api_key)):
    return await db.workers.find({}, {'_id': 0}).to_list(500)

@api_router.get('/public/certifications')
async def public_list_certs(tok=Depends(verify_api_key)):
    return await db.certifications.find({}, {'_id': 0}).to_list(2000)

@api_router.get('/public/locations')
async def public_list_locations(tok=Depends(verify_api_key)):
    return await db.locations.find({}, {'_id': 0}).to_list(500)

# ============== CHAT / NOTIFICATIONS ==============
class ChatMessage(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sender_id: str
    sender_name: str
    sender_role: str = 'worker'
    channel: str = 'general'  # general | broadcast | direct:<user_id>
    body: str
    location_id: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/chat/messages')
async def list_messages(channel: str = 'general', limit: int = 100, user=Depends(get_current_user)):
    items = await db.chat_messages.find({'channel': channel}, {'_id': 0}).sort('created_at', -1).limit(limit).to_list(limit)
    return list(reversed(items))

@api_router.post('/chat/messages')
async def post_message(body: dict, user=Depends(get_current_user)):
    msg = ChatMessage(
        sender_id=user['id'],
        sender_name=user.get('name', 'Unknown'),
        sender_role=user.get('role', 'worker'),
        channel=body.get('channel', 'general'),
        body=body.get('body', '').strip(),
        location_id=body.get('location_id'),
    ).model_dump()
    if not msg['body']:
        raise HTTPException(400, 'Empty message')
    await db.chat_messages.insert_one(msg)
    msg.pop('_id', None)
    # Add a notification fan-out
    await db.notifications.insert_one({
        'id': str(uuid.uuid4()),
        'type': 'chat',
        'title': f"New message from {msg['sender_name']}",
        'body': msg['body'][:120],
        'channel': msg['channel'],
        'created_at': now_iso(),
        'read': False,
    })
    return msg

@api_router.get('/notifications')
async def list_notifications(user=Depends(get_current_user)):
    items = await db.notifications.find({}, {'_id': 0}).sort('created_at', -1).limit(50).to_list(50)
    return items

@api_router.post('/notifications/mark-read')
async def mark_read(user=Depends(get_current_user)):
    await db.notifications.update_many({}, {'$set': {'read': True}})
    return {'ok': True}

# ============== PDF EXPORT ==============
@api_router.get('/submissions/{sid}/pdf')
async def submission_pdf(sid: str, authorization: Optional[str] = Header(None), token: Optional[str] = None):
    # Accept either Authorization header or ?token=... query (for download links)
    user = None
    auth_token = None
    if authorization and authorization.startswith('Bearer '):
        auth_token = authorization.split(' ', 1)[1]
    elif token:
        auth_token = token
    if not auth_token:
        raise HTTPException(401, 'Missing token')
    try:
        payload = jwt.decode(auth_token, JWT_SECRET, algorithms=[JWT_ALG])
        user = await db.users.find_one({'id': payload['sub']}, {'_id': 0, 'password': 0})
    except Exception:
        raise HTTPException(401, 'Invalid token')
    if not user:
        raise HTTPException(401, 'User not found')

    sub = await db.submissions.find_one({'id': sid}, {'_id': 0})
    if not sub:
        raise HTTPException(404, 'Submission not found')

    # Build PDF
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=0.6*inch, rightMargin=0.6*inch,
                            topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    PT_YELLOW = colors.HexColor('#FBBF24')
    PT_BLACK = colors.HexColor('#0B0B0F')

    title_style = ParagraphStyle('Title', parent=styles['Title'], textColor=PT_BLACK, fontSize=20, spaceAfter=4, alignment=0)
    sub_style = ParagraphStyle('Sub', parent=styles['Normal'], textColor=colors.HexColor('#6B7280'), fontSize=10, spaceAfter=12)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], textColor=PT_BLACK, fontSize=13, spaceBefore=10, spaceAfter=6)
    body = styles['BodyText']

    story = []
    # Branded header
    header_data = [[
        Paragraph('<b><font color="#FBBF24">PANELTEC</font></b><br/><font size=8 color="#6B7280">SAFETY PORTAL · CIVIL CONTRACTORS</font>', styles['Normal']),
        Paragraph(f'<para align=right><font size=9 color="#6B7280">Generated: {datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")}<br/>Document ID: {sub["id"][:8]}</font></para>', styles['Normal']),
    ]]
    header_tbl = Table(header_data, colWidths=[3.5*inch, 3.5*inch])
    header_tbl.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,-1), 2, PT_YELLOW),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(header_tbl)
    story.append(Spacer(1, 12))

    story.append(Paragraph(sub.get('template_name', 'Safety Form'), title_style))
    cat_label = sub.get('category', 'general').replace('_', ' ').title()
    story.append(Paragraph(f"Category: <b>{cat_label}</b>", sub_style))

    # Metadata table
    meta = [
        ['Worker:', sub.get('worker_name') or '—', 'Submitted:', (sub.get('submitted_at') or '')[:19].replace('T', ' ')],
        ['Job Site:', sub.get('location_name') or '—', 'GPS:', f"{sub.get('gps_lat'):.4f}, {sub.get('gps_lng'):.4f}" if sub.get('gps_lat') else '—'],
    ]
    meta_tbl = Table(meta, colWidths=[1.0*inch, 2.5*inch, 1.0*inch, 2.5*inch])
    meta_tbl.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F9FAFB')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#E5E7EB')),
        ('TEXTCOLOR', (0,0), (0,-1), colors.HexColor('#6B7280')),
        ('TEXTCOLOR', (2,0), (2,-1), colors.HexColor('#6B7280')),
        ('FONTNAME', (1,0), (1,-1), 'Helvetica-Bold'),
        ('FONTNAME', (3,0), (3,-1), 'Helvetica-Bold'),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(meta_tbl)

    # AI Summary
    if sub.get('ai_summary'):
        story.append(Paragraph('AI Safety Insights', h2))
        story.append(Paragraph(sub['ai_summary'].replace('\n', '<br/>'), body))

    # Responses
    story.append(Paragraph('Form Responses', h2))
    for k, v in (sub.get('answers') or {}).items():
        story.append(Paragraph(f"<b>{k}</b>", body))
        story.append(Paragraph(str(v), body))
        story.append(Spacer(1, 6))

    # Photos
    photos = sub.get('photos_b64') or []
    if photos:
        story.append(Paragraph('Photo Evidence', h2))
        for p in photos[:6]:
            try:
                if ',' in p:
                    p = p.split(',', 1)[1]
                img_bytes = base64.b64decode(p)
                img_buf = io.BytesIO(img_bytes)
                story.append(RLImage(img_buf, width=4.5*inch, height=3*inch, kind='proportional'))
                story.append(Spacer(1, 6))
            except Exception as e:
                logger.warning(f"PDF photo error: {e}")

    # Signature
    if sub.get('signature_b64'):
        story.append(Paragraph('Worker Signature', h2))
        try:
            sig = sub['signature_b64']
            if ',' in sig:
                sig = sig.split(',', 1)[1]
            sig_bytes = base64.b64decode(sig)
            story.append(RLImage(io.BytesIO(sig_bytes), width=3*inch, height=1*inch, kind='proportional'))
        except Exception as e:
            logger.warning(f"Sig render error: {e}")

    story.append(Spacer(1, 24))
    story.append(Paragraph(f'<para align=center><font size=8 color="#9CA3AF">© Paneltec Civil Contractors · This is an electronically signed record · {sub["id"]}</font></para>', styles['Normal']))

    doc.build(story)
    buf.seek(0)
    filename = f"paneltec-{sub.get('category','form')}-{sub['id'][:8]}.pdf"
    return StreamingResponse(buf, media_type='application/pdf',
                             headers={'Content-Disposition': f'inline; filename="{filename}"'})



# ============== RISK REGISTER ==============
RISK_RATING_LABELS = {
    1: 'Low', 2: 'Low', 3: 'Moderate', 4: 'Moderate',
    6: 'High', 8: 'High', 9: 'Critical', 12: 'Critical', 16: 'Critical', 25: 'Critical'
}

def risk_label(score: int) -> str:
    if score <= 2: return 'Low'
    if score <= 4: return 'Moderate'
    if score <= 8: return 'High'
    return 'Critical'

class Risk(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    location_id: Optional[str] = None
    hazard_category: str  # Gravity/fall, ICT, Administrative, Lighting, Biological, Chemical
    hazard: str
    risk_details: str
    likelihood: int = 3  # 1..5
    consequence: int = 3  # 1..5
    current_controls: List[str] = []
    residual_likelihood: int = 2
    residual_consequence: int = 2
    new_controls: List[str] = []
    owner_id: Optional[str] = None
    status: str = 'open'  # open, mitigated, closed
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/risks')
async def list_risks(user=Depends(get_current_user)):
    items = await db.risks.find({}, {'_id': 0}).sort('created_at', -1).to_list(2000)
    return items

@api_router.post('/risks')
async def create_risk(r: Risk, user=Depends(get_current_user)):
    doc = r.model_dump()
    await db.risks.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/risks/{rid}')
async def update_risk(rid: str, r: Risk, user=Depends(get_current_user)):
    doc = r.model_dump(); doc['id'] = rid
    await db.risks.update_one({'id': rid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/risks/{rid}')
async def delete_risk(rid: str, user=Depends(require_admin)):
    await db.risks.delete_one({'id': rid})
    return {'ok': True}

# ============== ACTION PLANS ==============
class ActionItem(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: Optional[str] = None
    assignee_id: Optional[str] = None
    assignee_name: Optional[str] = None
    submission_id: Optional[str] = None
    risk_id: Optional[str] = None
    location_id: Optional[str] = None
    priority: str = 'medium'  # low, medium, high, critical
    status: str = 'open'  # open, in_progress, done, overdue
    due_date: Optional[str] = None
    completed_at: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/actions')
async def list_actions(user=Depends(get_current_user)):
    items = await db.actions.find({}, {'_id': 0}).sort('due_date', 1).to_list(2000)
    return items

@api_router.post('/actions')
async def create_action(a: ActionItem, user=Depends(get_current_user)):
    doc = a.model_dump()
    await db.actions.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/actions/{aid}')
async def update_action(aid: str, a: ActionItem, user=Depends(get_current_user)):
    doc = a.model_dump(); doc['id'] = aid
    await db.actions.update_one({'id': aid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/actions/{aid}')
async def delete_action(aid: str, user=Depends(get_current_user)):
    await db.actions.delete_one({'id': aid})
    return {'ok': True}

# ============== CHEMICALS / SDS ==============
class Chemical(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    manufacturer: Optional[str] = None
    cas_number: Optional[str] = None
    hazard_class: str = 'general'  # flammable, corrosive, toxic, oxidizer, explosive, general
    sds_url: Optional[str] = None
    sds_b64: Optional[str] = None  # base64 PDF content
    location_id: Optional[str] = None
    quantity: Optional[str] = None
    notes: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/chemicals')
async def list_chemicals(user=Depends(get_current_user)):
    items = await db.chemicals.find({}, {'_id': 0}).to_list(2000)
    return items

@api_router.post('/chemicals')
async def create_chemical(c: Chemical, user=Depends(get_current_user)):
    doc = c.model_dump()
    await db.chemicals.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/chemicals/{cid}')
async def delete_chemical(cid: str, user=Depends(require_admin)):
    await db.chemicals.delete_one({'id': cid})
    return {'ok': True}

# ============== ASSETS ==============
class Asset(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    asset_type: str = 'equipment'  # equipment, vehicle, tool, ppe
    serial_number: Optional[str] = None
    location_id: Optional[str] = None
    assigned_worker_id: Optional[str] = None
    status: str = 'active'  # active, maintenance, retired
    next_inspection: Optional[str] = None
    last_inspection: Optional[str] = None
    notes: Optional[str] = None
    photo_b64: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/assets')
async def list_assets(user=Depends(get_current_user)):
    return await db.assets.find({}, {'_id': 0}).to_list(2000)

@api_router.post('/assets')
async def create_asset(a: Asset, user=Depends(get_current_user)):
    doc = a.model_dump()
    await db.assets.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/assets/{aid}')
async def delete_asset(aid: str, user=Depends(require_admin)):
    await db.assets.delete_one({'id': aid})
    return {'ok': True}

# ============== CONTRACTORS ==============
class Contractor(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    contact_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    trade: Optional[str] = None
    insurance_expiry: Optional[str] = None
    prequalified: bool = False
    workers_count: int = 0
    notes: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/contractors')
async def list_contractors(user=Depends(get_current_user)):
    return await db.contractors.find({}, {'_id': 0}).to_list(2000)

@api_router.post('/contractors')
async def create_contractor(c: Contractor, user=Depends(get_current_user)):
    doc = c.model_dump()
    await db.contractors.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/contractors/{cid}')
async def delete_contractor(cid: str, user=Depends(require_admin)):
    await db.contractors.delete_one({'id': cid})
    return {'ok': True}

# ============== HAZARD QUICK REPORT ==============
class HazardReport(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: str
    severity: str = 'medium'  # low, medium, high, critical
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    reporter_id: Optional[str] = None
    reporter_name: Optional[str] = None
    photo_b64: Optional[str] = None
    gps_lat: Optional[float] = None
    gps_lng: Optional[float] = None
    status: str = 'open'
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/hazards')
async def list_hazards(user=Depends(get_current_user)):
    return await db.hazards.find({}, {'_id': 0}).sort('created_at', -1).to_list(500)

@api_router.post('/hazards')
async def create_hazard(h: HazardReport, user=Depends(get_current_user)):
    doc = h.model_dump()
    doc['reporter_id'] = user['id']
    doc['reporter_name'] = user.get('name')
    await db.hazards.insert_one(doc)
    # Auto-create action for high/critical
    if doc.get('severity') in ('high', 'critical'):
        action = ActionItem(
            title=f"Address hazard: {doc['title']}",
            description=doc['description'],
            priority=doc['severity'],
            location_id=doc.get('location_id'),
            due_date=(datetime.now(timezone.utc) + timedelta(days=3 if doc['severity']=='critical' else 7)).date().isoformat(),
        ).model_dump()
        await db.actions.insert_one(action)
    doc.pop('_id', None)
    return doc

# ============== EXTEND SEED ==============
@api_router.post('/seed-extras')
async def seed_extras():
    """Seed risk register, chemicals, assets, contractors, actions for demo."""
    import random
    random.seed(7)
    await db.risks.delete_many({})
    await db.actions.delete_many({})
    await db.chemicals.delete_many({})
    await db.assets.delete_many({})
    await db.contractors.delete_many({})
    await db.hazards.delete_many({})

    locations = await db.locations.find({}, {'_id': 0}).to_list(100)
    workers = await db.workers.find({}, {'_id': 0}).to_list(100)

    # RISKS — civil-contractor specific
    risk_seed = [
        ('Gravity / Fall', 'Fall from height — scaffold without barrier on slab edge', 4, 5, ['Toe boards', 'PPE harness'], ['Install permanent guardrails', 'Daily inspection'], 2, 3),
        ('Excavation', 'Trench collapse — unshored excavation > 1.2m', 4, 5, ['Sloping at 1:1'], ['Hydraulic shoring boxes', 'Competent person daily check'], 1, 3),
        ('ICT', 'Loss of internet — no access to safe work procedures on jobsite', 3, 3, ['Paper backup of SWPs'], ['Offline mobile sync'], 1, 2),
        ('Administrative', 'Hearing damage — exposure to compactor / jackhammer noise', 4, 3, ['Hearing protection issued'], ['Mandatory double-hearing protection above 85dB'], 2, 2),
        ('Lighting', 'Insufficient lighting in stairwell of partially built structure', 3, 2, ['Temporary LED strings'], ['Permanent lighting install'], 1, 2),
        ('Biological', 'Possibility of exposure to flu/seasonal illness on crowded crew bus', 3, 2, ['Hand sanitizer at site entry'], ['Stagger crew shifts'], 2, 2),
        ('Plant / Equipment', 'Excavator swing-radius struck-by hazard for ground crew', 4, 5, ['Banksman / spotter'], ['Spatial awareness alarms', 'Exclusion zone tape'], 2, 4),
        ('Chemical', 'Exposure to silica dust during concrete cutting', 4, 4, ['Wet cutting', 'P2 respirators'], ['Engineered local exhaust ventilation'], 2, 3),
        ('Chemical', 'Asbestos exposure during demolition of legacy structures', 3, 5, ['Pre-demo asbestos survey'], ['Licensed asbestos removal contractor'], 1, 4),
        ('Manual Handling', 'Back strain — lifting heavy formwork timbers manually', 4, 3, ['2-person lift policy'], ['Mechanical lifting aid procurement'], 3, 2),
        ('Traffic', 'Public vehicle intrusion into work zone on Highway 401 expansion', 3, 5, ['TCP / cones / signage'], ['Concrete barriers', 'Variable message signs'], 2, 4),
        ('Electrical', 'Contact with overhead power lines during crane operations', 2, 5, ['Spotter / minimum distance'], ['De-energize lines via utility coordination'], 1, 4),
    ]
    risk_docs = []
    for hc, det, l, c, cur, new, rl, rc in risk_seed:
        r = Risk(
            location_id=random.choice(locations)['id'] if locations else None,
            hazard_category=hc,
            hazard=det.split(' — ')[0],
            risk_details=det,
            likelihood=l, consequence=c,
            current_controls=cur,
            new_controls=new,
            residual_likelihood=rl, residual_consequence=rc,
            owner_id=random.choice(workers)['id'] if workers else None,
        ).model_dump()
        risk_docs.append(r)
    await db.risks.insert_many([dict(x) for x in risk_docs])

    # ACTIONS
    action_seed = [
        ('Install permanent guardrails — Level 3 slab', 'high', 5),
        ('Procure hydraulic shoring boxes', 'critical', 2),
        ('Schedule licensed asbestos removal vendor', 'critical', 1),
        ('Roll out P2 fit-testing for concrete cutters', 'high', 14),
        ('Coordinate utility de-energization for crane lifts', 'high', 7),
        ('Toolbox talk: hearing protection above 85dB', 'medium', 3),
        ('Replace damaged barricades on Highway 401 zone', 'high', 4),
        ('Refresh first-aid kits at all 4 sites', 'medium', 10),
        ('Investigate near-miss: ground crew swing-radius', 'critical', -2),  # overdue
        ('Update lighting in Downtown Bridge stairwell', 'medium', 21),
    ]
    today = datetime.now(timezone.utc)
    action_docs = []
    for title, pri, days in action_seed:
        due = (today + timedelta(days=days)).date().isoformat()
        st = 'overdue' if days < 0 else random.choice(['open', 'open', 'in_progress'])
        w = random.choice(workers) if workers else None
        a = ActionItem(
            title=title,
            description=f"Follow-up corrective action — assigned during risk review.",
            assignee_id=w['id'] if w else None,
            assignee_name=w['name'] if w else None,
            priority=pri,
            status=st,
            due_date=due,
            location_id=random.choice(locations)['id'] if locations else None,
        ).model_dump()
        action_docs.append(a)
    await db.actions.insert_many([dict(x) for x in action_docs])

    # CHEMICALS / SDS
    chemicals = [
        ('Diesel Fuel #2', 'Petro-Canada', 'flammable', '68476-34-6', '500 L drum'),
        ('Portland Cement Type GU', 'Lafarge', 'general', '65997-15-1', '50 bags'),
        ('Concrete Sealer Solvent', 'Sika', 'flammable', '64742-95-6', '20 L'),
        ('Hydraulic Oil ISO 46', 'Shell Tellus', 'general', None, '200 L drum'),
        ('Acetylene Gas', 'Air Liquide', 'explosive', '74-86-2', '4 cylinders'),
        ('Oxygen Gas', 'Air Liquide', 'oxidizer', '7782-44-7', '4 cylinders'),
        ('Hydrochloric Acid 10%', 'Acme Chem', 'corrosive', '7647-01-0', '5 L'),
        ('Form Release Agent', 'BASF', 'general', None, '20 L'),
        ('Spray Paint - Marking', 'Krylon', 'flammable', None, '24 cans'),
        ('Sodium Hypochlorite (Bleach)', 'Univar', 'corrosive', '7681-52-9', '10 L'),
    ]
    chem_docs = []
    for n, m, hc, cas, qty in chemicals:
        c = Chemical(
            name=n, manufacturer=m, hazard_class=hc, cas_number=cas, quantity=qty,
            location_id=random.choice(locations)['id'] if locations else None,
            notes='SDS reviewed and on-site.'
        ).model_dump()
        chem_docs.append(c)
    await db.chemicals.insert_many([dict(x) for x in chem_docs])

    # ASSETS
    asset_seed = [
        ('CAT 320 Excavator', 'equipment', 'CAT320-2023-001', 14),
        ('Komatsu D65 Bulldozer', 'equipment', 'KOM-D65-022', 30),
        ('Manitowoc 18000 Crane', 'equipment', 'MTC-18K-08', 5),
        ('Ford F-550 Service Truck', 'vehicle', 'F550-PT-12', 60),
        ('Wacker DPU 6555 Compactor', 'equipment', 'WK-6555-04', -3),  # overdue
        ('Genie Z-45 Boom Lift', 'equipment', 'GZ45-19', 21),
        ('Stihl Concrete Saw', 'tool', 'STL-TS420-07', 45),
        ('Total Station Surveying Kit', 'tool', 'TS-LEICA-01', 90),
    ]
    asset_docs = []
    for n, t, sn, days in asset_seed:
        nxt = (today + timedelta(days=days)).date().isoformat()
        a = Asset(
            name=n, asset_type=t, serial_number=sn,
            location_id=random.choice(locations)['id'] if locations else None,
            next_inspection=nxt,
            last_inspection=(today - timedelta(days=180)).date().isoformat(),
            status='active' if days > 0 else 'maintenance',
        ).model_dump()
        asset_docs.append(a)
    await db.assets.insert_many([dict(x) for x in asset_docs])

    # CONTRACTORS
    contractor_seed = [
        ('Northern Steel Erectors Inc', 'Mark Lawson', 'mark@nse.ca', '416-555-1010', 'Structural Steel', True, 45, 200),
        ('Apex Excavation Ltd', 'Diane Walters', 'd.walters@apex.ca', '905-555-2222', 'Earthworks', True, 28, 90),
        ('Stellar Electrical Co', 'Raj Singh', 'raj@stellarelec.ca', '647-555-3333', 'Electrical', True, 12, -10),  # expired
        ('CleanSite Asbestos Removal', 'Hannah Patel', 'hannah@cleansite.ca', '416-555-4444', 'Hazmat / Asbestos', True, 6, 365),
        ('GreenScape Landscaping', 'Tom Chen', 'tom@greenscape.ca', '905-555-5555', 'Landscaping', False, 8, -45),  # not prequalified
    ]
    contractor_docs = []
    for company, contact, email, phone, trade, prequal, count, days in contractor_seed:
        exp = (today + timedelta(days=days)).date().isoformat()
        c = Contractor(
            company_name=company, contact_name=contact, email=email, phone=phone,
            trade=trade, prequalified=prequal, workers_count=count, insurance_expiry=exp,
            notes='Active on multiple jobsites.' if prequal else 'Pending prequalification review.',
        ).model_dump()
        contractor_docs.append(c)
    await db.contractors.insert_many([dict(x) for x in contractor_docs])

    return {
        'ok': True,
        'counts': {
            'risks': len(risk_docs), 'actions': len(action_docs),
            'chemicals': len(chem_docs), 'assets': len(asset_docs),
            'contractors': len(contractor_docs),
        }
    }

# ============== SWMS (Safe Work Method Statements) ==============
class SWMSStep(BaseModel):
    step_no: float = 1.0
    activity: str
    hazards: List[str] = []
    risk_class: int = 3  # 1-critical, 2-high, 3-moderate, 4-low
    controls: List[str] = []
    responsible: List[str] = []
    residual_risk: int = 4

class SWMS(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    high_risk_activity: str  # Mechanical Excavation, Working at Heights, Hot Work, etc.
    project_name: Optional[str] = None
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    supervisor_name: Optional[str] = None
    prepared_by: Optional[str] = None
    description: Optional[str] = None
    ppe_required: List[str] = []
    plant_equipment: List[str] = []
    legislation: List[str] = []
    training_required: List[str] = []
    steps: List[Dict[str, Any]] = []
    signoffs: List[Dict[str, Any]] = []  # [{name, signature_b64, company, date}]
    status: str = 'draft'  # draft, approved, active, retired
    valid_from: Optional[str] = None
    valid_until: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/swms')
async def list_swms(user=Depends(get_current_user)):
    return await db.swms.find({}, {'_id': 0}).sort('created_at', -1).to_list(500)

@api_router.get('/swms/{sid}')
async def get_swms(sid: str, user=Depends(get_current_user)):
    s = await db.swms.find_one({'id': sid}, {'_id': 0})
    if not s: raise HTTPException(404, 'Not found')
    return s

@api_router.post('/swms')
async def create_swms(s: SWMS, user=Depends(get_current_user)):
    doc = s.model_dump()
    await db.swms.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/swms/{sid}')
async def update_swms(sid: str, s: SWMS, user=Depends(get_current_user)):
    doc = s.model_dump(); doc['id'] = sid
    await db.swms.update_one({'id': sid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/swms/{sid}')
async def delete_swms(sid: str, user=Depends(require_admin)):
    await db.swms.delete_one({'id': sid})
    return {'ok': True}

@api_router.post('/swms/{sid}/signoff')
async def signoff_swms(sid: str, body: dict, user=Depends(get_current_user)):
    signoff = {
        'name': body.get('name', user.get('name')),
        'signature_b64': body.get('signature_b64'),
        'company': body.get('company', 'Paneltec'),
        'date': now_iso(),
        'user_id': user['id'],
    }
    await db.swms.update_one({'id': sid}, {'$push': {'signoffs': signoff}})
    return signoff

class AISwmsIn(BaseModel):
    title: str
    activity: str
    description: Optional[str] = None
    location_name: Optional[str] = None

@api_router.post('/ai/swms-draft')
async def ai_swms_draft(body: AISwmsIn, user=Depends(get_current_user)):
    """AI generates a draft SWMS from a task description."""
    prompt = f"""Generate a Safe Work Method Statement (SWMS) for a civil contractor.

Title: {body.title}
High Risk Activity: {body.activity}
Location: {body.location_name or 'Construction site'}
Description: {body.description or ''}

Return a valid JSON object with this exact structure (no markdown, no commentary, just JSON):
{{
  "ppe_required": ["item1", "item2", ...],
  "plant_equipment": ["item1", ...],
  "legislation": ["WHS Act 2011", "Code of Practice ..."],
  "training_required": ["White Card", "..."],
  "steps": [
    {{
      "step_no": 1.0,
      "activity": "Setup and pre-start checks",
      "hazards": ["Hazard 1", "Hazard 2"],
      "risk_class": 2,
      "controls": ["Control 1", "Control 2"],
      "responsible": ["Supervisor", "Operator"],
      "residual_risk": 4
    }},
    ...
  ]
}}

risk_class scale: 1=Critical, 2=High, 3=Moderate, 4=Low.
Provide 5-8 sequential steps covering setup, execution, and pack-down. Be specific to civil contractor work."""

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"swms-{uuid.uuid4()}",
            system_message="You are an Australian/Canadian WHS expert helping a civil contractor draft Safe Work Method Statements. Always return only valid JSON, no commentary, no markdown code fences."
        ).with_model('openai', 'gpt-4o-mini')
        result = await chat.send_message(UserMessage(text=prompt))
        text = str(result).strip()
        if text.startswith('```'):
            text = text.split('```', 2)[1]
            if text.startswith('json'): text = text[4:]
            text = text.strip()
        import json
        data = json.loads(text)
        return data
    except Exception as e:
        logger.error(f"AI SWMS error: {e}")
        # Fallback skeleton
        return {
            'ppe_required': ['Hard Hat', 'Safety Boots', 'High-Vis Vest', 'Safety Glasses', 'Gloves'],
            'plant_equipment': ['As required by activity'],
            'legislation': ['WHS Act', 'Code of Practice: Construction Work'],
            'training_required': ['Construction Induction Card (White Card)'],
            'steps': [
                {'step_no': 1.0, 'activity': 'Site setup and pre-start inspection', 'hazards': ['Slips, trips, falls'], 'risk_class': 3, 'controls': ['Walk site', 'Toolbox talk'], 'responsible': ['Supervisor'], 'residual_risk': 4},
                {'step_no': 2.0, 'activity': f'Execute {body.activity}', 'hazards': ['Activity-specific'], 'risk_class': 2, 'controls': ['Follow SWMS', 'PPE'], 'responsible': ['Operators'], 'residual_risk': 3},
                {'step_no': 3.0, 'activity': 'Pack-down and housekeeping', 'hazards': ['Manual handling'], 'risk_class': 3, 'controls': ['Team lift'], 'responsible': ['Crew'], 'residual_risk': 4},
            ],
            '_fallback': True,
        }

# ============== ITP (Inspection & Test Plans) ==============
class ITPItem(BaseModel):
    item_no: float = 1.0
    description: str
    reference: Optional[str] = None  # spec / standard
    inspection_type: str = 'check'  # check, hold_point, witness_point, survey, test
    frequency: Optional[str] = None
    acceptance_criteria: Optional[str] = None
    operations_by: Optional[str] = 'Paneltec'
    verification_by: Optional[str] = 'Client'
    status: str = 'pending'  # pending, in_progress, passed, failed, n_a
    completed_at: Optional[str] = None
    notes: Optional[str] = None

class ITP(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    project_name: Optional[str] = None
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    discipline: str = 'civil'  # civil, structural, mechanical, electrical
    description: Optional[str] = None
    items: List[Dict[str, Any]] = []
    status: str = 'open'
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/itps')
async def list_itps(user=Depends(get_current_user)):
    return await db.itps.find({}, {'_id': 0}).sort('created_at', -1).to_list(500)

@api_router.get('/itps/{iid}')
async def get_itp(iid: str, user=Depends(get_current_user)):
    i = await db.itps.find_one({'id': iid}, {'_id': 0})
    if not i: raise HTTPException(404, 'Not found')
    return i

@api_router.post('/itps')
async def create_itp(i: ITP, user=Depends(get_current_user)):
    doc = i.model_dump()
    await db.itps.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/itps/{iid}')
async def update_itp(iid: str, i: ITP, user=Depends(get_current_user)):
    doc = i.model_dump(); doc['id'] = iid
    await db.itps.update_one({'id': iid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/itps/{iid}')
async def delete_itp(iid: str, user=Depends(require_admin)):
    await db.itps.delete_one({'id': iid})
    return {'ok': True}

# ============== PERMITS TO WORK ==============
class Permit(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    permit_type: str = 'hot_work'  # hot_work, confined_space, excavation, working_at_heights, electrical_isolation
    project_name: Optional[str] = None
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    description: str
    valid_from: str
    valid_until: str
    issued_by: Optional[str] = None
    issued_to: Optional[str] = None
    precautions: List[str] = []
    checklist: List[Dict[str, Any]] = []  # [{label, checked}]
    status: str = 'active'  # draft, active, expired, closed, cancelled
    signoffs: List[Dict[str, Any]] = []
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/permits')
async def list_permits(user=Depends(get_current_user)):
    items = await db.permits.find({}, {'_id': 0}).sort('valid_from', -1).to_list(500)
    # Auto-expire
    now = datetime.now(timezone.utc).isoformat()
    for p in items:
        if p.get('status') == 'active' and p.get('valid_until', '') < now:
            await db.permits.update_one({'id': p['id']}, {'$set': {'status': 'expired'}})
            p['status'] = 'expired'
    return items

@api_router.post('/permits')
async def create_permit(p: Permit, user=Depends(get_current_user)):
    doc = p.model_dump()
    await db.permits.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/permits/{pid}')
async def update_permit(pid: str, p: Permit, user=Depends(get_current_user)):
    doc = p.model_dump(); doc['id'] = pid
    await db.permits.update_one({'id': pid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/permits/{pid}')
async def delete_permit(pid: str, user=Depends(require_admin)):
    await db.permits.delete_one({'id': pid})
    return {'ok': True}

# ============== ENVIRONMENTAL ASPECTS ==============
class EnvAspect(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    aspect_code: str  # EM-01..EM-09
    category: str  # air_quality, noise, soil_water, flora_fauna, heritage, fuels_chemicals, community, waste, biodiversity
    title: str
    description: Optional[str] = None
    control_measures: List[str] = []
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    monitoring_frequency: Optional[str] = None
    responsible: Optional[str] = None
    status: str = 'active'
    created_at: str = Field(default_factory=now_iso)

class EnvMonitoringLog(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    aspect_id: str
    aspect_title: Optional[str] = None
    reading: str  # e.g. "75 dB", "pH 7.2"
    threshold: Optional[str] = None
    status: str = 'within_limits'  # within_limits, exceeded, action_required
    notes: Optional[str] = None
    photo_b64: Optional[str] = None
    recorded_by: Optional[str] = None
    recorded_at: str = Field(default_factory=now_iso)

@api_router.get('/env/aspects')
async def list_env_aspects(user=Depends(get_current_user)):
    return await db.env_aspects.find({}, {'_id': 0}).to_list(500)

@api_router.post('/env/aspects')
async def create_env_aspect(a: EnvAspect, user=Depends(get_current_user)):
    doc = a.model_dump()
    await db.env_aspects.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/env/aspects/{aid}')
async def delete_env_aspect(aid: str, user=Depends(require_admin)):
    await db.env_aspects.delete_one({'id': aid})
    return {'ok': True}

@api_router.get('/env/logs')
async def list_env_logs(user=Depends(get_current_user)):
    return await db.env_logs.find({}, {'_id': 0}).sort('recorded_at', -1).to_list(2000)

@api_router.post('/env/logs')
async def create_env_log(l: EnvMonitoringLog, user=Depends(get_current_user)):
    doc = l.model_dump()
    doc['recorded_by'] = user.get('name')
    await db.env_logs.insert_one(doc)
    doc.pop('_id', None)
    return doc

# ============== INSURANCE POLICIES ==============
class InsurancePolicy(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    contractor_id: Optional[str] = None  # if linked to contractor; if null = own/Paneltec
    insurance_type: str  # combined_business, employers_liability, motor_vehicle, public_liability, professional_indemnity
    company: str
    policy_number: str
    coverage_amount: Optional[str] = None
    issued_date: Optional[str] = None
    expiry_date: str
    document_b64: Optional[str] = None
    notes: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/insurance')
async def list_insurance(user=Depends(get_current_user)):
    return await db.insurance.find({}, {'_id': 0}).sort('expiry_date', 1).to_list(500)

@api_router.post('/insurance')
async def create_insurance(p: InsurancePolicy, user=Depends(get_current_user)):
    doc = p.model_dump()
    await db.insurance.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/insurance/{pid}')
async def delete_insurance(pid: str, user=Depends(require_admin)):
    await db.insurance.delete_one({'id': pid})
    return {'ok': True}

# ============== AUDIT REGISTER ==============
class AuditFinding(BaseModel):
    description: str
    severity: str = 'observation'  # major_nc, minor_nc, observation, opportunity
    corrective_action: Optional[str] = None
    closed: bool = False

class Audit(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    audit_type: str = 'internal'  # internal, external, client, regulatory
    scope: Optional[str] = None
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    auditor: Optional[str] = None
    planned_date: Optional[str] = None
    actual_date: Optional[str] = None
    status: str = 'planned'  # planned, in_progress, completed, cancelled
    findings: List[Dict[str, Any]] = []
    report_b64: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/audits')
async def list_audits(user=Depends(get_current_user)):
    return await db.audits.find({}, {'_id': 0}).sort('planned_date', -1).to_list(500)

@api_router.post('/audits')
async def create_audit(a: Audit, user=Depends(get_current_user)):
    doc = a.model_dump()
    await db.audits.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/audits/{aid}')
async def update_audit(aid: str, a: Audit, user=Depends(get_current_user)):
    doc = a.model_dump(); doc['id'] = aid
    await db.audits.update_one({'id': aid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/audits/{aid}')
async def delete_audit(aid: str, user=Depends(require_admin)):
    await db.audits.delete_one({'id': aid})
    return {'ok': True}

# ============== SEED COMPLIANCE EXTRAS (SWMS/ITP/Permits/Env/Insurance/Audits + TasWater project) ==============
@api_router.post('/seed-compliance')
async def seed_compliance():
    import random
    random.seed(11)
    await db.swms.delete_many({})
    await db.itps.delete_many({})
    await db.permits.delete_many({})
    await db.env_aspects.delete_many({})
    await db.env_logs.delete_many({})
    await db.insurance.delete_many({})
    await db.audits.delete_many({})

    # Add TasWater realistic location if not present
    locations = await db.locations.find({}, {'_id': 0}).to_list(100)
    taswater = next((l for l in locations if 'Tasman' in l.get('name','')), None)
    if not taswater:
        taswater = Location(
            name='Tasman Hwy Scottsdale — WTP Sledge Track',
            address='35480-35530 Tasman Highway, Scottsdale TAS',
            project_code='PT-TW-WTP-001',
        ).model_dump()
        await db.locations.insert_one(dict(taswater))
        locations.append(taswater)

    workers = await db.workers.find({}, {'_id': 0}).to_list(100)
    today = datetime.now(timezone.utc)

    # SWMS — TasWater project examples
    swms_seed = [
        {
            'title': 'SWMS-003 — Water Main Construction & Traffic Management',
            'high_risk_activity': 'Mechanical Excavation in Road Reserve',
            'project_name': 'TasWater Tasman Hwy Scottsdale',
            'description': 'Replacement of 605m DN200 AC raw water main in and adjacent to road reserve. Covers excavation, asbestos pipe removal, traffic management.',
            'ppe_required': ['Hard Hat', 'Hi-Vis Vest (white reflective for night)', 'Safety Boots', 'Safety Glasses', 'Hearing Protection', 'P2 Dust Masks', 'Razor Shield Gloves'],
            'plant_equipment': ['Vacuum Truck', 'Excavator CAT 320', 'HDD Drill Rig', 'Dump Truck', 'Traffic Cones / Bollards'],
            'legislation': ['WHS Act 2012 (Tas)', 'WHS Regulations 2012', 'AS1742.3 Traffic Control', 'TasWater Supplement WSAA Code', 'Construction Work CoP 2013'],
            'training_required': ['White Card Construction Induction', 'TasWater Induction', 'Traffic Control Ticket', 'Class B Asbestos Awareness', 'Confined Space (if applicable)'],
            'steps': [
                {'step_no': 1.0, 'activity': 'General safety requirements & site induction', 'hazards': ['Personal injury to operator/workers'], 'risk_class': 2, 'controls': ['Current TasWater induction', 'Site-specific JSEA daily review', 'PPE per matrix'], 'responsible': ['Compliance Manager', 'Supervisor'], 'residual_risk': 4},
                {'step_no': 2.0, 'activity': 'Setting out / taking down roadwork signage', 'hazards': ['Struck by oncoming traffic'], 'risk_class': 1, 'controls': ['Verify signage per TGS', 'Hi-vis at all times', 'Traffic controllers at both ends'], 'responsible': ['Traffic Controllers'], 'residual_risk': 3},
                {'step_no': 3.0, 'activity': 'Underground asset locating', 'hazards': ['Strike on services', 'Contact electrical lines'], 'risk_class': 1, 'controls': ['DBYD lookup', 'Vacuum excavation for proving', 'Trained locator'], 'responsible': ['Supervisor', 'Asset Locator'], 'residual_risk': 3},
                {'step_no': 4.0, 'activity': 'Mechanical excavation of trench', 'hazards': ['Trench collapse', 'Falling into trench', 'Plant rollover'], 'risk_class': 1, 'controls': ['Shoring boxes >1.2m', 'Trench barriers', 'Spotter for plant', 'Daily competent person inspection'], 'responsible': ['Supervisor', 'Excavator Operator'], 'residual_risk': 3},
                {'step_no': 5.0, 'activity': 'Class B asbestos AC pipe removal', 'hazards': ['Asbestos fibre exposure'], 'risk_class': 1, 'controls': ['Wet methods', 'P2 respirators', 'Licensed removalist', 'Air monitoring', 'Decontamination'], 'responsible': ['Asbestos Supervisor'], 'residual_risk': 3},
                {'step_no': 6.0, 'activity': 'Pipe laying & connection', 'hazards': ['Manual handling', 'Crush injury'], 'risk_class': 2, 'controls': ['Mechanical lifting aids', 'Pipe slings', 'Clear communication'], 'responsible': ['Crew', 'Crane Operator'], 'residual_risk': 4},
                {'step_no': 7.0, 'activity': 'Backfill & compaction', 'hazards': ['Compactor injury', 'Noise > 85dB'], 'risk_class': 2, 'controls': ['Double hearing protection', 'Compactor pre-start'], 'responsible': ['Operator'], 'residual_risk': 4},
                {'step_no': 8.0, 'activity': 'Pack-down & site reinstatement', 'hazards': ['Residual road hazards'], 'risk_class': 3, 'controls': ['Sweep roadway', 'Remove signage in correct order'], 'responsible': ['Supervisor'], 'residual_risk': 4},
            ],
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'supervisor_name': 'Mathew Loone',
            'prepared_by': 'Patrick Monaghan',
            'status': 'active',
            'valid_from': (today - timedelta(days=14)).isoformat(),
            'valid_until': (today + timedelta(days=90)).isoformat(),
        },
        {
            'title': 'SWMS-004 — Working at Heights (Scaffold Erection)',
            'high_risk_activity': 'Working at Heights >2m',
            'project_name': 'Downtown Bridge Reconstruction',
            'description': 'Erection and use of scaffold for bridge soffit works.',
            'ppe_required': ['Hard Hat', 'Hi-Vis Vest', 'Safety Boots', 'Fall Arrest Harness', 'Lanyard'],
            'plant_equipment': ['Modular scaffold system', 'Edge protection'],
            'legislation': ['Managing Risk of Falls at Workplaces CoP 2015'],
            'training_required': ['Working at Heights ticket', 'Scaffold High Risk Work Licence'],
            'steps': [
                {'step_no': 1.0, 'activity': 'Scaffold design verification', 'hazards': ['Scaffold collapse'], 'risk_class': 1, 'controls': ['Engineer-certified design', 'Competent erector'], 'responsible': ['Scaffold Supervisor'], 'residual_risk': 3},
                {'step_no': 2.0, 'activity': 'Erection of scaffold base', 'hazards': ['Falls from height', 'Falling objects'], 'risk_class': 1, 'controls': ['Edge protection', 'Tool tethers', 'Exclusion zone'], 'responsible': ['Scaffolders'], 'residual_risk': 3},
                {'step_no': 3.0, 'activity': 'Daily inspection & tagging', 'hazards': ['Use of defective scaffold'], 'risk_class': 2, 'controls': ['Scafftag system', 'Pre-use inspection'], 'responsible': ['Site Supervisor'], 'residual_risk': 4},
                {'step_no': 4.0, 'activity': 'Working on scaffold platform', 'hazards': ['Fall from edge'], 'risk_class': 2, 'controls': ['100% tie-off above 2m', 'Toe boards'], 'responsible': ['All Workers'], 'residual_risk': 4},
                {'step_no': 5.0, 'activity': 'Dismantling', 'hazards': ['Material falling'], 'risk_class': 1, 'controls': ['Reverse erection sequence', 'Drop zone'], 'responsible': ['Scaffolders'], 'residual_risk': 3},
            ],
            'location_id': locations[1]['id'] if len(locations) > 1 else None,
            'location_name': locations[1]['name'] if len(locations) > 1 else None,
            'status': 'active',
            'valid_from': (today - timedelta(days=5)).isoformat(),
            'valid_until': (today + timedelta(days=60)).isoformat(),
        },
        {
            'title': 'SWMS-005 — Hot Work (Welding & Cutting)',
            'high_risk_activity': 'Hot Work',
            'description': 'Welding and oxy-cutting of structural steel on site.',
            'ppe_required': ['Welding Helmet', 'Leather Gloves', 'Welding Jacket', 'Safety Boots', 'Hi-Vis'],
            'plant_equipment': ['Welder', 'Oxy-acetylene set', 'Fire extinguisher (CO2 + dry chem)'],
            'legislation': ['Welding Processes CoP 2016'],
            'training_required': ['Welding ticket', 'Hot Work Permit Holder training'],
            'steps': [
                {'step_no': 1.0, 'activity': 'Hot Work Permit issued', 'hazards': ['Fire ignition'], 'risk_class': 1, 'controls': ['Permit valid', 'Combustibles removed 11m radius'], 'responsible': ['Permit Issuer'], 'residual_risk': 3},
                {'step_no': 2.0, 'activity': 'Fire watch posted', 'hazards': ['Smouldering ignition post-work'], 'risk_class': 2, 'controls': ['Fire watch 60min post-work'], 'responsible': ['Fire Watch'], 'residual_risk': 4},
                {'step_no': 3.0, 'activity': 'Welding/cutting activity', 'hazards': ['Arc flash', 'UV burns', 'Fume inhalation'], 'risk_class': 2, 'controls': ['Welding screens', 'Extraction fans'], 'responsible': ['Welder'], 'residual_risk': 4},
            ],
            'status': 'active',
            'valid_from': (today - timedelta(days=2)).isoformat(),
            'valid_until': (today + timedelta(days=30)).isoformat(),
        },
    ]
    swms_docs = []
    for s in swms_seed:
        d = SWMS(**s).model_dump()
        swms_docs.append(d)
    await db.swms.insert_many([dict(x) for x in swms_docs])

    # ITPs — TasWater inspections
    itp_seed = [
        {
            'name': 'Hydrostatic Pressure Testing — Mains & Services',
            'project_name': 'TasWater Tasman Hwy Scottsdale',
            'discipline': 'civil',
            'description': 'Per TasWater guidelines & WSA 03-2011-3.1 Water Supply Code',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'items': [
                {'item_no': 1.0, 'description': 'Pre-test visual inspection of joints', 'reference': 'WSA 03-2011 §13', 'inspection_type': 'check', 'frequency': 'Each section', 'acceptance_criteria': 'All joints visually sound', 'operations_by': 'Paneltec', 'verification_by': 'Paneltec', 'status': 'passed'},
                {'item_no': 2.0, 'description': 'Fill pipeline & vent', 'reference': 'WSA 03-2011', 'inspection_type': 'check', 'frequency': 'Each section', 'acceptance_criteria': 'No air pockets', 'operations_by': 'Paneltec', 'verification_by': 'Paneltec', 'status': 'passed'},
                {'item_no': 3.0, 'description': 'Pressurise to 1.5× design pressure', 'inspection_type': 'hold_point', 'frequency': 'Each test', 'acceptance_criteria': 'Hold pressure 2 hours, drop ≤ 5kPa', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'in_progress'},
                {'item_no': 4.0, 'description': 'Witness pressure decay test', 'inspection_type': 'witness_point', 'frequency': 'Each test', 'acceptance_criteria': 'TasWater witness present', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'pending'},
                {'item_no': 5.0, 'description': 'Pressure test certificate issued', 'inspection_type': 'check', 'acceptance_criteria': 'Signed by both parties', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'pending'},
            ],
        },
        {
            'name': 'Asphalt Reinstatement',
            'project_name': 'TasWater Tasman Hwy Scottsdale',
            'discipline': 'civil',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'items': [
                {'item_no': 1.0, 'description': 'Subgrade preparation', 'inspection_type': 'check', 'acceptance_criteria': 'CBR ≥ 5%', 'operations_by': 'Paneltec', 'verification_by': 'Paneltec', 'status': 'passed'},
                {'item_no': 2.0, 'description': 'Sub-base compaction', 'inspection_type': 'hold_point', 'acceptance_criteria': '95% MMDD', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'passed'},
                {'item_no': 3.0, 'description': 'Prime coat applied', 'inspection_type': 'check', 'acceptance_criteria': 'Coverage even', 'operations_by': 'Paneltec', 'verification_by': 'Paneltec', 'status': 'passed'},
                {'item_no': 4.0, 'description': 'Asphalt laid to spec thickness', 'inspection_type': 'witness_point', 'acceptance_criteria': '50mm AC10', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'in_progress'},
                {'item_no': 5.0, 'description': 'Final survey & line marking', 'inspection_type': 'survey', 'operations_by': 'Paneltec', 'verification_by': 'TasWater', 'status': 'pending'},
            ],
        },
        {
            'name': 'Concrete Foundation Pour — Industrial Park',
            'discipline': 'civil',
            'items': [
                {'item_no': 1.0, 'description': 'Formwork inspection', 'inspection_type': 'hold_point', 'acceptance_criteria': 'Per drawing', 'operations_by': 'Paneltec', 'verification_by': 'Engineer', 'status': 'passed'},
                {'item_no': 2.0, 'description': 'Reinforcement placement', 'inspection_type': 'hold_point', 'acceptance_criteria': 'Cover ≥ 50mm', 'operations_by': 'Paneltec', 'verification_by': 'Engineer', 'status': 'passed'},
                {'item_no': 3.0, 'description': 'Concrete delivery slump test', 'inspection_type': 'test', 'acceptance_criteria': 'Slump 80±20mm', 'operations_by': 'Paneltec', 'verification_by': 'Paneltec', 'status': 'in_progress'},
                {'item_no': 4.0, 'description': '28-day cylinder strength', 'inspection_type': 'test', 'acceptance_criteria': '32 MPa min', 'operations_by': 'Lab', 'verification_by': 'Engineer', 'status': 'pending'},
            ],
        },
    ]
    itp_docs = []
    for i in itp_seed:
        d = ITP(**i).model_dump()
        itp_docs.append(d)
    await db.itps.insert_many([dict(x) for x in itp_docs])

    # PERMITS
    permit_seed = [
        {
            'permit_type': 'hot_work',
            'description': 'Oxy-cutting of redundant steel pipe brackets',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'valid_from': today.isoformat(),
            'valid_until': (today + timedelta(hours=8)).isoformat(),
            'issued_by': 'Mathew Loone',
            'issued_to': 'John Carpenter',
            'precautions': ['Fire extinguisher within 5m', 'Fire watch 60min after', 'Combustibles removed 11m radius', 'Hot work area screened'],
            'checklist': [
                {'label': 'Combustible materials removed within 11m', 'checked': True},
                {'label': 'Fire extinguisher available', 'checked': True},
                {'label': 'Fire watch assigned', 'checked': True},
                {'label': 'SDS reviewed for materials being cut', 'checked': True},
                {'label': 'PPE inspected', 'checked': True},
            ],
            'status': 'active',
        },
        {
            'permit_type': 'excavation',
            'description': 'Trench excavation 1.8m deep for water main replacement Ch 0+200 to 0+250',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'valid_from': (today - timedelta(days=2)).isoformat(),
            'valid_until': (today + timedelta(days=5)).isoformat(),
            'issued_by': 'Patrick Monaghan',
            'issued_to': 'Mike Rodriguez',
            'precautions': ['DBYD completed', 'Shoring boxes installed', 'Edge protection', 'Daily competent person inspection'],
            'checklist': [
                {'label': 'Dial Before You Dig (DBYD) completed', 'checked': True},
                {'label': 'Underground services visually proven via vacuum excavation', 'checked': True},
                {'label': 'Shoring or benching plan in place', 'checked': True},
                {'label': 'Safe access/egress every 7.5m', 'checked': True},
                {'label': 'Emergency rescue plan briefed', 'checked': False},
            ],
            'status': 'active',
        },
        {
            'permit_type': 'confined_space',
            'description': 'Entry into existing valve chamber for inspection',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'valid_from': (today - timedelta(days=15)).isoformat(),
            'valid_until': (today - timedelta(days=14)).isoformat(),
            'issued_by': 'Patrick Monaghan',
            'issued_to': 'David Chen',
            'precautions': ['Gas test before entry', 'Stand-by person', 'Tripod rescue setup'],
            'checklist': [
                {'label': 'Atmospheric testing (O₂, LEL, H₂S, CO)', 'checked': True},
                {'label': 'Stand-by person posted', 'checked': True},
                {'label': 'Rescue tripod & harness in place', 'checked': True},
            ],
            'status': 'expired',
        },
        {
            'permit_type': 'working_at_heights',
            'description': 'Working on scaffold deck at Level 2 of bridge soffit',
            'location_id': locations[1]['id'] if len(locations) > 1 else None,
            'location_name': locations[1]['name'] if len(locations) > 1 else None,
            'valid_from': today.isoformat(),
            'valid_until': (today + timedelta(days=14)).isoformat(),
            'issued_by': 'Liam OConnor',
            'issued_to': 'Sarah Thompson',
            'precautions': ['Scaffold tagged green', 'Harness inspected', '100% tie-off above 2m'],
            'checklist': [
                {'label': 'Scaffold inspected & tagged today', 'checked': True},
                {'label': 'Harness & lanyard inspected', 'checked': True},
                {'label': 'Anchor points identified', 'checked': True},
                {'label': 'Rescue plan briefed', 'checked': True},
            ],
            'status': 'active',
        },
    ]
    permit_docs = []
    for p in permit_seed:
        d = Permit(**p).model_dump()
        permit_docs.append(d)
    await db.permits.insert_many([dict(x) for x in permit_docs])

    # ENVIRONMENTAL ASPECTS (EM-01 .. EM-09)
    env_seed = [
        ('EM-01', 'air_quality', 'Air Quality (Dust Control & Plant Emissions)', 'Manage dust from earthworks and emissions from plant', ['Water carts for dust suppression', 'Regular plant maintenance', 'Tarp loaded trucks'], 'Daily'),
        ('EM-02', 'air_quality', 'Prime & Bitumen', 'Manage fumes from bituminous works', ['Apply at correct temperature', 'PPE for crew', 'Avoid sensitive receptors'], 'Per pour'),
        ('EM-03', 'community', 'Community Relations', 'Manage community impact and complaints', ['Letterbox drop pre-works', 'Public complaints to TasWater CSC', 'Site signage with contact'], 'Ongoing'),
        ('EM-04', 'flora_fauna', 'Flora & Fauna', 'Protect biodiversity', ['No-go zones marked', 'Spotter for tree felling', 'Wildlife rescue plan'], 'Ongoing'),
        ('EM-05', 'heritage', 'Heritage & Archaeology', 'Protect cultural heritage', ['Unexpected finds procedure', 'Aboriginal cultural awareness'], 'Ongoing'),
        ('EM-06', 'noise', 'Noise Pollution', 'Manage construction noise to neighbours', ['Work during normal hours', 'Quieter equipment selection', 'Noise monitoring'], 'Weekly'),
        ('EM-07', 'soil_water', 'Soil Management (Inc. Contaminated Soil)', 'Prevent erosion & manage spoil', ['Silt fences', 'Stabilised site entry', 'Contaminated soil segregation'], 'Daily'),
        ('EM-08', 'fuels_chemicals', 'Storage of Fuels & Chemicals on Site', 'Prevent spills', ['Bunded storage', 'Spill kits at each site', 'SDS available'], 'Weekly'),
        ('EM-09', 'soil_water', 'Water Quality Monitoring', 'Protect waterways', ['No-discharge to waterways', 'Wash-down bay', 'pH monitoring'], 'Weekly'),
    ]
    env_docs = []
    for code, cat, title, desc, ctrl, freq in env_seed:
        a = EnvAspect(
            aspect_code=code, category=cat, title=title, description=desc,
            control_measures=ctrl, monitoring_frequency=freq,
            location_id=taswater['id'], location_name=taswater['name'],
            responsible='Site Supervisor'
        ).model_dump()
        env_docs.append(a)
    await db.env_aspects.insert_many([dict(x) for x in env_docs])

    # Environmental monitoring logs
    log_docs = []
    log_samples = [
        ('Noise Pollution', '72 dB', '75 dB (daytime limit)', 'within_limits', 'Background reading at boundary'),
        ('Noise Pollution', '83 dB', '75 dB', 'exceeded', 'Compactor in operation, residents notified'),
        ('Air Quality (Dust Control & Plant Emissions)', 'PM10 45 µg/m³', '50 µg/m³', 'within_limits', 'Water cart in use'),
        ('Water Quality Monitoring', 'pH 7.4', 'pH 6.5-8.5', 'within_limits', 'Wash-bay discharge'),
        ('Water Quality Monitoring', 'pH 9.1', 'pH 6.5-8.5', 'exceeded', 'Concrete washout — investigated, washout bin emptied'),
        ('Storage of Fuels & Chemicals on Site', 'No leaks', 'No leaks', 'within_limits', 'Weekly bund check OK'),
    ]
    for title, reading, threshold, status, notes in log_samples:
        aspect = next((a for a in env_docs if a['title'] == title), env_docs[0])
        days_ago = random.randint(0, 30)
        log_docs.append(EnvMonitoringLog(
            aspect_id=aspect['id'], aspect_title=aspect['title'],
            reading=reading, threshold=threshold, status=status, notes=notes,
            recorded_by='Site Supervisor',
            recorded_at=(today - timedelta(days=days_ago)).isoformat(),
        ).model_dump())
    await db.env_logs.insert_many([dict(x) for x in log_docs])

    # INSURANCE POLICIES (Paneltec own + on contractors)
    insurance_seed = [
        # Paneltec own
        ('combined_business', 'QBE Insurance Australia Ltd', '183A327842BPK', '$20M', 365),
        ('employers_liability', 'QBE Insurance Australia Ltd', '1HO1962502GWC', '$50M', 365),
        ('motor_vehicle', 'National Transport Insurance', '9924502', '$2M per vehicle', 200),
        ('public_liability', 'Allianz', 'PL-PT-2025-008', '$20M', 14),  # expiring soon!
        ('professional_indemnity', 'CGU', 'PI-PT-2025-002', '$5M', 60),
    ]
    insurance_docs = []
    for itype, company, polno, cov, days in insurance_seed:
        p = InsurancePolicy(
            contractor_id=None,
            insurance_type=itype, company=company, policy_number=polno,
            coverage_amount=cov,
            issued_date=(today - timedelta(days=365 - days)).date().isoformat(),
            expiry_date=(today + timedelta(days=days)).date().isoformat(),
        ).model_dump()
        insurance_docs.append(p)
    # Link insurance to existing contractors
    contractors = await db.contractors.find({}, {'_id': 0}).to_list(100)
    for i, c in enumerate(contractors[:3]):
        days = [180, -10, 30][i]
        p = InsurancePolicy(
            contractor_id=c['id'],
            insurance_type='public_liability',
            company='Various',
            policy_number=f'CTR-{c["id"][:6]}-PL',
            coverage_amount='$10M',
            issued_date=(today - timedelta(days=365)).date().isoformat(),
            expiry_date=(today + timedelta(days=days)).date().isoformat(),
        ).model_dump()
        insurance_docs.append(p)
    await db.insurance.insert_many([dict(x) for x in insurance_docs])

    # AUDITS
    audit_seed = [
        {
            'title': 'Internal WHS System Audit Q1',
            'audit_type': 'internal',
            'scope': 'Full WHS management system review',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'auditor': 'Emily Brooks',
            'planned_date': (today - timedelta(days=20)).date().isoformat(),
            'actual_date': (today - timedelta(days=18)).date().isoformat(),
            'status': 'completed',
            'findings': [
                {'description': 'SWMS not signed by all crew before commencement on 2 occasions', 'severity': 'minor_nc', 'corrective_action': 'Supervisor verification at pre-start', 'closed': True},
                {'description': 'First aid kit at site office missing eye-wash solution', 'severity': 'observation', 'corrective_action': 'Restock from supplier', 'closed': True},
                {'description': 'Opportunity: Digitise sign-in register via QR code', 'severity': 'opportunity', 'closed': False},
            ],
        },
        {
            'title': 'TasWater Client Compliance Audit',
            'audit_type': 'client',
            'scope': 'Audit against TasWater Contractor Compliance Standards',
            'location_id': taswater['id'],
            'location_name': taswater['name'],
            'auditor': 'TasWater HSE Team',
            'planned_date': (today + timedelta(days=14)).date().isoformat(),
            'status': 'planned',
            'findings': [],
        },
        {
            'title': 'External ISO 45001 Surveillance Audit',
            'audit_type': 'external',
            'scope': 'ISO 45001:2018 Occupational Health & Safety surveillance',
            'auditor': 'SAI Global',
            'planned_date': (today - timedelta(days=90)).date().isoformat(),
            'actual_date': (today - timedelta(days=88)).date().isoformat(),
            'status': 'completed',
            'findings': [
                {'description': 'Minor: Risk register review frequency not consistently documented', 'severity': 'minor_nc', 'corrective_action': 'Monthly review cadence with sign-off', 'closed': True},
            ],
        },
        {
            'title': 'Site Inspection — Highway 401',
            'audit_type': 'internal',
            'scope': 'PPE & site setup audit',
            'location_id': locations[0]['id'] if locations else None,
            'location_name': locations[0]['name'] if locations else None,
            'auditor': 'Liam OConnor',
            'planned_date': (today + timedelta(days=3)).date().isoformat(),
            'status': 'planned',
            'findings': [],
        },
    ]
    audit_docs = []
    for a in audit_seed:
        d = Audit(**a).model_dump()
        audit_docs.append(d)
    await db.audits.insert_many([dict(x) for x in audit_docs])

    return {
        'ok': True,
        'counts': {
            'swms': len(swms_docs), 'itps': len(itp_docs), 'permits': len(permit_docs),
            'env_aspects': len(env_docs), 'env_logs': len(log_docs),
            'insurance': len(insurance_docs), 'audits': len(audit_docs),
        }
    }





# ============== DOCUMENT LIBRARY ==============
# Pre-defined categories based on Paneltec Risk & Compliance structure
DEFAULT_DOC_CATEGORIES = [
    {'slug': 'alcohol-drug', 'name': 'Alcohol & Drug Screening', 'icon': 'flask', 'color': '#8B5CF6'},
    {'slug': 'asbestos', 'name': 'Asbestos', 'icon': 'shield-alert', 'color': '#DC2626'},
    {'slug': 'audits', 'name': 'Audits', 'icon': 'clipboard-check', 'color': '#0EA5E9'},
    {'slug': 'australian-standards', 'name': 'Australian Standards', 'icon': 'book', 'color': '#0891B2'},
    {'slug': 'barriers', 'name': 'Barriers', 'icon': 'fence', 'color': '#EA580C'},
    {'slug': 'byda', 'name': 'BYDA (Before You Dig)', 'icon': 'shovel', 'color': '#CA8A04'},
    {'slug': 'calibration', 'name': 'Calibration Certificates', 'icon': 'gauge', 'color': '#0D9488'},
    {'slug': 'carbon-reduction', 'name': 'Carbon Reduction', 'icon': 'leaf', 'color': '#16A34A'},
    {'slug': 'ccf', 'name': 'CCF (Civil Contractors Federation)', 'icon': 'building', 'color': '#2563EB'},
    {'slug': 'checklists', 'name': 'Checklists', 'icon': 'list-checks', 'color': '#7C3AED'},
    {'slug': 'chemical-storage', 'name': 'Chemical Storage & Bunding', 'icon': 'flask', 'color': '#B91C1C'},
    {'slug': 'codesafe', 'name': 'CodeSafe', 'icon': 'shield', 'color': '#059669'},
    {'slug': 'committees', 'name': 'Committees & Memberships', 'icon': 'users', 'color': '#9333EA'},
    {'slug': 'competencies', 'name': 'Competencies Matrices', 'icon': 'grid', 'color': '#0284C7'},
    {'slug': 'confined-space', 'name': 'Confined Space', 'icon': 'box', 'color': '#DC2626'},
    {'slug': 'contract-mgmt', 'name': 'Contract Management Plans', 'icon': 'file-text', 'color': '#0F172A'},
    {'slug': 'electrical', 'name': 'Electrical Safety', 'icon': 'zap', 'color': '#F59E0B'},
    {'slug': 'emergency', 'name': 'Emergency Management', 'icon': 'siren', 'color': '#DC2626'},
    {'slug': 'environmental', 'name': 'Environmental Management', 'icon': 'leaf', 'color': '#16A34A'},
    {'slug': 'first-aid', 'name': 'First Aid', 'icon': 'heart-pulse', 'color': '#EF4444'},
    {'slug': 'forms', 'name': 'Forms', 'icon': 'file-text', 'color': '#6366F1'},
    {'slug': 'heights', 'name': 'Working at Heights', 'icon': 'arrow-up', 'color': '#F97316'},
    {'slug': 'hot-work', 'name': 'Hot Work', 'icon': 'flame', 'color': '#DC2626'},
    {'slug': 'incidents', 'name': 'Incident Reports', 'icon': 'alert-triangle', 'color': '#DC2626'},
    {'slug': 'inductions', 'name': 'Inductions', 'icon': 'user-check', 'color': '#0891B2'},
    {'slug': 'insurance', 'name': 'Insurance', 'icon': 'shield-check', 'color': '#2563EB'},
    {'slug': 'itp', 'name': 'ITPs (Inspection & Test Plans)', 'icon': 'clipboard-list', 'color': '#0D9488'},
    {'slug': 'jsea', 'name': 'JSEA / Risk Assessments', 'icon': 'alert-octagon', 'color': '#F59E0B'},
    {'slug': 'licences', 'name': 'Licences & Tickets', 'icon': 'award', 'color': '#CA8A04'},
    {'slug': 'manuals', 'name': 'Manuals & Procedures', 'icon': 'book-open', 'color': '#475569'},
    {'slug': 'permits', 'name': 'Permits to Work', 'icon': 'key', 'color': '#9333EA'},
    {'slug': 'plant-equipment', 'name': 'Plant & Equipment', 'icon': 'truck', 'color': '#0EA5E9'},
    {'slug': 'policies', 'name': 'Company Policies', 'icon': 'file-badge', 'color': '#1E40AF'},
    {'slug': 'ppe', 'name': 'PPE', 'icon': 'hard-hat', 'color': '#FBBF24'},
    {'slug': 'procurement', 'name': 'Procurement', 'icon': 'shopping-cart', 'color': '#7C3AED'},
    {'slug': 'rehab', 'name': 'Rehabilitation & RTW', 'icon': 'activity', 'color': '#0D9488'},
    {'slug': 'reports', 'name': 'Reports', 'icon': 'bar-chart', 'color': '#3B82F6'},
    {'slug': 'sds', 'name': 'SDS (Safety Data Sheets)', 'icon': 'flask-conical', 'color': '#B91C1C'},
    {'slug': 'site-management', 'name': 'Site Management', 'icon': 'map-pin', 'color': '#0891B2'},
    {'slug': 'subcontractors', 'name': 'Subcontractor Management', 'icon': 'briefcase', 'color': '#7C3AED'},
    {'slug': 'swms', 'name': 'SWMS', 'icon': 'shield-alert', 'color': '#EA580C'},
    {'slug': 'toolbox', 'name': 'Toolbox Talks', 'icon': 'message-square', 'color': '#10B981'},
    {'slug': 'traffic', 'name': 'Traffic Management', 'icon': 'traffic-cone', 'color': '#F97316'},
    {'slug': 'training', 'name': 'Training Records', 'icon': 'graduation-cap', 'color': '#0284C7'},
    {'slug': 'whs-acts', 'name': 'WHS Acts & Regulations', 'icon': 'scale', 'color': '#1E40AF'},
    {'slug': 'uncategorized', 'name': 'Uncategorized', 'icon': 'folder', 'color': '#64748B'},
]

class DocCategory(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str
    name: str
    icon: str = 'folder'
    color: str = '#64748B'
    created_at: str = Field(default_factory=now_iso)

class Document(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    category_slug: str = 'uncategorized'
    subfolder_id: Optional[str] = None  # nested 2-level folder support
    file_type: str = 'unknown'
    mime_type: Optional[str] = None
    size_bytes: int = 0
    content_b64: Optional[str] = None
    ai_summary: Optional[str] = None
    ai_tags: List[str] = []
    ai_doc_type: Optional[str] = None
    is_form: bool = False
    extracted_fields: List[Dict[str, Any]] = []
    uploaded_by: Optional[str] = None
    description: Optional[str] = None
    # NEW: versioning
    version: int = 1
    parent_doc_id: Optional[str] = None  # if this is a version of another doc, points to current latest
    is_latest: bool = True
    # NEW: expiry & review
    expiry_date: Optional[str] = None
    review_date: Optional[str] = None
    review_frequency_months: Optional[int] = None
    # NEW: acknowledgement
    requires_ack: bool = False
    ack_assignee_ids: List[str] = []  # worker IDs; empty list = all workers
    # NEW: embeddings for semantic search (small float array)
    embedding: List[float] = []
    # NEW: source tracking
    source: str = 'manual'  # manual, dropbox, onedrive
    source_ref: Optional[str] = None  # e.g. Dropbox path
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/doc-categories')
async def list_doc_categories(user=Depends(get_current_user)):
    items = await db.doc_categories.find({}, {'_id': 0}).to_list(200)
    if not items:
        # Auto-seed on first call
        for c in DEFAULT_DOC_CATEGORIES:
            doc = DocCategory(**c).model_dump()
            await db.doc_categories.insert_one(doc)
        items = await db.doc_categories.find({}, {'_id': 0}).to_list(200)
    # Add document counts
    docs = await db.documents.find({}, {'_id': 0, 'category_slug': 1}).to_list(5000)
    counts = {}
    for d in docs:
        slug = d.get('category_slug', 'uncategorized')
        counts[slug] = counts.get(slug, 0) + 1
    for c in items:
        c['doc_count'] = counts.get(c['slug'], 0)
    return items

@api_router.post('/doc-categories')
async def create_doc_category(c: DocCategory, user=Depends(require_admin)):
    doc = c.model_dump()
    await db.doc_categories.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/doc-categories/{cid}')
async def delete_doc_category(cid: str, user=Depends(require_admin)):
    await db.doc_categories.delete_one({'id': cid})
    return {'ok': True}

@api_router.get('/documents')
async def list_documents(category: Optional[str] = None, subfolder: Optional[str] = None,
                         search: Optional[str] = None, include_versions: bool = False,
                         user=Depends(get_current_user)):
    q = {}
    if not include_versions:
        q['is_latest'] = True
    if category and category != 'all':
        q['category_slug'] = category
    if subfolder == 'root':
        q['subfolder_id'] = None
    elif subfolder:
        q['subfolder_id'] = subfolder
    if search:
        q['$or'] = [
            {'name': {'$regex': search, '$options': 'i'}},
            {'ai_tags': {'$regex': search, '$options': 'i'}},
            {'ai_summary': {'$regex': search, '$options': 'i'}},
        ]
    items = await db.documents.find(q, {'_id': 0, 'content_b64': 0, 'embedding': 0}).sort('created_at', -1).to_list(2000)
    return items

@api_router.get('/documents/{did}')
async def get_document(did: str, user=Depends(get_current_user)):
    d = await db.documents.find_one({'id': did}, {'_id': 0, 'embedding': 0})
    if not d: raise HTTPException(404, 'Not found')
    return d

async def _generate_embedding(text: str) -> List[float]:
    """Embeddings not available via Emergent proxy. Returns empty — semantic search falls back to AI keyword expansion."""
    return []

async def _expand_query_ai(query: str) -> List[str]:
    """Use chat LLM to expand a query into related search terms (semantic-ish)."""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"qexpand-{uuid.uuid4()}",
            system_message="You expand search queries for a civil contractor safety document library. Return a JSON array of 5-10 related search terms (synonyms, abbreviations, related concepts). Output only the JSON array, no commentary."
        ).with_model('openai', 'gpt-4o-mini')
        result = await chat.send_message(UserMessage(text=f"Query: \"{query}\"\nReturn ONLY a JSON array like [\"term1\", \"term2\", ...]"))
        raw = str(result).strip()
        if raw.startswith('```'):
            raw = raw.split('```', 2)[1]
            if raw.startswith('json'): raw = raw[4:]
            raw = raw.strip()
        terms = json.loads(raw)
        if isinstance(terms, list):
            return [str(t) for t in terms if t][:10]
    except Exception as e:
        logger.warning(f"Query expansion failed: {e}")
    return []

def _cosine_similarity(a: List[float], b: List[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    import math
    dot = sum(x*y for x, y in zip(a, b))
    na = math.sqrt(sum(x*x for x in a))
    nb = math.sqrt(sum(y*y for y in b))
    if na == 0 or nb == 0: return 0.0
    return dot / (na * nb)

@api_router.post('/documents')
async def create_document(d: Document, user=Depends(get_current_user)):
    doc = d.model_dump()
    doc['uploaded_by'] = user.get('name')
    # Generate embedding from name + summary + tags for semantic search
    embed_text = f"{doc.get('name','')}\n{doc.get('ai_summary','')}\n{' '.join(doc.get('ai_tags',[]))}\n{doc.get('ai_doc_type','')}"
    doc['embedding'] = await _generate_embedding(embed_text)
    await db.documents.insert_one(doc)
    doc.pop('_id', None)
    doc.pop('content_b64', None)
    doc.pop('embedding', None)
    return doc

@api_router.put('/documents/{did}')
async def update_document(did: str, body: dict, user=Depends(get_current_user)):
    allowed_keys = ('name', 'category_slug', 'subfolder_id', 'description',
                    'ai_summary', 'ai_tags', 'ai_doc_type', 'is_form', 'extracted_fields',
                    'expiry_date', 'review_date', 'review_frequency_months',
                    'requires_ack', 'ack_assignee_ids')
    allowed = {k: v for k, v in body.items() if k in allowed_keys}
    await db.documents.update_one({'id': did}, {'$set': allowed})
    d = await db.documents.find_one({'id': did}, {'_id': 0, 'content_b64': 0, 'embedding': 0})
    return d

@api_router.delete('/documents/{did}')
async def delete_document(did: str, user=Depends(get_current_user)):
    await db.documents.delete_one({'id': did})
    return {'ok': True}

# AI classifier
class AIClassifyIn(BaseModel):
    filename: str
    content_b64: Optional[str] = None  # data: prefix or raw base64
    text_excerpt: Optional[str] = None  # if client extracted text already
    file_type: Optional[str] = None

def _detect_filetype(name: str, mime: Optional[str] = None) -> str:
    n = name.lower()
    if n.endswith('.pdf'): return 'pdf'
    if n.endswith('.docx') or n.endswith('.doc'): return 'docx'
    if n.endswith('.xlsx') or n.endswith('.xls') or n.endswith('.csv'): return 'xlsx'
    if n.endswith('.pptx') or n.endswith('.ppt'): return 'pptx'
    if any(n.endswith(e) for e in ('.png', '.jpg', '.jpeg', '.webp', '.gif')): return 'image'
    if n.endswith('.txt') or n.endswith('.md'): return 'txt'
    return 'other'

def _extract_text_from_b64(content_b64: str, file_type: str) -> str:
    """Best-effort text extraction from a base64 file."""
    try:
        # Strip data: prefix
        if ',' in content_b64:
            content_b64 = content_b64.split(',', 1)[1]
        raw = base64.b64decode(content_b64)
        if file_type == 'txt':
            return raw.decode('utf-8', errors='ignore')[:8000]
        if file_type == 'docx':
            try:
                from docx import Document as DocxDoc
                buf = io.BytesIO(raw)
                d = DocxDoc(buf)
                txt = []
                for p in d.paragraphs:
                    if p.text.strip(): txt.append(p.text.strip())
                for t in d.tables[:20]:
                    for r in t.rows[:20]:
                        for c in r.cells:
                            if c.text.strip(): txt.append(c.text.strip())
                return '\n'.join(txt)[:8000]
            except Exception as e:
                logger.warning(f"docx extract failed: {e}")
                return ''
        if file_type == 'pdf':
            try:
                # Lightweight: try pypdf if available, else skip
                try:
                    from pypdf import PdfReader
                except ImportError:
                    return ''
                buf = io.BytesIO(raw)
                reader = PdfReader(buf)
                txt = []
                for p in reader.pages[:10]:
                    try: txt.append(p.extract_text() or '')
                    except: pass
                return '\n'.join(txt)[:8000]
            except Exception as e:
                logger.warning(f"pdf extract failed: {e}")
                return ''
        if file_type == 'xlsx':
            try:
                from openpyxl import load_workbook
                buf = io.BytesIO(raw)
                wb = load_workbook(buf, read_only=True, data_only=True)
                txt = []
                for sh in wb.sheetnames[:5]:
                    ws = wb[sh]
                    txt.append(f"# Sheet: {sh}")
                    for row in ws.iter_rows(max_row=50, values_only=True):
                        cells = [str(c) for c in row if c is not None]
                        if cells: txt.append(' | '.join(cells))
                return '\n'.join(txt)[:8000]
            except Exception as e:
                logger.warning(f"xlsx extract failed: {e}")
                return ''
        return ''
    except Exception as e:
        logger.warning(f"extract_text error: {e}")
        return ''

@api_router.post('/ai/classify-document')
async def ai_classify_document(body: AIClassifyIn, user=Depends(get_current_user)):
    """AI classifies a document into a category, generates tags + summary, detects if it's a form."""
    file_type = body.file_type or _detect_filetype(body.filename)
    text = body.text_excerpt or ''
    if not text and body.content_b64:
        text = _extract_text_from_b64(body.content_b64, file_type)

    # Build category catalog for the AI
    cat_list = '\n'.join([f"- {c['slug']}: {c['name']}" for c in DEFAULT_DOC_CATEGORIES])

    prompt = f"""You are a document classification assistant for a civil contractor's risk & compliance library.

Filename: {body.filename}
File type: {file_type}

Available categories (use the slug):
{cat_list}

Document text excerpt (first 8000 chars):
\"\"\"
{text[:7500] if text else '(no text extracted — classify based on filename only)'}
\"\"\"

Return ONLY a valid JSON object with this structure (no markdown, no commentary):
{{
  "category_slug": "one-of-the-slugs-above",
  "doc_type": "Policy|Procedure|SWMS|Permit|ITP|Form|Checklist|Risk Assessment|SDS|Standard|Manual|Report|Certificate|Audit|Training Record|Other",
  "summary": "2-3 sentence summary of what this document is about",
  "tags": ["tag1", "tag2", "tag3", "tag4"],
  "is_form": true_or_false_whether_this_is_a_fillable_form_or_checklist,
  "extracted_fields": [
    {{"label": "Field label from form", "type": "text|textarea|date|number|select|checkbox|signature|photo", "required": true_or_false}},
    ...
  ]
}}

Only include extracted_fields if is_form is true. Extract realistic form fields if you can identify them in the text. Be concise."""

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"doc-classify-{uuid.uuid4()}",
            system_message="You classify civil-contractor risk & compliance documents. Return only valid JSON, no commentary, no markdown code fences."
        ).with_model('openai', 'gpt-4o-mini')
        result = await chat.send_message(UserMessage(text=prompt))
        raw = str(result).strip()
        if raw.startswith('```'):
            raw = raw.split('```', 2)[1]
            if raw.startswith('json'): raw = raw[4:]
            raw = raw.strip()
        import json
        data = json.loads(raw)
        # Validate category exists
        valid_slugs = {c['slug'] for c in DEFAULT_DOC_CATEGORIES}
        if data.get('category_slug') not in valid_slugs:
            data['category_slug'] = 'uncategorized'
        return data
    except Exception as e:
        logger.error(f"AI classify error: {e}")
        # Heuristic fallback based on filename
        name_l = body.filename.lower()
        guess_slug = 'uncategorized'
        for c in DEFAULT_DOC_CATEGORIES:
            kw = c['name'].lower().split()[0]
            if kw in name_l:
                guess_slug = c['slug']; break
        return {
            'category_slug': guess_slug,
            'doc_type': 'Other',
            'summary': f"Document {body.filename} — automatic classification unavailable, please review.",
            'tags': [file_type],
            'is_form': False,
            'extracted_fields': [],
            '_fallback': True,
        }

@api_router.post('/documents/{did}/to-template')
async def doc_to_template(did: str, user=Depends(require_admin)):
    """Convert a classified form document into a fillable form template."""
    d = await db.documents.find_one({'id': did}, {'_id': 0})
    if not d: raise HTTPException(404, 'Not found')
    if not d.get('extracted_fields'):
        raise HTTPException(400, 'No extracted fields — run AI classify first or document is not a form')

    tpl = FormTemplate(
        name=d['name'].rsplit('.', 1)[0],
        description=f"Auto-generated from document: {d['name']}",
        category='general',
        fields=[{'id': f"f{i+1}", **f} for i, f in enumerate(d['extracted_fields'])],
    ).model_dump()
    await db.form_templates.insert_one(tpl)
    tpl.pop('_id', None)
    return tpl



# ============== SUBFOLDERS ==============
class Subfolder(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    category_slug: str
    name: str
    color: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/subfolders')
async def list_subfolders(category: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if category: q['category_slug'] = category
    items = await db.subfolders.find(q, {'_id': 0}).sort('name', 1).to_list(500)
    # Add doc counts per subfolder
    docs = await db.documents.find({'is_latest': True}, {'_id': 0, 'subfolder_id': 1}).to_list(5000)
    counts = {}
    for d in docs:
        sid = d.get('subfolder_id')
        if sid:
            counts[sid] = counts.get(sid, 0) + 1
    for s in items:
        s['doc_count'] = counts.get(s['id'], 0)
    return items

@api_router.post('/subfolders')
async def create_subfolder(s: Subfolder, user=Depends(get_current_user)):
    doc = s.model_dump()
    await db.subfolders.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/subfolders/{sid}')
async def delete_subfolder(sid: str, user=Depends(require_admin)):
    await db.documents.update_many({'subfolder_id': sid}, {'$set': {'subfolder_id': None}})
    await db.subfolders.delete_one({'id': sid})
    return {'ok': True}

# ============== DOCUMENT VERSIONING ==============
@api_router.get('/documents/{did}/versions')
async def list_doc_versions(did: str, user=Depends(get_current_user)):
    """Return all versions of a document (current + history)."""
    current = await db.documents.find_one({'id': did}, {'_id': 0, 'embedding': 0, 'content_b64': 0})
    if not current: raise HTTPException(404, 'Not found')
    # Find current "head" (the latest version of this lineage)
    head_id = current['id'] if current.get('is_latest') else current.get('parent_doc_id')
    if not head_id: head_id = current['id']
    # All docs in lineage = head + all with parent_doc_id pointing to head
    head = await db.documents.find_one({'id': head_id}, {'_id': 0, 'embedding': 0, 'content_b64': 0})
    history = await db.documents.find(
        {'parent_doc_id': head_id, 'is_latest': False},
        {'_id': 0, 'embedding': 0, 'content_b64': 0}
    ).sort('version', -1).to_list(100)
    return {'current': head, 'history': history}

class NewVersionIn(BaseModel):
    name: Optional[str] = None
    content_b64: str
    file_type: Optional[str] = None
    mime_type: Optional[str] = None
    size_bytes: int = 0
    change_note: Optional[str] = None

@api_router.post('/documents/{did}/new-version')
async def upload_new_version(did: str, body: NewVersionIn, user=Depends(get_current_user)):
    """Upload a new version of an existing document. Old version is archived as is_latest=false."""
    current = await db.documents.find_one({'id': did}, {'_id': 0})
    if not current: raise HTTPException(404, 'Document not found')
    head_id = current['id']  # current is always treated as the head

    # Archive current
    archive_id = str(uuid.uuid4())
    archive = dict(current)
    archive['id'] = archive_id
    archive['parent_doc_id'] = head_id
    archive['is_latest'] = False
    await db.documents.insert_one(archive)

    # Update head with new file
    new_version = current.get('version', 1) + 1
    update = {
        'content_b64': body.content_b64,
        'size_bytes': body.size_bytes,
        'version': new_version,
        'created_at': now_iso(),
        'uploaded_by': user.get('name'),
    }
    if body.name: update['name'] = body.name
    if body.file_type: update['file_type'] = body.file_type
    if body.mime_type: update['mime_type'] = body.mime_type
    if body.change_note: update['description'] = body.change_note
    await db.documents.update_one({'id': did}, {'$set': update})

    refreshed = await db.documents.find_one({'id': did}, {'_id': 0, 'embedding': 0, 'content_b64': 0})
    return refreshed

# ============== DOCUMENT EXPIRY OVERVIEW ==============
@api_router.get('/documents/expiring')
async def list_expiring_docs(days: int = 90, user=Depends(get_current_user)):
    """Documents expiring or due for review within X days."""
    today = datetime.now(timezone.utc).date()
    cutoff = today + timedelta(days=days)
    docs = await db.documents.find(
        {'is_latest': True, '$or': [
            {'expiry_date': {'$ne': None, '$exists': True}},
            {'review_date': {'$ne': None, '$exists': True}},
        ]},
        {'_id': 0, 'content_b64': 0, 'embedding': 0}
    ).to_list(2000)
    result = []
    for d in docs:
        for field in ('expiry_date', 'review_date'):
            v = d.get(field)
            if not v: continue
            try:
                dt = datetime.fromisoformat(v).date()
                if dt <= cutoff:
                    days_left = (dt - today).days
                    status = 'expired' if days_left < 0 else 'critical' if days_left <= 14 else 'warning'
                    result.append({
                        **d,
                        '_field': field,
                        '_field_date': v,
                        '_days_left': days_left,
                        '_status': status,
                    })
                    break
            except Exception:
                pass
    result.sort(key=lambda x: x.get('_days_left', 9999))
    return result

# ============== ACKNOWLEDGEMENTS ==============
class Acknowledgement(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    user_id: str
    user_name: Optional[str] = None
    acknowledged_at: str = Field(default_factory=now_iso)

@api_router.get('/acknowledgements/document/{did}')
async def list_ack_for_doc(did: str, user=Depends(get_current_user)):
    items = await db.acknowledgements.find({'document_id': did}, {'_id': 0}).to_list(2000)
    return items

@api_router.get('/acknowledgements/required')
async def list_my_required_acks(user=Depends(get_current_user)):
    """List documents this user needs to acknowledge but hasn't yet."""
    # Get all docs requiring ack
    docs = await db.documents.find(
        {'requires_ack': True, 'is_latest': True},
        {'_id': 0, 'content_b64': 0, 'embedding': 0}
    ).to_list(2000)
    # Find user's worker_id
    worker = await db.workers.find_one({'email': user.get('email')}, {'_id': 0})
    worker_id = worker['id'] if worker else None

    # User's existing acks
    my_acks = await db.acknowledgements.find({'user_id': user['id']}, {'_id': 0}).to_list(2000)
    acked_ids = {a['document_id'] for a in my_acks}

    required = []
    for d in docs:
        assignees = d.get('ack_assignee_ids', [])
        # Empty assignees = everyone
        if assignees and (not worker_id or worker_id not in assignees):
            continue
        if d['id'] in acked_ids:
            continue
        required.append(d)
    return required

@api_router.post('/acknowledgements')
async def create_ack(body: dict, user=Depends(get_current_user)):
    did = body.get('document_id')
    if not did: raise HTTPException(400, 'document_id required')
    # Prevent duplicates
    existing = await db.acknowledgements.find_one({'document_id': did, 'user_id': user['id']})
    if existing:
        return {'ok': True, 'already_acked': True}
    ack = Acknowledgement(
        document_id=did,
        user_id=user['id'],
        user_name=user.get('name'),
    ).model_dump()
    await db.acknowledgements.insert_one(ack)
    ack.pop('_id', None)
    return ack

@api_router.get('/acknowledgements/compliance')
async def ack_compliance(user=Depends(get_current_user)):
    """Compliance % per document requiring ack."""
    docs = await db.documents.find(
        {'requires_ack': True, 'is_latest': True},
        {'_id': 0, 'content_b64': 0, 'embedding': 0}
    ).to_list(2000)
    workers = await db.workers.find({}, {'_id': 0}).to_list(2000)
    worker_id_by_email = {w.get('email'): w['id'] for w in workers if w.get('email')}
    all_worker_ids = {w['id'] for w in workers}
    result = []
    for d in docs:
        assignees = set(d.get('ack_assignee_ids') or [])
        if not assignees:
            assignees = all_worker_ids
        acks = await db.acknowledgements.find({'document_id': d['id']}, {'_id': 0}).to_list(2000)
        # Acks are by user_id; map via email to worker_id
        users = await db.users.find({'id': {'$in': [a['user_id'] for a in acks]}}, {'_id': 0}).to_list(2000)
        ack_worker_ids = set()
        for u in users:
            wid = worker_id_by_email.get(u.get('email'))
            if wid: ack_worker_ids.add(wid)
        completed = len(assignees & ack_worker_ids)
        total = len(assignees)
        result.append({
            'document_id': d['id'],
            'document_name': d['name'],
            'category_slug': d.get('category_slug'),
            'total': total,
            'completed': completed,
            'pct': round(100 * completed / total, 1) if total else 0,
        })
    return result

# ============== SEMANTIC SEARCH ==============
class SemanticSearchIn(BaseModel):
    query: str
    top_k: int = 10

@api_router.post('/documents/search/semantic')
async def semantic_search(body: SemanticSearchIn, user=Depends(get_current_user)):
    """AI-powered search: expand query into related terms, score by keyword + tag overlap."""
    q = body.query.strip()
    if not q:
        return {'results': [], 'mode': 'empty'}

    # Expand query with AI
    expanded = await _expand_query_ai(q)
    all_terms = [q.lower()] + [t.lower() for t in expanded]

    # Score every latest doc
    docs = await db.documents.find(
        {'is_latest': True},
        {'_id': 0, 'content_b64': 0, 'embedding': 0}
    ).to_list(5000)

    scored = []
    for d in docs:
        haystack = ' '.join([
            d.get('name', ''),
            d.get('ai_summary', '') or '',
            ' '.join(d.get('ai_tags') or []),
            d.get('ai_doc_type', '') or '',
            d.get('category_slug', '') or '',
        ]).lower()
        score = 0.0
        for term in all_terms:
            if not term: continue
            if term in haystack:
                # Boost name + tags hits
                if term in d.get('name', '').lower(): score += 2.0
                if term in ' '.join(d.get('ai_tags') or []).lower(): score += 1.5
                score += 1.0
        if score > 0:
            scored.append((score, d))

    scored.sort(key=lambda x: -x[0])
    top = scored[:body.top_k]
    return {
        'results': [{**d, '_score': round(s, 2), '_mode': 'ai'} for s, d in top],
        'mode': 'ai',
        'expanded_terms': expanded,
    }

@api_router.post('/documents/reindex')
async def reindex_embeddings(user=Depends(require_admin)):
    """Stub — embeddings not in use (Emergent proxy doesn't support them)."""
    return {'ok': True, 'note': 'AI keyword expansion mode — no embedding index needed.'}

# ============== DROPBOX SYNC (paste-token approach) ==============
class DropboxConfig(BaseModel):
    access_token: str
    folder_path: str = ''  # empty = root; eg '/Risk & Compliance'

@api_router.get('/integrations/dropbox/status')
async def dropbox_status(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'dropbox'}, {'_id': 0})
    if not cfg:
        return {'connected': False}
    return {
        'connected': True,
        'folder_path': cfg.get('folder_path', ''),
        'last_sync': cfg.get('last_sync'),
        'last_sync_count': cfg.get('last_sync_count', 0),
    }

@api_router.post('/integrations/dropbox/connect')
async def dropbox_connect(body: DropboxConfig, user=Depends(require_admin)):
    """Save Dropbox token + verify by listing the folder."""
    import httpx
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.post(
            'https://api.dropboxapi.com/2/users/get_current_account',
            headers={'Authorization': f'Bearer {body.access_token}'},
        )
        if r.status_code != 200:
            raise HTTPException(400, f'Dropbox auth failed: {r.text[:200]}')
        account = r.json()
    await db.integrations.update_one(
        {'id': 'dropbox'},
        {'$set': {
            'id': 'dropbox',
            'access_token': body.access_token,
            'folder_path': body.folder_path,
            'connected_at': now_iso(),
            'account_email': account.get('email'),
            'account_name': account.get('name', {}).get('display_name'),
        }},
        upsert=True,
    )
    return {'ok': True, 'account': account.get('email')}

@api_router.delete('/integrations/dropbox')
async def dropbox_disconnect(user=Depends(require_admin)):
    await db.integrations.delete_one({'id': 'dropbox'})
    return {'ok': True}

@api_router.post('/integrations/dropbox/sync')
async def dropbox_sync(classify: bool = True, max_files: int = 50, user=Depends(require_admin)):
    """Pull files from configured Dropbox folder, classify with AI, ingest as documents."""
    cfg = await db.integrations.find_one({'id': 'dropbox'}, {'_id': 0})
    if not cfg: raise HTTPException(400, 'Dropbox not connected')
    token = cfg['access_token']
    folder = cfg.get('folder_path', '')

    import httpx
    synced = []
    errors = []
    async with httpx.AsyncClient(timeout=60.0) as client:
        # List files recursively
        cursor = None
        all_entries = []
        while True:
            url = 'https://api.dropboxapi.com/2/files/list_folder' if cursor is None else 'https://api.dropboxapi.com/2/files/list_folder/continue'
            payload = {'cursor': cursor} if cursor else {'path': folder, 'recursive': True, 'limit': 200}
            r = await client.post(url, headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'}, json=payload)
            if r.status_code != 200:
                errors.append(f'list_folder failed: {r.text[:200]}')
                break
            data = r.json()
            all_entries.extend([e for e in data.get('entries', []) if e.get('.tag') == 'file'])
            if not data.get('has_more'): break
            cursor = data.get('cursor')
            if len(all_entries) > max_files * 3: break

        # Get already-synced source_refs to skip
        existing_refs = set()
        async for d in db.documents.find({'source': 'dropbox'}, {'_id': 0, 'source_ref': 1, 'name': 1, 'size_bytes': 1}):
            if d.get('source_ref'): existing_refs.add(d['source_ref'])

        # Process up to max_files new files
        to_process = [e for e in all_entries if e.get('path_lower') not in existing_refs][:max_files]

        for entry in to_process:
            try:
                path = entry['path_lower']
                name = entry['name']
                size = entry.get('size', 0)
                if size > 25 * 1024 * 1024:  # skip files > 25MB
                    errors.append(f'{name}: too large ({size} bytes)')
                    continue
                # Download
                r = await client.post(
                    'https://content.dropboxapi.com/2/files/download',
                    headers={
                        'Authorization': f'Bearer {token}',
                        'Dropbox-API-Arg': json.dumps({'path': entry['path_display']}),
                    },
                )
                if r.status_code != 200:
                    errors.append(f'{name}: download failed {r.status_code}')
                    continue
                raw = r.content
                file_type = _detect_filetype(name)
                # Build data URL
                mime = {
                    'pdf': 'application/pdf', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    'image': 'image/png', 'txt': 'text/plain', 'other': 'application/octet-stream',
                }.get(file_type, 'application/octet-stream')
                b64 = f"data:{mime};base64,{base64.b64encode(raw).decode()}"

                # AI classify
                classification = {'category_slug': 'uncategorized', 'doc_type': 'Other', 'summary': '', 'tags': [], 'is_form': False, 'extracted_fields': []}
                if classify:
                    try:
                        text = _extract_text_from_b64(b64, file_type)
                        ai_in = AIClassifyIn(filename=name, content_b64=b64, file_type=file_type, text_excerpt=text)
                        ai_result = await ai_classify_document(ai_in, user)
                        classification.update(ai_result)
                    except Exception as ce:
                        errors.append(f'{name}: classify failed {str(ce)[:100]}')

                # Build doc
                doc = Document(
                    name=name,
                    category_slug=classification.get('category_slug', 'uncategorized'),
                    file_type=file_type,
                    mime_type=mime,
                    size_bytes=len(raw),
                    content_b64=b64,
                    ai_summary=classification.get('summary'),
                    ai_tags=classification.get('tags', []),
                    ai_doc_type=classification.get('doc_type'),
                    is_form=classification.get('is_form', False),
                    extracted_fields=classification.get('extracted_fields', []),
                    source='dropbox',
                    source_ref=path,
                ).model_dump()
                doc['uploaded_by'] = user.get('name')
                # Embedding
                embed_text = f"{name}\n{doc.get('ai_summary','')}\n{' '.join(doc.get('ai_tags',[]))}"
                doc['embedding'] = await _generate_embedding(embed_text)
                await db.documents.insert_one(doc)
                synced.append({'name': name, 'category': doc['category_slug'], 'doc_type': doc.get('ai_doc_type')})
            except Exception as ee:
                errors.append(f"{entry.get('name','?')}: {str(ee)[:200]}")

    await db.integrations.update_one(
        {'id': 'dropbox'},
        {'$set': {'last_sync': now_iso(), 'last_sync_count': len(synced)}},
        upsert=True,
    )
    return {'ok': True, 'synced_count': len(synced), 'synced': synced[:20], 'errors': errors[:20], 'total_found': len(all_entries)}



# ============== CLIENTS (companies/projects from Simpro, with local fallback) ==============
class Client(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    state: Optional[str] = None
    parent_client_id: Optional[str] = None
    member_ids: List[str] = []  # worker IDs assigned to this client
    location_id: Optional[str] = None
    status: str = 'active'  # active, inactive
    created_by: Optional[str] = None
    simpro_id: Optional[str] = None
    simpro_refs: List[str] = []  # all SimPRO references (cross-company tracking)
    notes: Optional[str] = None
    source: str = 'manual'  # manual, simpro
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/clients')
async def list_clients(search: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if search:
        q['name'] = {'$regex': search, '$options': 'i'}
    items = await db.clients.find(q, {'_id': 0}).sort('name', 1).to_list(2000)
    return items

@api_router.post('/clients')
async def create_client(c: Client, user=Depends(require_admin)):
    doc = c.model_dump()
    doc['created_by'] = user.get('name')
    await db.clients.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/clients/{cid}')
async def update_client(cid: str, c: Client, user=Depends(require_admin)):
    doc = c.model_dump(); doc['id'] = cid
    await db.clients.update_one({'id': cid}, {'$set': doc}, upsert=True)
    return doc

@api_router.get('/clients/{cid}')
async def get_client(cid: str, user=Depends(get_current_user)):
    c = await db.clients.find_one({'id': cid}, {'_id': 0})
    if not c: raise HTTPException(404, 'Not found')
    return c

@api_router.delete('/clients/{cid}')
async def delete_client(cid: str, user=Depends(require_admin)):
    await db.clients.delete_one({'id': cid})
    return {'ok': True}

# ============== SKILLS ==============
class Skill(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    category: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/skills')
async def list_skills(user=Depends(get_current_user)):
    items = await db.skills.find({}, {'_id': 0}).sort('name', 1).to_list(500)
    if not items:
        # seed defaults
        defaults = [
            'Concrete Finishing', 'Excavator Operation', 'Crane Operation', 'Welding',
            'Pipe Laying', 'Asphalt Laying', 'Traffic Control', 'Surveying',
            'Carpentry', 'Steel Fixing', 'Formwork', 'Asbestos Removal (Class B)',
            'Confined Space Entry', 'First Aid', 'Working at Heights',
            'Backhoe Loader', 'Bulldozer', 'Compactor', 'Dump Truck', 'Grader',
            'Forklift', 'EWP (Boom/Scissor)', 'Telehandler', 'Vacuum Truck',
            'HDD (Directional Drilling)',
        ]
        for n in defaults:
            await db.skills.insert_one(Skill(name=n).model_dump())
        items = await db.skills.find({}, {'_id': 0}).sort('name', 1).to_list(500)
    return items

@api_router.post('/skills')
async def create_skill(s: Skill, user=Depends(require_admin)):
    doc = s.model_dump()
    await db.skills.insert_one(doc)
    doc.pop('_id', None)
    return doc

# ============== WORKER AVAILABILITY ==============
@api_router.put('/workers/{wid}/availability')
async def update_availability(wid: str, body: dict, user=Depends(get_current_user)):
    """Update a worker's weekly availability. Body: {monday: {enabled, start, end}, ...}"""
    await db.workers.update_one({'id': wid}, {'$set': {'availability': body.get('availability', body)}})
    w = await db.workers.find_one({'id': wid}, {'_id': 0})
    return w

# ============== SIMPRO INTEGRATION ==============
class SimproConfig(BaseModel):
    base_url: str  # e.g. https://yourcompany.simprosuite.com
    api_token: Optional[str] = None  # legacy basic-token (Bearer)
    company_id: Optional[str] = None  # e.g. "0"
    client_id: Optional[str] = None  # OAuth2 Build option
    client_secret: Optional[str] = None
    # OAuth tokens stored after exchange
    access_token: Optional[str] = None
    refresh_token: Optional[str] = None
    token_expires_at: Optional[str] = None

@api_router.get('/integrations/simpro/status')
async def simpro_status(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg:
        return {'connected': False}
    return {
        'connected': True,
        'base_url': cfg.get('base_url'),
        'company_id': cfg.get('company_id'),
        'has_token': bool(cfg.get('api_token') or cfg.get('access_token')),
        'auth_method': 'oauth2' if cfg.get('client_id') else 'token',
        'last_sync': cfg.get('last_sync'),
        'last_sync_count': cfg.get('last_sync_count', 0),
    }

@api_router.post('/integrations/simpro/connect')
async def simpro_connect(body: SimproConfig, user=Depends(require_admin)):
    """Save Simpro config + verify by hitting /api/v1.0/companies/"""
    base = (body.base_url or '').strip().rstrip('/')
    if not base:
        raise HTTPException(400, 'base_url required')
    # Strip whitespace from all credentials
    api_token = (body.api_token or '').strip() or None
    access_token = (body.access_token or '').strip() or None
    client_id = (body.client_id or '').strip() or None
    client_secret = (body.client_secret or '').strip() or None

    headers = {}
    if api_token: headers['Authorization'] = f'Bearer {api_token}'
    elif access_token: headers['Authorization'] = f'Bearer {access_token}'

    verify_ok = False
    verify_error = None
    if headers:
        try:
            import httpx
            async with httpx.AsyncClient(timeout=15.0) as client:
                r = await client.get(f"{base}/api/v1.0/companies/", headers=headers)
                if r.status_code in (200, 304):
                    verify_ok = True
                else:
                    verify_error = f'HTTP {r.status_code}: {r.text[:200] or "no body"}'
        except Exception as e:
            verify_error = str(e)[:200]

    await db.integrations.update_one(
        {'id': 'simpro'},
        {'$set': {
            'id': 'simpro',
            'base_url': base,
            'api_token': api_token,
            'company_id': (body.company_id or '0').strip(),
            'client_id': client_id,
            'client_secret': client_secret,
            'access_token': access_token,
            'refresh_token': body.refresh_token,
            'token_expires_at': body.token_expires_at,
            'connected_at': now_iso(),
            'verified': verify_ok,
        }},
        upsert=True,
    )
    return {'ok': True, 'verified': verify_ok, 'error': verify_error}

@api_router.delete('/integrations/simpro')
async def simpro_disconnect(user=Depends(require_admin)):
    await db.integrations.delete_one({'id': 'simpro'})
    return {'ok': True}

async def _simpro_get(path: str, params: Optional[dict] = None):
    """Helper to call Simpro API with stored credentials."""
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg:
        raise HTTPException(400, 'Simpro not connected. Configure in Settings → Simpro.')
    base = (cfg['base_url'] or '').strip().rstrip('/')
    token = (cfg.get('api_token') or cfg.get('access_token') or '').strip()
    if not token:
        raise HTTPException(400, 'No Simpro token configured.')
    import httpx
    async with httpx.AsyncClient(timeout=30.0) as client:
        r = await client.get(f"{base}{path}", headers={'Authorization': f'Bearer {token}'}, params=params)
        if r.status_code >= 400:
            raise HTTPException(r.status_code, f'Simpro API error ({r.status_code}): {r.text[:300] or "no body"}')
        return r.json()

@api_router.post('/integrations/simpro/sync/employees')
async def simpro_sync_employees(company_ids: Optional[str] = None, user=Depends(require_admin)):
    """Pull employees from Simpro for one or more companies.
    company_ids: comma-separated list (e.g. '2,3'). Falls back to configured company_id."""
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg:
        raise HTTPException(400, 'Simpro not connected')

    # Determine which companies to sync
    if company_ids:
        target_companies = [c.strip() for c in company_ids.split(',') if c.strip()]
    else:
        target_companies = [cfg.get('company_id', '0')]

    # Map company IDs to friendly names (configurable, with defaults)
    company_name_map = cfg.get('company_name_map') or {
        '2': 'Paneltec Pty Ltd',
        '3': 'Viatec Traffic Solutions',
    }

    synced = []
    errors = []
    for cid in target_companies:
        company_name = company_name_map.get(cid, f'Company {cid}')
        try:
            employees = await _simpro_get(f'/api/v1.0/companies/{cid}/employees/', params={'pageSize': 250})
        except HTTPException as e:
            errors.append(f'Company {cid}: {str(e.detail)[:200]}')
            continue

        if not isinstance(employees, list):
            employees = employees.get('items', []) if isinstance(employees, dict) else []

        for emp_summary in employees:
            emp_id = str(emp_summary.get('ID') or emp_summary.get('id') or '')
            if not emp_id: continue

            # Fetch full employee details (list endpoint only returns basic fields)
            try:
                emp = await _simpro_get(f'/api/v1.0/companies/{cid}/employees/{emp_id}')
            except HTTPException:
                emp = emp_summary  # Fall back to summary if detail fetch fails

            simpro_key = f"{cid}:{emp_id}"  # company-scoped key

            # SimPRO structures: Name is full, PrimaryContact has email/phones, Address has nested fields
            full_name = emp.get('Name') or emp_summary.get('Name') or ''
            first = emp.get('GivenName') or emp.get('FirstName') or ''
            last = emp.get('FamilyName') or emp.get('LastName') or emp.get('Surname') or ''
            name = (first + ' ' + last).strip() or full_name or 'Unknown'
            # If only combined name, split it
            if not first and not last and full_name and ' ' in full_name:
                parts = full_name.split(' ', 1)
                first, last = parts[0], parts[1]
            elif not first and not last and full_name:
                first = full_name

            # PrimaryContact nested object
            primary = emp.get('PrimaryContact') or {}
            email = (primary.get('Email') or primary.get('SecondaryEmail') or
                     emp.get('Email') or emp.get('email') or '')
            phone = (primary.get('CellPhone') or primary.get('WorkPhone') or
                     emp.get('Phone') or emp.get('phone') or '')

            position = emp.get('Position') or emp.get('JobTitle') or emp.get('TypeName') or ''

            # Address nested
            addr = emp.get('Address') or {}
            street_address = addr.get('Address') or ''
            suburb = addr.get('City') or ''
            state_full = addr.get('State') or ''
            postal_code = addr.get('PostalCode') or ''
            country = addr.get('Country') or 'AUSTRALIA'
            # Map state abbreviations to full names
            state_map = {'TAS':'Tasmania','VIC':'Victoria','NSW':'New South Wales','QLD':'Queensland',
                         'SA':'South Australia','WA':'Western Australia','NT':'Northern Territory','ACT':'Australian Capital Territory'}
            state = state_map.get(state_full.upper(), state_full) if state_full else ''

            birth_date = emp.get('DateOfBirth') or None

            existing = await db.workers.find_one({'simpro_id': simpro_key}, {'_id': 0})
            worker_data = {
                'simpro_id': simpro_key,
                'simpro_company_id': cid,
                'simpro_company_name': company_name,
                'name': name,
                'first_name': first or None,
                'last_name': last or None,
                'email': email or None,
                'phone': phone or None,
                'trade': position or None,
                'street_address': street_address or None,
                'suburb': suburb or None,
                'state': state or None,
                'postal_code': postal_code or None,
                'country': country.upper() if country else 'AUSTRALIA',
                'birth_date': birth_date,
                'role': 'worker',
                'source': 'simpro',
                'status': 'inactive' if emp.get('Archived') else 'active',
            }
            if existing:
                worker_data['id'] = existing['id']
                # Preserve manually-set fields like availability, license, client_ids, skills
                preserve = ['availability','license_allocated','license_allocated_by','is_manager','client_ids','skills','location_ids','signature_b64','photo_b64','additional_notes']
                for k in preserve:
                    if existing.get(k) and not worker_data.get(k):
                        worker_data[k] = existing[k]
                await db.workers.update_one({'simpro_id': simpro_key}, {'$set': worker_data})
            else:
                worker_data['id'] = str(uuid.uuid4())
                worker_data['created_at'] = now_iso()
                await db.workers.insert_one(Worker(**worker_data).model_dump())
            synced.append({'name': name, 'email': email, 'phone': phone, 'position': position, 'company': company_name})

    await db.integrations.update_one(
        {'id': 'simpro'},
        {'$set': {'last_sync': now_iso(), 'last_sync_count': len(synced), 'last_sync_type': 'employees'}},
    )
    return {'ok': True, 'synced_count': len(synced), 'synced': synced[:100],
            'errors': errors, 'companies_synced': target_companies}

@api_router.delete('/workers-bulk/seed-data')
async def clear_seed_workers(user=Depends(require_admin)):
    """Remove all non-Simpro (test/seeded/manual) workers. Useful before first Simpro sync."""
    # Delete only workers that are NOT from Simpro
    result = await db.workers.delete_many({
        '$or': [
            {'source': {'$exists': False}},
            {'source': {'$ne': 'simpro'}},
            {'simpro_id': {'$in': [None, '']}},
            {'simpro_id': {'$exists': False}},
        ]
    })
    return {'ok': True, 'deleted': result.deleted_count}

@api_router.post('/integrations/simpro/sync/clients')
async def simpro_sync_clients(company_ids: Optional[str] = None, include_cost_centers: bool = True, user=Depends(require_admin)):
    """Pull customer companies + cost centers from Simpro across one or more companies."""
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg:
        raise HTTPException(400, 'Simpro not connected')

    if company_ids:
        target_companies = [c.strip() for c in company_ids.split(',') if c.strip()]
    else:
        target_companies = [cfg.get('company_id', '0')]

    company_name_map = cfg.get('company_name_map') or {
        '2': 'Paneltec Pty Ltd',
        '3': 'Viatec Traffic Solutions',
    }

    synced = []
    errors = []

    async def _upsert_client(name: str, kind: str, ext_id: str, comp_label: str, cid: str,
                             email: Optional[str] = None, phone: Optional[str] = None):
        """Upsert a client by normalized name. Tracks which SimPRO companies/types reference it.
        Prevents duplicates when the same customer exists in multiple SimPRO companies."""
        norm_name = name.strip()
        if not norm_name: return
        # Look for an existing client with same name OR matching simpro_id
        existing = await db.clients.find_one({
            '$or': [
                {'name': norm_name},
                {'simpro_id': f"{cid}:{kind}:{ext_id}"},
                {'simpro_refs': f"{cid}:{kind}:{ext_id}"},
            ]
        }, {'_id': 0})
        new_ref = f"{cid}:{kind}:{ext_id}"
        if existing:
            refs = set(existing.get('simpro_refs') or [])
            if existing.get('simpro_id'): refs.add(existing['simpro_id'])
            refs.add(new_ref)
            update = {
                'simpro_refs': list(refs),
                'source': 'simpro',
            }
            # Only fill empty fields (preserve manual edits)
            if email and not existing.get('contact_email'): update['contact_email'] = email
            if phone and not existing.get('phone'): update['phone'] = phone
            await db.clients.update_one({'id': existing['id']}, {'$set': update})
            synced.append({'name': norm_name, 'type': kind, 'company': comp_label, 'merged': True})
        else:
            doc = {
                'id': str(uuid.uuid4()),
                'simpro_id': new_ref,
                'simpro_refs': [new_ref],
                'name': norm_name,
                'contact_email': email,
                'phone': phone,
                'source': 'simpro',
                'notes': f"SimPRO {kind} from {comp_label}" if kind == 'cc' else None,
                'created_at': now_iso(),
            }
            await db.clients.insert_one(Client(**doc).model_dump())
            synced.append({'name': norm_name, 'type': kind, 'company': comp_label, 'merged': False})

    for cid in target_companies:
        comp_label = company_name_map.get(cid, f'Company {cid}')

        # 1) Cost centers
        if include_cost_centers:
            try:
                cost_centers = await _simpro_get(f'/api/v1.0/companies/{cid}/setup/accounts/costCenters/', params={'pageSize': 250})
                if not isinstance(cost_centers, list):
                    cost_centers = cost_centers.get('items', []) if isinstance(cost_centers, dict) else []
                for c in cost_centers:
                    await _upsert_client(
                        name=c.get('Name') or 'Unnamed CC',
                        kind='cc',
                        ext_id=str(c.get('ID') or ''),
                        comp_label=comp_label,
                        cid=cid,
                    )
            except HTTPException as e:
                errors.append(f'Company {cid} cost centers: {str(e.detail)[:200]}')

        # 2) Customer companies — paginated
        try:
            page = 1
            while True:
                customers = await _simpro_get(f'/api/v1.0/companies/{cid}/customers/companies/', params={'pageSize': 250, 'page': page})
                if not isinstance(customers, list):
                    customers = customers.get('items', []) if isinstance(customers, dict) else []
                if not customers: break
                for c in customers:
                    await _upsert_client(
                        name=c.get('CompanyName') or c.get('Name') or 'Unnamed',
                        kind='cust',
                        ext_id=str(c.get('ID') or ''),
                        comp_label=comp_label,
                        cid=cid,
                        email=c.get('Email'),
                        phone=c.get('Phone'),
                    )
                if len(customers) < 250: break
                page += 1
                if page > 20: break
        except HTTPException as e:
            errors.append(f'Company {cid} customers: {str(e.detail)[:200]}')

    # Calculate stats
    unique_count = len({s['name'] for s in synced})
    new_count = sum(1 for s in synced if not s.get('merged'))
    merged_count = sum(1 for s in synced if s.get('merged'))

    return {
        'ok': True,
        'synced_count': new_count,
        'merged_count': merged_count,
        'unique_clients': unique_count,
        'total_refs': len(synced),
        'companies_synced': target_companies,
        'errors': errors,
        'synced': synced[:100],
    }

# Seed local fallback clients (so the UI has data before Simpro is connected)
@api_router.post('/clients/seed-defaults')
async def seed_default_clients(user=Depends(require_admin)):
    count = await db.clients.count_documents({})
    if count > 0:
        return {'ok': True, 'note': 'Clients already exist', 'count': count}
    defaults = [
        'Active Tree Services', 'AETV Pty Ltd', 'Andrew Foley Plumbing',
        'Andrew Walters Constructions', 'Anglicare Tasmania Inc.', 'Aus Flight Handling',
        'Barwick Developments Pty Ltd', 'Batchelor Civil Contracting', 'Binc Premix Concrete',
        'Boags', 'Boral Concrete', "Break O'Day Council", 'Bridge Pro Engineering P/L',
        'Burnie City Council', 'TasWater', 'Tasmanian Government', 'Hydro Tasmania',
        'Launceston City Council', 'Devonport City Council', 'Department of State Growth',
    ]
    for n in defaults:
        await db.clients.insert_one(Client(name=n).model_dump())
    return {'ok': True, 'created': len(defaults)}



# ============== CLIENT FOLDERS (per-client document folders) ==============
class ClientFolder(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    name: str
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/client-folders')
async def list_client_folders(client_id: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if client_id: q['client_id'] = client_id
    items = await db.client_folders.find(q, {'_id': 0}).sort('name', 1).to_list(500)
    return items

@api_router.post('/client-folders')
async def create_client_folder(f: ClientFolder, user=Depends(get_current_user)):
    doc = f.model_dump()
    doc['created_by'] = user.get('name')
    await db.client_folders.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/client-folders/{fid}')
async def delete_client_folder(fid: str, user=Depends(get_current_user)):
    await db.client_folders.delete_one({'id': fid})
    return {'ok': True}

# ============== NOTES ==============
class Note(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    body: str = ""
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    folder_id: Optional[str] = None
    file_ref: Optional[str] = None  # optional document_id link
    sub_location: Optional[str] = None
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/notes')
async def list_notes(client_id: Optional[str] = None, folder_id: Optional[str] = None,
                     search: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if client_id: q['client_id'] = client_id
    if folder_id: q['folder_id'] = folder_id
    if search:
        q['$or'] = [
            {'title': {'$regex': search, '$options': 'i'}},
            {'body': {'$regex': search, '$options': 'i'}},
        ]
    items = await db.notes.find(q, {'_id': 0}).sort('created_at', -1).to_list(2000)
    return items

@api_router.get('/notes/{nid}')
async def get_note(nid: str, user=Depends(get_current_user)):
    n = await db.notes.find_one({'id': nid}, {'_id': 0})
    if not n: raise HTTPException(404, 'Not found')
    return n

@api_router.post('/notes')
async def create_note(n: Note, user=Depends(get_current_user)):
    doc = n.model_dump()
    doc['created_by'] = user.get('name')
    if doc.get('client_id') and not doc.get('client_name'):
        c = await db.clients.find_one({'id': doc['client_id']}, {'_id': 0, 'name': 1})
        if c: doc['client_name'] = c['name']
    await db.notes.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/notes/{nid}')
async def update_note(nid: str, n: Note, user=Depends(get_current_user)):
    doc = n.model_dump(); doc['id'] = nid
    await db.notes.update_one({'id': nid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/notes/{nid}')
async def delete_note(nid: str, user=Depends(get_current_user)):
    await db.notes.delete_one({'id': nid})
    return {'ok': True}

# ============== CLIENT TASKS (separate from corrective-action ActionItem) ==============
class ClientTask(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    assignee_id: Optional[str] = None
    assignee_name: Optional[str] = None
    priority: str = 'medium'  # low, medium, high, critical
    status: str = 'open'  # open, in_progress, done, blocked
    due_date: Optional[str] = None
    completed_at: Optional[str] = None
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/tasks')
async def list_tasks(client_id: Optional[str] = None, status: Optional[str] = None,
                     user=Depends(get_current_user)):
    q = {}
    if client_id: q['client_id'] = client_id
    if status and status != 'all': q['status'] = status
    items = await db.client_tasks.find(q, {'_id': 0}).sort('due_date', 1).to_list(2000)
    return items

@api_router.post('/tasks')
async def create_task(t: ClientTask, user=Depends(get_current_user)):
    doc = t.model_dump()
    doc['created_by'] = user.get('name')
    if doc.get('client_id') and not doc.get('client_name'):
        c = await db.clients.find_one({'id': doc['client_id']}, {'_id': 0, 'name': 1})
        if c: doc['client_name'] = c['name']
    if doc.get('assignee_id') and not doc.get('assignee_name'):
        w = await db.workers.find_one({'id': doc['assignee_id']}, {'_id': 0, 'name': 1})
        if w: doc['assignee_name'] = w['name']
    await db.client_tasks.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/tasks/{tid}')
async def update_task(tid: str, t: ClientTask, user=Depends(get_current_user)):
    doc = t.model_dump(); doc['id'] = tid
    if doc.get('status') == 'done' and not doc.get('completed_at'):
        doc['completed_at'] = now_iso()
    await db.client_tasks.update_one({'id': tid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/tasks/{tid}')
async def delete_task(tid: str, user=Depends(get_current_user)):
    await db.client_tasks.delete_one({'id': tid})
    return {'ok': True}

@api_router.post('/tasks/{tid}/status')
async def quick_status(tid: str, body: dict, user=Depends(get_current_user)):
    new_status = body.get('status', 'open')
    update = {'status': new_status}
    if new_status == 'done': update['completed_at'] = now_iso()
    await db.client_tasks.update_one({'id': tid}, {'$set': update})
    t = await db.client_tasks.find_one({'id': tid}, {'_id': 0})
    return t



# ============== EXTENDED SIMPRO CONFIG ==============
class SimproConfigExt(BaseModel):
    base_url: Optional[str] = None
    company_id: Optional[str] = None
    api_token: Optional[str] = None
    staff_custom_field: Optional[str] = None
    staff_field_value: Optional[str] = None
    position_filter: List[str] = []
    sync_interval_minutes: int = 60
    auto_sync: bool = False
    completed_jobs_history_days: int = 90
    incremental_sync: bool = False

@api_router.put('/integrations/simpro/config')
async def simpro_update_config(body: SimproConfigExt, user=Depends(require_admin)):
    update = {k: v for k, v in body.model_dump().items() if v is not None}
    update['id'] = 'simpro'
    await db.integrations.update_one({'id': 'simpro'}, {'$set': update}, upsert=True)
    return {'ok': True}

@api_router.get('/integrations/simpro/companies')
async def simpro_list_companies(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg or not cfg.get('api_token') or not cfg.get('base_url'):
        raise HTTPException(400, 'Configure base_url + api_token first')
    import httpx
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.get(
            f"{cfg['base_url'].rstrip('/')}/api/v1.0/companies/",
            headers={'Authorization': f"Bearer {cfg['api_token']}"},
        )
        if r.status_code != 200:
            raise HTTPException(r.status_code, r.text[:300])
        data = r.json()
    if isinstance(data, dict): data = data.get('items', [])
    return [{'id': c.get('ID'), 'name': c.get('Name')} for c in data]

@api_router.get('/integrations/simpro/custom-fields')
async def simpro_list_custom_fields(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'simpro'}, {'_id': 0})
    if not cfg or not cfg.get('api_token'): raise HTTPException(400, 'Connect Simpro first')
    cid = cfg.get('company_id', '0')
    import httpx
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            r = await client.get(
                f"{cfg['base_url'].rstrip('/')}/api/v1.0/companies/{cid}/setup/customFields/employees/",
                headers={'Authorization': f"Bearer {cfg['api_token']}"},
            )
            if r.status_code != 200:
                return {'items': [], 'error': r.text[:200]}
            data = r.json()
    except Exception as e:
        return {'items': [], 'error': str(e)[:200]}
    items = data if isinstance(data, list) else data.get('items', [])
    return {'items': [{'id': i.get('ID'), 'name': i.get('Name'), 'options': i.get('ListItems', [])} for i in items]}

# ============== NAVIXY GPS ==============
class NavixyConfig(BaseModel):
    email: str
    password: Optional[str] = None
    api_key: Optional[str] = None  # session hash
    api_base_url: str = 'api.us.navixy.com'
    account_id: Optional[str] = None
    tag_filter: List[str] = []
    poll_interval_seconds: int = 30
    auto_poll: bool = False

@api_router.get('/integrations/navixy/status')
async def navixy_status(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'navixy'}, {'_id': 0})
    if not cfg: return {'connected': False}
    return {
        'connected': True,
        'email': cfg.get('email'),
        'api_base_url': cfg.get('api_base_url'),
        'has_key': bool(cfg.get('api_key')),
        'tag_filter': cfg.get('tag_filter', []),
        'poll_interval_seconds': cfg.get('poll_interval_seconds', 30),
        'auto_poll': cfg.get('auto_poll', False),
        'verified': cfg.get('verified', False),
        'last_sync': cfg.get('last_sync'),
    }

@api_router.post('/integrations/navixy/connect')
async def navixy_connect(body: NavixyConfig, user=Depends(require_admin)):
    """Save Navixy config. If password given but no api_key, attempt to fetch session hash."""
    import httpx
    verified = False
    error = None
    api_key = body.api_key
    base = body.api_base_url.replace('https://', '').replace('http://', '').rstrip('/')
    if not api_key and body.password:
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                r = await client.post(
                    f"https://{base}/user/auth",
                    data={'login': body.email, 'password': body.password},
                )
                if r.status_code == 200:
                    data = r.json()
                    if data.get('success'):
                        api_key = data.get('hash')
                        verified = True
                    else: error = data.get('description', 'auth failed')
                else: error = f"HTTP {r.status_code}"
        except Exception as e:
            error = str(e)[:200]
    elif api_key:
        # Verify with user/get_info
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                r = await client.post(f"https://{base}/user/get_info", data={'hash': api_key})
                verified = (r.status_code == 200 and r.json().get('success'))
                if not verified: error = r.json().get('description', 'invalid key')
        except Exception as e:
            error = str(e)[:200]

    await db.integrations.update_one(
        {'id': 'navixy'},
        {'$set': {
            'id': 'navixy',
            'email': body.email,
            'api_key': api_key,
            'api_base_url': base,
            'account_id': body.account_id,
            'tag_filter': body.tag_filter,
            'poll_interval_seconds': body.poll_interval_seconds,
            'auto_poll': body.auto_poll,
            'verified': verified,
            'connected_at': now_iso(),
        }},
        upsert=True,
    )
    return {'ok': True, 'verified': verified, 'error': error, 'api_key': api_key}

@api_router.delete('/integrations/navixy')
async def navixy_disconnect(user=Depends(require_admin)):
    await db.integrations.delete_one({'id': 'navixy'})
    return {'ok': True}

@api_router.post('/integrations/navixy/test')
async def navixy_test(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'navixy'}, {'_id': 0})
    if not cfg or not cfg.get('api_key'): raise HTTPException(400, 'Not connected')
    import httpx
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(
                f"https://{cfg['api_base_url']}/user/get_info",
                data={'hash': cfg['api_key']},
            )
            data = r.json()
            return {'ok': data.get('success', False), 'detail': data}
    except Exception as e:
        return {'ok': False, 'error': str(e)[:200]}

@api_router.get('/integrations/navixy/tags')
async def navixy_list_tags(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'navixy'}, {'_id': 0})
    if not cfg or not cfg.get('api_key'): raise HTTPException(400, 'Not connected')
    import httpx
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(
                f"https://{cfg['api_base_url']}/tag/list",
                data={'hash': cfg['api_key']},
            )
            data = r.json()
            if data.get('success'):
                return {'tags': [{'id': t.get('id'), 'name': t.get('name'), 'color': t.get('color')} for t in data.get('list', [])]}
            return {'tags': [], 'error': data.get('description')}
    except Exception as e:
        return {'tags': [], 'error': str(e)[:200]}

@api_router.get('/integrations/navixy/trackers')
async def navixy_list_trackers(user=Depends(require_admin)):
    """Return current vehicle positions (filtered by tag if configured)."""
    cfg = await db.integrations.find_one({'id': 'navixy'}, {'_id': 0})
    if not cfg or not cfg.get('api_key'): raise HTTPException(400, 'Not connected')
    import httpx
    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            # List trackers
            r = await client.post(
                f"https://{cfg['api_base_url']}/tracker/list",
                data={'hash': cfg['api_key']},
            )
            data = r.json()
            if not data.get('success'):
                return {'trackers': [], 'error': data.get('description')}
            trackers = data.get('list', [])
            # Filter by tag if set
            tag_filter = cfg.get('tag_filter', [])
            if tag_filter:
                trackers = [t for t in trackers if any(tid in (t.get('tag_bindings') or []) for tid in tag_filter)]
            # Get latest positions
            ids = [t['id'] for t in trackers]
            if ids:
                r2 = await client.post(
                    f"https://{cfg['api_base_url']}/tracker/get_states",
                    data={'hash': cfg['api_key'], 'trackers': str(ids).replace("'", '"')},
                )
                states_data = r2.json()
                states_map = {s['source_id']: s for s in states_data.get('states', [])}
                for t in trackers:
                    s = states_map.get(t['id'], {})
                    gps = s.get('gps', {})
                    t['latitude'] = gps.get('location', {}).get('lat')
                    t['longitude'] = gps.get('location', {}).get('lng')
                    t['speed'] = gps.get('speed')
                    t['updated'] = s.get('actual_track_update')
                    t['movement'] = s.get('movement_status')
        await db.integrations.update_one({'id': 'navixy'}, {'$set': {'last_sync': now_iso()}})
        return {'trackers': trackers}
    except Exception as e:
        return {'trackers': [], 'error': str(e)[:200]}

# ============== TEXTMAGIC SMS ==============
class TextmagicConfig(BaseModel):
    api_username: str
    api_key: str
    default_sender: Optional[str] = None

@api_router.get('/integrations/textmagic/status')
async def textmagic_status(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'textmagic'}, {'_id': 0})
    if not cfg: return {'connected': False}
    return {
        'connected': True,
        'api_username': cfg.get('api_username'),
        'default_sender': cfg.get('default_sender'),
        'verified': cfg.get('verified', False),
        'last_sent': cfg.get('last_sent'),
        'webhook_delivery': f"{os.environ.get('PUBLIC_BASE_URL', '')}/api/integrations/textmagic/webhook/delivery",
        'webhook_inbound': f"{os.environ.get('PUBLIC_BASE_URL', '')}/api/integrations/textmagic/webhook/inbound",
    }

@api_router.post('/integrations/textmagic/connect')
async def textmagic_connect(body: TextmagicConfig, user=Depends(require_admin)):
    import httpx
    verified = False
    error = None
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.get(
                'https://rest.textmagic.com/api/v2/user',
                headers={'X-TM-Username': body.api_username, 'X-TM-Key': body.api_key},
            )
            verified = (r.status_code == 200)
            if not verified: error = r.text[:200]
    except Exception as e:
        error = str(e)[:200]
    await db.integrations.update_one(
        {'id': 'textmagic'},
        {'$set': {
            'id': 'textmagic',
            'api_username': body.api_username,
            'api_key': body.api_key,
            'default_sender': body.default_sender,
            'verified': verified,
            'connected_at': now_iso(),
        }},
        upsert=True,
    )
    return {'ok': True, 'verified': verified, 'error': error}

@api_router.delete('/integrations/textmagic')
async def textmagic_disconnect(user=Depends(require_admin)):
    await db.integrations.delete_one({'id': 'textmagic'})
    return {'ok': True}

class SendSMSIn(BaseModel):
    phones: List[str]
    text: str

@api_router.post('/integrations/textmagic/send')
async def textmagic_send(body: SendSMSIn, user=Depends(get_current_user)):
    cfg = await db.integrations.find_one({'id': 'textmagic'}, {'_id': 0})
    if not cfg: raise HTTPException(400, 'Textmagic not configured')
    import httpx
    payload = {'phones': ','.join(body.phones), 'text': body.text}
    if cfg.get('default_sender'): payload['from'] = cfg['default_sender']
    async with httpx.AsyncClient(timeout=20.0) as client:
        r = await client.post(
            'https://rest.textmagic.com/api/v2/messages',
            headers={'X-TM-Username': cfg['api_username'], 'X-TM-Key': cfg['api_key']},
            data=payload,
        )
        result = {'status_code': r.status_code, 'response': r.json() if r.headers.get('content-type','').startswith('application/json') else r.text}
    await db.integrations.update_one({'id': 'textmagic'}, {'$set': {'last_sent': now_iso()}})
    # Log
    await db.sms_log.insert_one({
        'id': str(uuid.uuid4()),
        'to': body.phones, 'text': body.text,
        'sent_by': user.get('name'),
        'sent_at': now_iso(),
        'status': 'sent' if r.status_code < 400 else 'failed',
        'response': str(result.get('response'))[:500],
    })
    return result

@api_router.get('/integrations/textmagic/log')
async def textmagic_log(user=Depends(get_current_user)):
    return await db.sms_log.find({}, {'_id': 0}).sort('sent_at', -1).limit(100).to_list(100)

@api_router.post('/integrations/textmagic/webhook/delivery')
async def textmagic_webhook_delivery(body: dict):
    await db.sms_callbacks.insert_one({
        'id': str(uuid.uuid4()),
        'type': 'delivery',
        'payload': body,
        'received_at': now_iso(),
    })
    return {'ok': True}

@api_router.post('/integrations/textmagic/webhook/inbound')
async def textmagic_webhook_inbound(body: dict):
    await db.sms_callbacks.insert_one({
        'id': str(uuid.uuid4()),
        'type': 'inbound',
        'payload': body,
        'received_at': now_iso(),
    })
    return {'ok': True}

# ============== MICROSOFT 365 EMAIL (Graph SendMail) ==============
class M365Config(BaseModel):
    tenant_id: str
    client_id: str
    client_secret: str
    send_from_mailbox: str
    reply_to: Optional[str] = None

@api_router.get('/integrations/m365/status')
async def m365_status(user=Depends(require_admin)):
    cfg = await db.integrations.find_one({'id': 'm365'}, {'_id': 0})
    if not cfg: return {'connected': False}
    return {
        'connected': True,
        'tenant_id': cfg.get('tenant_id'),
        'client_id': cfg.get('client_id'),
        'send_from_mailbox': cfg.get('send_from_mailbox'),
        'reply_to': cfg.get('reply_to'),
        'verified': cfg.get('verified', False),
        'last_sent': cfg.get('last_sent'),
    }

async def _m365_get_token():
    cfg = await db.integrations.find_one({'id': 'm365'}, {'_id': 0})
    if not cfg: raise HTTPException(400, 'M365 not configured')
    # Check cache
    cache_exp = cfg.get('token_expires_at')
    if cfg.get('access_token') and cache_exp:
        try:
            if datetime.fromisoformat(cache_exp) > datetime.now(timezone.utc) + timedelta(minutes=2):
                return cfg['access_token']
        except: pass
    # Get new token
    import httpx
    async with httpx.AsyncClient(timeout=15.0) as client:
        r = await client.post(
            f"https://login.microsoftonline.com/{cfg['tenant_id']}/oauth2/v2.0/token",
            data={
                'client_id': cfg['client_id'],
                'client_secret': cfg['client_secret'],
                'scope': 'https://graph.microsoft.com/.default',
                'grant_type': 'client_credentials',
            },
        )
        if r.status_code != 200:
            raise HTTPException(r.status_code, f"M365 token error: {r.text[:300]}")
        td = r.json()
        token = td['access_token']
        exp = (datetime.now(timezone.utc) + timedelta(seconds=td.get('expires_in', 3600))).isoformat()
        await db.integrations.update_one({'id': 'm365'}, {'$set': {'access_token': token, 'token_expires_at': exp}})
        return token

@api_router.post('/integrations/m365/connect')
async def m365_connect(body: M365Config, user=Depends(require_admin)):
    await db.integrations.update_one(
        {'id': 'm365'},
        {'$set': {
            'id': 'm365',
            'tenant_id': body.tenant_id,
            'client_id': body.client_id,
            'client_secret': body.client_secret,
            'send_from_mailbox': body.send_from_mailbox,
            'reply_to': body.reply_to,
            'connected_at': now_iso(),
            'access_token': None, 'token_expires_at': None,  # reset cache
        }},
        upsert=True,
    )
    # Try token
    verified = False
    error = None
    try:
        await _m365_get_token()
        verified = True
    except HTTPException as e:
        error = str(e.detail)
    except Exception as e:
        error = str(e)[:200]
    await db.integrations.update_one({'id': 'm365'}, {'$set': {'verified': verified}})
    return {'ok': True, 'verified': verified, 'error': error}

@api_router.delete('/integrations/m365')
async def m365_disconnect(user=Depends(require_admin)):
    await db.integrations.delete_one({'id': 'm365'})
    return {'ok': True}

class SendEmailIn(BaseModel):
    to: List[str]
    subject: str
    body: str  # html
    cc: List[str] = []

async def _send_m365_email(payload: SendEmailIn) -> dict:
    cfg = await db.integrations.find_one({'id': 'm365'}, {'_id': 0})
    if not cfg: return {'ok': False, 'error': 'M365 not configured'}
    token = await _m365_get_token()
    mailbox = cfg['send_from_mailbox']
    reply_to = cfg.get('reply_to')
    msg = {
        'message': {
            'subject': payload.subject,
            'body': {'contentType': 'HTML', 'content': payload.body},
            'toRecipients': [{'emailAddress': {'address': e}} for e in payload.to],
        },
        'saveToSentItems': True,
    }
    if payload.cc:
        msg['message']['ccRecipients'] = [{'emailAddress': {'address': e}} for e in payload.cc]
    if reply_to:
        msg['message']['replyTo'] = [{'emailAddress': {'address': reply_to}}]
    import httpx
    async with httpx.AsyncClient(timeout=30.0) as client:
        r = await client.post(
            f"https://graph.microsoft.com/v1.0/users/{mailbox}/sendMail",
            headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
            json=msg,
        )
        ok = (r.status_code in (200, 202))
        return {'ok': ok, 'status_code': r.status_code, 'response': r.text[:500]}

@api_router.post('/integrations/m365/send')
async def m365_send(body: SendEmailIn, user=Depends(get_current_user)):
    result = await _send_m365_email(body)
    await db.email_log.insert_one({
        'id': str(uuid.uuid4()),
        'to': body.to, 'subject': body.subject,
        'sent_by': user.get('name'),
        'sent_at': now_iso(),
        'status': 'sent' if result.get('ok') else 'failed',
        'response': str(result)[:500],
    })
    if result.get('ok'):
        await db.integrations.update_one({'id': 'm365'}, {'$set': {'last_sent': now_iso()}})
    return result

@api_router.get('/integrations/m365/log')
async def m365_log(user=Depends(get_current_user)):
    return await db.email_log.find({}, {'_id': 0}).sort('sent_at', -1).limit(100).to_list(100)



# ============== PROJECT BOOKINGS ==============
class ProjectBooking(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    member_id: str
    member_name: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    location_id: Optional[str] = None
    sub_location: Optional[str] = None
    start_date: str  # YYYY-MM-DD
    end_date: str
    notes: Optional[str] = None
    status: str = 'active'  # active, completed, cancelled
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/project-bookings')
async def list_project_bookings(member_id: Optional[str] = None, client_id: Optional[str] = None,
                                 user=Depends(get_current_user)):
    q = {}
    if member_id: q['member_id'] = member_id
    if client_id: q['client_id'] = client_id
    items = await db.project_bookings.find(q, {'_id': 0}).sort('start_date', -1).to_list(1000)
    return items

@api_router.post('/project-bookings')
async def create_project_booking(b: ProjectBooking, user=Depends(get_current_user)):
    doc = b.model_dump()
    doc['created_by'] = user.get('name')
    if doc.get('member_id') and not doc.get('member_name'):
        w = await db.workers.find_one({'id': doc['member_id']}, {'_id': 0, 'name': 1})
        if w: doc['member_name'] = w['name']
    if doc.get('client_id') and not doc.get('client_name'):
        c = await db.clients.find_one({'id': doc['client_id']}, {'_id': 0, 'name': 1})
        if c: doc['client_name'] = c['name']
    await db.project_bookings.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/project-bookings/{bid}')
async def update_project_booking(bid: str, b: ProjectBooking, user=Depends(get_current_user)):
    doc = b.model_dump(); doc['id'] = bid
    await db.project_bookings.update_one({'id': bid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/project-bookings/{bid}')
async def delete_project_booking(bid: str, user=Depends(get_current_user)):
    await db.project_bookings.delete_one({'id': bid})
    return {'ok': True}

# ============== ALLOCATIONS (Day-level worker rostering) ==============
class Allocation(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    member_id: str
    member_name: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    skill: Optional[str] = None
    booking_date: str  # YYYY-MM-DD
    start_time: str = '07:00'  # HH:MM
    end_time: str = '15:00'
    sub_location: Optional[str] = None
    notes: Optional[str] = None
    notified_at: Optional[str] = None
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/allocations')
async def list_allocations(start: Optional[str] = None, end: Optional[str] = None,
                            member_id: Optional[str] = None, client_id: Optional[str] = None,
                            user=Depends(get_current_user)):
    q = {}
    if start and end:
        q['booking_date'] = {'$gte': start, '$lte': end}
    elif start:
        q['booking_date'] = {'$gte': start}
    if member_id: q['member_id'] = member_id
    if client_id: q['client_id'] = client_id
    items = await db.allocations.find(q, {'_id': 0}).sort('booking_date', 1).to_list(3000)
    return items

@api_router.post('/allocations')
async def create_allocation(a: Allocation, user=Depends(get_current_user)):
    doc = a.model_dump()
    doc['created_by'] = user.get('name')
    if doc.get('member_id') and not doc.get('member_name'):
        w = await db.workers.find_one({'id': doc['member_id']}, {'_id': 0, 'name': 1})
        if w: doc['member_name'] = w['name']
    if doc.get('client_id') and not doc.get('client_name'):
        c = await db.clients.find_one({'id': doc['client_id']}, {'_id': 0, 'name': 1})
        if c: doc['client_name'] = c['name']
    await db.allocations.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/allocations/{aid}')
async def update_allocation(aid: str, a: Allocation, user=Depends(get_current_user)):
    doc = a.model_dump(); doc['id'] = aid
    await db.allocations.update_one({'id': aid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/allocations/{aid}')
async def delete_allocation(aid: str, user=Depends(get_current_user)):
    await db.allocations.delete_one({'id': aid})
    return {'ok': True}

@api_router.post('/allocations/notify')
async def notify_allocations(body: dict, user=Depends(require_admin)):
    """Send SMS + Email notifications for all allocations in a date range.
    Body: {start: 'YYYY-MM-DD', end: 'YYYY-MM-DD'}"""
    start = body.get('start')
    end = body.get('end')
    if not start or not end:
        raise HTTPException(400, 'start and end dates required')

    allocs = await db.allocations.find({
        'booking_date': {'$gte': start, '$lte': end}
    }, {'_id': 0}).to_list(2000)

    # Group by member
    by_member = {}
    for a in allocs:
        if a.get('member_id'):
            by_member.setdefault(a['member_id'], []).append(a)

    workers = {w['id']: w for w in await db.workers.find({}, {'_id': 0}).to_list(2000)}
    sms_sent = 0; email_sent = 0; errors = []

    textmagic_ready = bool(await db.integrations.find_one({'id': 'textmagic', 'verified': True}))
    m365_ready = bool(await db.integrations.find_one({'id': 'm365', 'verified': True}))

    for member_id, allocs_for_member in by_member.items():
        worker = workers.get(member_id)
        if not worker: continue

        # Build a readable schedule
        lines = ["Hi " + (worker.get('first_name') or worker.get('name', '')) + ", your Paneltec schedule:"]
        for a in sorted(allocs_for_member, key=lambda x: x['booking_date']):
            line = f"  {a['booking_date']} {a.get('start_time','')}-{a.get('end_time','')}"
            if a.get('client_name'): line += f" @ {a['client_name']}"
            if a.get('skill'): line += f" ({a['skill']})"
            lines.append(line)
        sms_body = '\n'.join(lines)[:1000]
        html_body = '<h3>Your Paneltec Schedule</h3><table border=1 cellpadding=5 style="border-collapse:collapse"><tr><th>Date</th><th>Time</th><th>Client</th><th>Skill</th><th>Notes</th></tr>'
        for a in sorted(allocs_for_member, key=lambda x: x['booking_date']):
            html_body += f"<tr><td>{a['booking_date']}</td><td>{a.get('start_time','')}-{a.get('end_time','')}</td><td>{a.get('client_name') or '-'}</td><td>{a.get('skill') or '-'}</td><td>{a.get('notes') or ''}</td></tr>"
        html_body += '</table>'

        # SMS
        phone = worker.get('phone')
        if phone and textmagic_ready:
            try:
                # Reuse the existing send function
                cfg = await db.integrations.find_one({'id': 'textmagic'}, {'_id': 0})
                import httpx
                async with httpx.AsyncClient(timeout=20.0) as client:
                    r = await client.post(
                        'https://rest.textmagic.com/api/v2/messages',
                        headers={'X-TM-Username': cfg['api_username'], 'X-TM-Key': cfg['api_key']},
                        data={'phones': phone.replace(' ', ''), 'text': sms_body,
                              'from': cfg.get('default_sender', 'PANELTEC')},
                    )
                    if r.status_code < 400: sms_sent += 1
                    else: errors.append(f"SMS to {worker.get('name')}: HTTP {r.status_code}")
            except Exception as e:
                errors.append(f"SMS to {worker.get('name')}: {str(e)[:100]}")

        # Email
        email = worker.get('email')
        if email and m365_ready:
            try:
                result = await _send_m365_email(SendEmailIn(
                    to=[email],
                    subject=f"[Paneltec] Your schedule {start} to {end}",
                    body=html_body,
                ))
                if result.get('ok'): email_sent += 1
                else: errors.append(f"Email to {worker.get('name')}: {result.get('response','')[:80]}")
            except Exception as e:
                errors.append(f"Email to {worker.get('name')}: {str(e)[:100]}")

        # Mark allocations as notified
        for a in allocs_for_member:
            await db.allocations.update_one({'id': a['id']}, {'$set': {'notified_at': now_iso()}})

    return {'ok': True, 'sms_sent': sms_sent, 'email_sent': email_sent,
            'members_notified': len(by_member), 'errors': errors[:20]}

class NotifSchedule(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    range_start: str
    range_end: str
    days_of_week: List[str] = []  # ['monday','tuesday'...]
    cadence: str = 'specific_day'  # previous_day, specific_day, weekly, monthly
    scope: str = 'this_week'  # this_week, all_weeks
    notification_time: str = '06:00'
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/allocations/notif-schedules')
async def list_notif_schedules(user=Depends(get_current_user)):
    return await db.notif_schedules.find({}, {'_id': 0}).sort('created_at', -1).to_list(200)

@api_router.post('/allocations/notif-schedules')
async def create_notif_schedule(s: NotifSchedule, user=Depends(require_admin)):
    doc = s.model_dump()
    doc['created_by'] = user.get('name')
    await db.notif_schedules.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.delete('/allocations/notif-schedules/{sid}')
async def delete_notif_schedule(sid: str, user=Depends(require_admin)):
    await db.notif_schedules.delete_one({'id': sid})
    return {'ok': True}

@api_router.post('/allocations/email-pdf')
async def email_allocations_pdf(body: dict, user=Depends(require_admin)):
    """Generate a PDF of the schedule and email it to a list of recipients."""
    start = body.get('start')
    end = body.get('end')
    recipients = body.get('recipients', [user.get('email')] if user.get('email') else [])
    if not start or not end:
        raise HTTPException(400, 'start and end required')
    if not recipients:
        raise HTTPException(400, 'recipients required')

    allocs = await db.allocations.find({'booking_date': {'$gte': start, '$lte': end}}, {'_id': 0}).sort('booking_date', 1).to_list(2000)

    # Build PDF
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=0.4*inch, rightMargin=0.4*inch, topMargin=0.4*inch, bottomMargin=0.4*inch)
    styles = getSampleStyleSheet()
    PT_YELLOW = colors.HexColor('#FBBF24')
    story = []
    story.append(Paragraph(f'<font color="#0B0B0F"><b>PANELTEC</b></font> <font color="#6B7280">SAFETY PORTAL</font>', styles['Title']))
    story.append(Paragraph(f'Allocation Schedule: <b>{start}</b> to <b>{end}</b>', styles['Heading2']))
    story.append(Spacer(1, 10))

    if allocs:
        data = [['Date', 'Time', 'Member', 'Client', 'Skill', 'Sub Loc', 'Notes']]
        for a in allocs:
            data.append([
                a.get('booking_date',''),
                f"{a.get('start_time','')}-{a.get('end_time','')}",
                a.get('member_name','')[:25],
                (a.get('client_name','') or '')[:25],
                (a.get('skill','') or '')[:20],
                (a.get('sub_location','') or '')[:15],
                (a.get('notes','') or '')[:30],
            ])
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), PT_YELLOW),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 0.3, colors.HexColor('#CBD5E1')),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
        ]))
        story.append(tbl)
    else:
        story.append(Paragraph('<i>No allocations in this date range.</i>', styles['Normal']))

    doc.build(story)
    buf.seek(0)
    pdf_b64 = base64.b64encode(buf.read()).decode()

    # Email it (attachment via Graph API)
    cfg = await db.integrations.find_one({'id': 'm365'}, {'_id': 0})
    if not cfg:
        return {'ok': False, 'error': 'M365 not configured. Connect in Integrations → M365.'}
    token = await _m365_get_token()
    msg = {
        'message': {
            'subject': f'[Paneltec] Allocation Schedule {start} to {end}',
            'body': {'contentType': 'HTML', 'content': f'<p>Attached is the Paneltec allocation schedule from {start} to {end}.</p>'},
            'toRecipients': [{'emailAddress': {'address': e}} for e in recipients],
            'attachments': [{
                '@odata.type': '#microsoft.graph.fileAttachment',
                'name': f'paneltec-schedule-{start}-to-{end}.pdf',
                'contentType': 'application/pdf',
                'contentBytes': pdf_b64,
            }],
        }
    }
    import httpx
    async with httpx.AsyncClient(timeout=30.0) as client:
        r = await client.post(
            f"https://graph.microsoft.com/v1.0/users/{cfg['send_from_mailbox']}/sendMail",
            headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
            json=msg,
        )
        ok = r.status_code in (200, 202)
        return {'ok': ok, 'recipients': recipients, 'response': r.text[:300] if not ok else None}

# ============== PERSONNEL REQUIRED ==============
class PersonnelRequired(BaseModel):
    model_config = ConfigDict(extra='ignore')
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    required_date: str  # YYYY-MM-DD
    location_id: Optional[str] = None
    location_name: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    skill: str
    number_required: int = 1
    number_filled: int = 0
    notes: Optional[str] = None
    status: str = 'open'  # open, filled, cancelled
    created_by: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)

@api_router.get('/personnel-required')
async def list_personnel_required(skill: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if skill: q['skill'] = {'$regex': skill, '$options': 'i'}
    items = await db.personnel_required.find(q, {'_id': 0}).sort('required_date', 1).to_list(1000)
    return items

@api_router.post('/personnel-required')
async def create_personnel_required(p: PersonnelRequired, user=Depends(get_current_user)):
    doc = p.model_dump()
    doc['created_by'] = user.get('name')
    await db.personnel_required.insert_one(doc)
    doc.pop('_id', None)
    return doc

@api_router.put('/personnel-required/{pid}')
async def update_personnel_required(pid: str, p: PersonnelRequired, user=Depends(get_current_user)):
    doc = p.model_dump(); doc['id'] = pid
    await db.personnel_required.update_one({'id': pid}, {'$set': doc}, upsert=True)
    return doc

@api_router.delete('/personnel-required/{pid}')
async def delete_personnel_required(pid: str, user=Depends(get_current_user)):
    await db.personnel_required.delete_one({'id': pid})
    return {'ok': True}

# ============== NOTE EXTENSIONS (member_ids, file_refs, additional_notes, sub_location) ==============
@api_router.put('/notes/{nid}/members')
async def update_note_members(nid: str, body: dict, user=Depends(get_current_user)):
    await db.notes.update_one({'id': nid}, {'$set': {'member_ids': body.get('member_ids', [])}})
    return {'ok': True}

@api_router.put('/notes/{nid}/files')
async def update_note_files(nid: str, body: dict, user=Depends(get_current_user)):
    await db.notes.update_one({'id': nid}, {'$set': {'file_refs': body.get('file_refs', [])}})
    return {'ok': True}

@api_router.put('/notes/{nid}/additional')
async def update_note_additional(nid: str, body: dict, user=Depends(get_current_user)):
    await db.notes.update_one({'id': nid}, {'$set': {'additional_notes_list': body.get('additional_notes_list', [])}})
    return {'ok': True}



@api_router.get('/')
async def root():
    return {'message': 'Paneltec Safety Portal API', 'status': 'ok'}

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@app.on_event('startup')
async def startup_seed():
    count = await db.users.count_documents({})
    if count == 0:
        logger.info('No users found, auto-seeding demo data...')
        try:
            await seed_demo()
            logger.info('Seed complete.')
        except Exception as e:
            logger.error(f'Seed failed: {e}')
    risk_count = await db.risks.count_documents({})
    if risk_count == 0:
        logger.info('Seeding extras (risks, actions, chemicals, assets, contractors)...')
        try:
            await seed_extras()
        except Exception as e:
            logger.error(f'Extras seed failed: {e}')
    swms_count = await db.swms.count_documents({})
    if swms_count == 0:
        logger.info('Seeding compliance (SWMS, ITPs, Permits, Env, Insurance, Audits)...')
        try:
            await seed_compliance()
        except Exception as e:
            logger.error(f'Compliance seed failed: {e}')
    # Seed default clients/skills
    if await db.clients.count_documents({}) == 0:
        try:
            defaults = [
                'Active Tree Services', 'AETV Pty Ltd', 'Andrew Foley Plumbing',
                'Andrew Walters Constructions', 'Anglicare Tasmania Inc.', 'Aus Flight Handling',
                'Barwick Developments Pty Ltd', 'Batchelor Civil Contracting', 'Binc Premix Concrete',
                'Boags', 'Boral Concrete', "Break O'Day Council", 'Bridge Pro Engineering P/L',
                'Burnie City Council', 'TasWater', 'Tasmanian Government', 'Hydro Tasmania',
                'Launceston City Council', 'Devonport City Council', 'Department of State Growth',
            ]
            for n in defaults:
                await db.clients.insert_one(Client(name=n).model_dump())
            logger.info(f'Seeded {len(defaults)} default clients')
        except Exception as e:
            logger.error(f'Client seed failed: {e}')

@app.on_event('shutdown')
async def shutdown_db_client():
    client.close()
