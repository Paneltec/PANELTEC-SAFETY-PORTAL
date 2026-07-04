"""Phase 3.10 — Universal PDF preview for any Document Library file.

Endpoints (all prefixed `/api`):
    GET  /files/{id}/pdf        — stream PDF (inline or attachment via ?dl=1)
    GET  /files/{id}/pdf.pdf    — same, path-disguised variant for ad-blocker
                                  compatibility (mirrors `forms_pdf` pattern)
    POST /files/pdf-bundle      — concatenate multiple files into one PDF
    POST /admin/install-libreoffice — admin-only install hook (stub trigger)

Conversion pipeline (Pragmatic Phase A — LibreOffice **not** installed):
    application/pdf                         → passthrough
    image/jpeg|png|webp                     → Pillow + reportlab A4 fit-to-page
    image/heic|heif                         → pillow-heif → JPG → reportlab
    text/csv | text/plain | text/markdown   → reportlab monospace paginated
    .docx (Word)                            → docx2pdf (best effort) → if <1KB
                                              output fall back to python-docx
                                              + reportlab plain-text renderer
                                              (lossy but never blank)
    .xlsx / .pptx / .odt / .rtf / other     → 415 "LibreOffice not installed —
                                              PDF preview not available for
                                              this format"

Cache: converted PDFs live in `doc_files_pdf_cache` keyed by
       (file_id, sha1, pipeline). Cache miss writes the row; subsequent calls
       stream from cache. Invalidated when the source file is replaced
       (different sha1).
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import hmac
import io
import json
import logging
import os
import subprocess
import tempfile
import time
import uuid
from collections import deque
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, Request, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import (
    Paragraph, Preformatted, SimpleDocTemplate, Spacer,
)

from db import db
from models import new_id, now_iso  # noqa: E402  — Phase 3.14 OCR-index timestamp helper
from auth import get_current_user


log = logging.getLogger("paneltec.files.pdf")
router = APIRouter(prefix="", tags=["files-pdf"])

UPLOAD_DIR = Path(__file__).parent / "uploads" / "document_library"
PIPELINES = {
    "passthrough", "image", "heic", "text",
    "docx_libreoffice", "docx_docx2pdf", "docx_text_fallback",
    "xlsx_libreoffice", "pptx_libreoffice", "odt_libreoffice", "rtf_libreoffice",
}

# Phase 3.13 — LibreOffice primary path. Override via env to fault-test.
LIBREOFFICE_BIN_OVERRIDE = os.environ.get("PANELTEC_LIBREOFFICE_BIN")
LIBREOFFICE_TIMEOUT_S = int(os.environ.get("PANELTEC_LIBREOFFICE_TIMEOUT_S", "60"))


def _libreoffice_binary() -> str | None:
    """Resolve the soffice/libreoffice executable, honouring an env override.
    Set PANELTEC_LIBREOFFICE_BIN to a non-existent path during fault tests to
    force the pragmatic fallback."""
    import shutil
    if LIBREOFFICE_BIN_OVERRIDE is not None:
        return LIBREOFFICE_BIN_OVERRIDE if Path(LIBREOFFICE_BIN_OVERRIDE).exists() else None
    return shutil.which("soffice") or shutil.which("libreoffice")


def _libreoffice_to_pdf(src_path: Path, out_dir: Path, timeout: int = LIBREOFFICE_TIMEOUT_S) -> Path:
    """Convert an office doc → PDF via headless LibreOffice. Returns the
    output PDF path or raises. Each call gets its own UserInstallation profile
    so concurrent requests don't clobber each other's lockfiles."""
    bin_path = _libreoffice_binary()
    if not bin_path:
        raise RuntimeError("LibreOffice not installed")
    profile = out_dir / f"_lo_profile_{os.getpid()}_{int(time.time()*1000)}"
    profile.mkdir(parents=True, exist_ok=True)
    cmd = [
        bin_path, "--headless",
        f"-env:UserInstallation=file://{profile}",
        "--convert-to", "pdf", "--outdir", str(out_dir), str(src_path),
    ]
    res = subprocess.run(cmd, capture_output=True, timeout=timeout)
    if res.returncode != 0:
        tail = (res.stderr or res.stdout or b"").decode("utf-8", errors="replace")[-300:]
        raise RuntimeError(f"libreoffice rc={res.returncode}: {tail}")
    expected = out_dir / (src_path.stem + ".pdf")
    if not expected.exists():
        raise RuntimeError("LibreOffice produced no PDF output file")
    return expected


def _office_to_pdf_via_lo(blob: bytes, ext: str, name: str) -> bytes:
    """Pipe blob → temp file → LibreOffice → PDF bytes. Raises on any
    failure; caller decides whether to fall back."""
    with tempfile.TemporaryDirectory() as td:
        td_p = Path(td)
        src = td_p / f"in.{ext.lstrip('.')}"
        src.write_bytes(blob)
        out = _libreoffice_to_pdf(src, td_p)
        return out.read_bytes()


# ────────────────── helpers ──────────────────

def _sha1_file(p: Path) -> str:
    h = hashlib.sha1()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _is_pdf(blob: bytes) -> bool:
    return blob[:5] == b"%PDF-"


async def _resolve_file(file_id: str, user: dict) -> tuple[dict, Path]:
    doc = await db.doc_files.find_one(
        {"id": file_id, "org_id": user["org_id"], "deleted_at": None},
        {"_id": 0},
    )
    if not doc:
        raise HTTPException(404, "File not found")
    path = UPLOAD_DIR / doc["folder_id"] / doc["stored_name"]
    if not path.exists():
        raise HTTPException(404, "File missing on disk")
    return doc, path


def _pipeline_for(mime: str, name: str) -> str:
    m = (mime or "").lower()
    n = (name or "").lower()
    if m == "application/pdf" or n.endswith(".pdf"):              return "passthrough"
    if m in {"image/jpeg", "image/png", "image/webp"} or n.split(".")[-1] in {"jpg", "jpeg", "png", "webp"}:
        return "image"
    if m in {"image/heic", "image/heif"} or n.endswith((".heic", ".heif")):
        return "heic"
    if m in {"text/csv", "text/plain", "text/markdown"} or n.endswith((".csv", ".txt", ".md")):
        return "text"
    # Phase 3.13 — LibreOffice primary path for all office formats.
    # .docx still has a pragmatic ReportLab text fallback for ultra-defensive
    # delivery; xlsx/pptx/odt/rtf are LO-only (raises 415 on LO failure).
    if n.endswith(".docx") or m == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx_libreoffice"
    if n.endswith(".xlsx") or m == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return "xlsx_libreoffice"
    if n.endswith(".pptx") or m == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
        return "pptx_libreoffice"
    if n.endswith(".odt") or m == "application/vnd.oasis.opendocument.text":
        return "odt_libreoffice"
    if n.endswith(".rtf") or m in {"application/rtf", "text/rtf"}:
        return "rtf_libreoffice"
    return ""  # unsupported


# ────────────────── conversion implementations ──────────────────

def _img_to_pdf(blob: bytes) -> bytes:
    """Wrap a JPG/PNG/WEBP in a single-page A4 PDF, fit-to-page with margins."""
    from PIL import Image
    img = Image.open(io.BytesIO(blob))
    if img.mode in {"RGBA", "P"}:
        img = img.convert("RGB")
    out = io.BytesIO()
    page_w, page_h = A4
    margin = 14 * mm
    avail_w, avail_h = page_w - 2 * margin, page_h - 2 * margin
    ratio = min(avail_w / img.width, avail_h / img.height)
    w, h = img.width * ratio, img.height * ratio
    x, y = (page_w - w) / 2, (page_h - h) / 2
    # Save img to a temp buffer reportlab can ingest.
    img_buf = io.BytesIO()
    img.save(img_buf, format="JPEG", quality=88)
    img_buf.seek(0)
    c = pdfcanvas.Canvas(out, pagesize=A4)
    from reportlab.lib.utils import ImageReader
    c.drawImage(ImageReader(img_buf), x, y, w, h, preserveAspectRatio=True, mask="auto")
    c.showPage(); c.save()
    return out.getvalue()


def _heic_to_pdf(blob: bytes) -> bytes:
    import pillow_heif
    pillow_heif.register_heif_opener()
    return _img_to_pdf(blob)


def _text_to_pdf(blob: bytes, name: str) -> bytes:
    text = blob.decode("utf-8", errors="replace")
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4,
                            leftMargin=14*mm, rightMargin=14*mm,
                            topMargin=14*mm, bottomMargin=14*mm)
    style = ParagraphStyle("Mono", fontName="Courier", fontSize=8.5, leading=11)
    head = ParagraphStyle("Head", fontName="Helvetica-Bold", fontSize=12, leading=14, spaceAfter=8)
    story = [Paragraph(name, head), Spacer(1, 4)]
    # Chunk into 1KB-ish blocks so reportlab can paginate.
    chunk: list[str] = []
    for line in text.splitlines():
        chunk.append(line)
        if len(chunk) >= 60:
            story.append(Preformatted("\n".join(chunk), style))
            chunk = []
    if chunk:
        story.append(Preformatted("\n".join(chunk), style))
    doc.build(story)
    return out.getvalue()


def _docx_text_fallback(blob: bytes, name: str) -> bytes:
    """Lossy but never-blank: pull paragraphs + tables from a .docx and render
    them as reportlab paragraphs. Tables are flattened to tab-separated lines."""
    from docx import Document
    d = Document(io.BytesIO(blob))
    lines: list[str] = []
    for p in d.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for tbl in d.tables:
        lines.append("")  # spacer
        for row in tbl.rows:
            lines.append(" | ".join((c.text or "").strip() for c in row.cells))
    return _text_to_pdf(("\n".join(lines)).encode("utf-8"), name)


def _docx_to_pdf(blob: bytes, name: str) -> tuple[bytes, str]:
    """Best-effort docx → PDF. Tries LibreOffice headless first (high fidelity),
    then docx2pdf (legacy Windows-on-PATH path), and finally the pragmatic
    ReportLab text fallback so we **never** return blank.

    Each fallback reason is logged at INFO so the production log shows the
    pipeline actually used (visible via `tail /var/log/supervisor/backend.out.log`)."""
    # 1. LibreOffice headless — primary path.
    try:
        pdf = _office_to_pdf_via_lo(blob, "docx", name)
        if len(pdf) >= 1024 and _is_pdf(pdf):
            log.info("libreoffice: ok docx=%s bytes=%d", name, len(pdf))
            return pdf, "docx_libreoffice"
        log.info("libreoffice fallback: docx produced %d bytes (<1KB or not PDF)", len(pdf))
    except Exception as e:
        log.info("libreoffice fallback: %s", e)

    # 2. docx2pdf — legacy path (only effective if Word/LO is on PATH, but
    # still gives us a third shot before the lossy fallback).
    with tempfile.TemporaryDirectory() as td:
        td_p = Path(td)
        src = td_p / "in.docx"
        src.write_bytes(blob)
        dst = td_p / "in.pdf"
        try:
            import docx2pdf
            docx2pdf.convert(str(src), str(dst))
            if dst.exists():
                data = dst.read_bytes()
                if len(data) >= 1024 and _is_pdf(data):
                    log.info("docx2pdf: ok docx=%s bytes=%d", name, len(data))
                    return data, "docx_docx2pdf"
        except Exception as e:
            log.info("docx2pdf fallback: %s", e)

    # 3. Pragmatic ReportLab text renderer — never blank.
    log.info("docx text fallback engaged for %s", name)
    return _docx_text_fallback(blob, name), "docx_text_fallback"


def _office_to_pdf_or_415(blob: bytes, ext: str, name: str, pipeline: str) -> tuple[bytes, str]:
    """xlsx / pptx / odt / rtf have no pragmatic fallback. Convert via LO or
    raise a 415 with a useful hint."""
    try:
        pdf = _office_to_pdf_via_lo(blob, ext, name)
        if _is_pdf(pdf) and len(pdf) >= 100:
            log.info("libreoffice: ok %s=%s bytes=%d", ext, name, len(pdf))
            return pdf, pipeline
        raise RuntimeError(f"produced invalid PDF: {len(pdf)} bytes")
    except subprocess.TimeoutExpired:
        log.info("libreoffice timeout for %s — giving up", name)
        raise HTTPException(504, f"LibreOffice timed out converting {name}. Try again or open locally.")
    except Exception as e:
        log.info("libreoffice failed for %s: %s", name, e)
        raise HTTPException(415, f"Couldn't render {ext.upper()} preview: {e}")


# ────────────────── OCR utility (opt-in) ──────────────────

def ocr_pdf_to_text(pdf_path: Path | str, lang: str = "eng", timeout: int = 90) -> str:
    """Extract plaintext from a PDF. Tries `pdftotext` first (fast — works for
    text-layer PDFs), falls back to Tesseract via Poppler's `pdftoppm` when
    the file is image-only. **Not** wired into the upload path — call this
    explicitly from a search-indexer job or admin tool.

    Raises FileNotFoundError if the binaries aren't installed."""
    import shutil
    src = Path(pdf_path)
    if not src.exists():
        raise FileNotFoundError(src)
    if not shutil.which("pdftotext"):
        raise FileNotFoundError("pdftotext (poppler-utils) not installed")
    # Fast path — pdftotext.
    res = subprocess.run(["pdftotext", "-layout", str(src), "-"],
                         capture_output=True, timeout=timeout)
    text = (res.stdout or b"").decode("utf-8", errors="replace").strip()
    if text:
        return text
    # Slow path — rasterise + tesseract OCR.
    if not shutil.which("pdftoppm") or not shutil.which("tesseract"):
        return ""
    with tempfile.TemporaryDirectory() as td:
        td_p = Path(td)
        subprocess.run(["pdftoppm", "-r", "200", str(src), str(td_p / "page"), "-png"],
                       capture_output=True, timeout=timeout)
        out_chunks: list[str] = []
        for img in sorted(td_p.glob("page-*.png")):
            r = subprocess.run(["tesseract", str(img), "-", "-l", lang],
                               capture_output=True, timeout=timeout)
            out_chunks.append((r.stdout or b"").decode("utf-8", errors="replace"))
        return "\n".join(out_chunks).strip()


# ────────────────── cache + dispatcher ──────────────────

async def _cache_lookup(file_id: str, sha1: str, pipeline: str) -> Optional[bytes]:
    row = await db.doc_files_pdf_cache.find_one(
        {"file_id": file_id, "sha1": sha1, "pipeline": pipeline},
        {"_id": 0, "pdf_b64": 1},
    )
    if not row:
        return None
    import base64
    return base64.b64decode(row["pdf_b64"])


async def _cache_store(file_id: str, sha1: str, pipeline: str, pdf: bytes) -> None:
    import base64
    await db.doc_files_pdf_cache.update_one(
        {"file_id": file_id, "sha1": sha1, "pipeline": pipeline},
        {"$set": {
            "file_id": file_id, "sha1": sha1, "pipeline": pipeline,
            "pdf_b64": base64.b64encode(pdf).decode("ascii"),
            "size": len(pdf),
        }},
        upsert=True,
    )


async def _convert(doc: dict, path: Path) -> tuple[bytes, str]:
    sha1 = _sha1_file(path)
    pipeline = _pipeline_for(doc.get("mime"), doc.get("filename") or "")
    if not pipeline:
        ctype = doc.get("mime") or "application/octet-stream"
        msg = (f"LibreOffice not installed — PDF preview not available for "
               f"this format ({ctype})") if "officedocument" in ctype else \
              f"PDF preview not available for {ctype}"
        raise HTTPException(415, msg)
    cached = await _cache_lookup(doc["id"], sha1, pipeline)
    if cached:
        return cached, pipeline
    blob = path.read_bytes()
    if pipeline == "passthrough":
        pdf = blob if _is_pdf(blob) else b""
        if not pdf:
            raise HTTPException(415, "File claims PDF but is not — refusing to serve.")
    elif pipeline == "image": pdf = _img_to_pdf(blob)
    elif pipeline == "heic":  pdf = _heic_to_pdf(blob)
    elif pipeline == "text":  pdf = _text_to_pdf(blob, doc.get("filename") or "Document")
    elif pipeline == "docx_libreoffice":
        pdf, pipeline = _docx_to_pdf(blob, doc.get("filename") or "Document")
    elif pipeline in {"xlsx_libreoffice", "pptx_libreoffice", "odt_libreoffice", "rtf_libreoffice"}:
        ext = pipeline.split("_", 1)[0]
        pdf, pipeline = _office_to_pdf_or_415(blob, ext, doc.get("filename") or f"Document.{ext}", pipeline)
    else:
        raise HTTPException(500, f"Unknown pipeline: {pipeline}")
    await _cache_store(doc["id"], sha1, pipeline, pdf)
    return pdf, pipeline


def _pdf_response(pdf: bytes, original_name: str, dl: bool, pipeline: str) -> Response:
    base = original_name.rsplit(".", 1)[0] or "document"
    fname = f"{base}.pdf"
    disp = "attachment" if dl else "inline"
    # Phase 3.10 hotfix — Chrome blocks cross-origin iframe loading without
    # explicit CSP frame-ancestors + same-site CORP. Stamp them on every PDF
    # response so the PdfPreviewModal iframe loads cleanly.
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'{disp}; filename="{fname}"',
            "X-Pipeline": pipeline,
            "Cache-Control": "private, max-age=3600",
            "X-Frame-Options": "SAMEORIGIN",
            "Content-Security-Policy": (
                "frame-ancestors 'self' https://*.emergentagent.com "
                "https://*.preview.emergentagent.com"
            ),
            "Cross-Origin-Resource-Policy": "same-site",
            "Cross-Origin-Opener-Policy": "same-origin-allow-popups",
        },
    )


# ────────────────── signed preview token (iframe auth) ──────────────────
#
# Iframes can't carry the Authorization header, so the PdfPreviewModal mints a
# short-lived signed token via POST /preview-token and passes it as `?t=` on
# the iframe src. The token is HMAC-SHA256 over {file_id, user_id, exp}, signed
# with the JWT secret. Audience is bound to file_id to prevent token reuse on
# a different file.

def _preview_secret() -> bytes:
    s = os.environ.get("JWT_SECRET") or os.environ.get("SECRET_KEY") or "paneltec-dev"
    return s.encode()


def _mint_preview_token(file_id: str, user_id: str, ttl_seconds: int = 300) -> str:
    payload = {"f": file_id, "u": user_id, "exp": int(time.time()) + ttl_seconds}
    body = base64.urlsafe_b64encode(json.dumps(payload, separators=(",", ":")).encode()).rstrip(b"=").decode()
    sig = hmac.new(_preview_secret(), body.encode(), hashlib.sha256).digest()
    sig_b64 = base64.urlsafe_b64encode(sig).rstrip(b"=").decode()
    return f"{body}.{sig_b64}"


def _verify_preview_token(token: str, file_id: str) -> Optional[str]:
    try:
        body, sig_b64 = token.split(".", 1)
        expected = hmac.new(_preview_secret(), body.encode(), hashlib.sha256).digest()
        got = base64.urlsafe_b64decode(sig_b64 + "==")
        if not hmac.compare_digest(expected, got):
            return None
        payload = json.loads(base64.urlsafe_b64decode(body + "==").decode())
        if payload.get("f") != file_id:
            return None
        if int(payload.get("exp", 0)) < int(time.time()):
            return None
        return payload.get("u")
    except Exception:
        return None


# ────────────────── endpoints ──────────────────

@router.get("/files/{file_id}/pdf")
async def file_pdf(file_id: str, request: Request,
                   background: BackgroundTasks,
                   dl: int = Query(0),
                   t: Optional[str] = Query(None, description="signed iframe token; alternative to Bearer auth")):
    """Stream PDF. Accepts EITHER an Authorization Bearer header (curl /
    download path) OR a short-lived signed `?t=` token (iframe path, since
    iframes can't carry custom headers).

    Phase 3.14 — first time a file is converted to PDF, we kick a background
    task that runs `ocr_pdf_to_text` against the PDF bytes and persists the
    result onto `doc_files.search_text` so the Smart Search indexer (existing
    or future) can pick it up. Fire-and-forget; the response returns
    immediately as it does today."""
    if t:
        user_id = _verify_preview_token(t, file_id)
        if not user_id:
            raise HTTPException(401, "Invalid or expired preview token")
        u = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not u:
            raise HTTPException(401, "Token user not found")
        user = u
    else:
        user = await get_current_user(request, creds=None)
    doc, path = await _resolve_file(file_id, user)
    pdf, pipeline = await _convert(doc, path)
    # Spool the OCR + index step into the background. We persist the PDF to a
    # short-lived temp file so `ocr_pdf_to_text` (which expects a Path) can
    # read it without re-converting. The doc id keeps us idempotent — see the
    # `already_indexed` short-circuit in `_ocr_index_file`.
    try:
        tmp = Path(tempfile.gettempdir()) / f"ocr_idx_{doc['id']}.pdf"
        tmp.write_bytes(pdf)
        background.add_task(_ocr_index_file, doc["id"], tmp)
    except Exception as e:
        log.warning("ocr scheduling failed file=%s err=%s", doc.get("id"), e)
    return _pdf_response(pdf, doc.get("filename") or file_id, bool(dl), pipeline)


@router.get("/files/{file_id}/pdf.pdf")
async def file_pdf_aliased(file_id: str, request: Request,
                           background: BackgroundTasks,
                           dl: int = Query(0), t: Optional[str] = Query(None)):
    return await file_pdf(file_id, request, background, dl=dl, t=t)


@router.post("/files/{file_id}/preview-token")
async def mint_file_preview_token(file_id: str, user: dict = Depends(get_current_user)):
    """Issue a 5-minute signed token bound to (file_id, user_id) so the iframe
    can fetch the PDF without an Authorization header."""
    doc, _ = await _resolve_file(file_id, user)  # access check + 404
    token = _mint_preview_token(doc["id"], user["id"])
    return {"token": token, "expires_in": 300}


# ────────────────── inline PDF stash (Phase 3.13.1) ──────────────────
#
# Some flows generate a PDF on the fly and pass it to the universal
# PdfPreviewModal (Site QR sheet, Supplier QR sheet, Induction Print, ...).
# Originally those flows wrapped the bytes in `URL.createObjectURL` and gave
# the iframe a `blob:` URL — which ad blockers and privacy extensions
# routinely refuse to load (ERR_BLOCKED_BY_CLIENT).
#
# To make those previews behave like normal HTTPS document loads, we stash
# the PDF bytes in-process and return a same-origin signed URL the iframe
# can hit directly. The stash is org-scoped, capped, and TTL'd.
_INLINE_STASH: dict[str, dict] = {}
_INLINE_STASH_TTL_SECONDS = 600   # 10 minutes
_INLINE_STASH_MAX_BYTES = 25 * 1024 * 1024   # 25 MB cap per blob
_INLINE_STASH_MAX_ENTRIES = 200


def _stash_prune() -> None:
    now = time.time()
    expired = [k for k, v in _INLINE_STASH.items() if v["exp"] < now]
    for k in expired:
        _INLINE_STASH.pop(k, None)
    # Hard cap — drop oldest if we've blown past the limit.
    while len(_INLINE_STASH) > _INLINE_STASH_MAX_ENTRIES:
        oldest = min(_INLINE_STASH, key=lambda k: _INLINE_STASH[k]["exp"])
        _INLINE_STASH.pop(oldest, None)


def stash_inline_pdf(pdf_bytes: bytes, user_id: str, org_id: str,
                      filename: str = "document.pdf",
                      ttl_seconds: int = _INLINE_STASH_TTL_SECONDS) -> str:
    """Persist `pdf_bytes` in-memory and return a stash id callers can hand to
    the frontend so it can pull the bytes via `GET /files/inline/{id}?t=...`
    instead of feeding a `blob:` URL to the iframe."""
    if not pdf_bytes:
        raise HTTPException(400, "Empty PDF body")
    if len(pdf_bytes) > _INLINE_STASH_MAX_BYTES:
        raise HTTPException(413, f"PDF too large to stash ({len(pdf_bytes)} bytes)")
    _stash_prune()
    sid = new_id()
    _INLINE_STASH[sid] = {
        "bytes": pdf_bytes,
        "user_id": user_id,
        "org_id": org_id,
        "filename": filename,
        "exp": time.time() + ttl_seconds,
    }
    return sid


@router.get("/files/inline/{stash_id}")
async def serve_inline_pdf(stash_id: str, request: Request,
                            t: Optional[str] = Query(None,
                                description="signed iframe token; alternative to Bearer auth")):
    """Stream a previously stashed PDF. Accepts either a Bearer header (curl /
    same-tab navigation) or a `?t=` signed token (iframe path).

    The stash entry pins (user_id, org_id) so a token bound to a different
    user can't pull it. TTL ~10 min; entry is left in place until it expires
    so the iframe can re-fetch if the page reloads."""
    _stash_prune()
    entry = _INLINE_STASH.get(stash_id)
    if not entry:
        raise HTTPException(404, "Inline preview not found or expired")
    if t:
        user_id = _verify_preview_token(t, stash_id)
        if not user_id:
            raise HTTPException(401, "Invalid or expired preview token")
    else:
        user = await get_current_user(request, creds=None)
        user_id = user["id"]
    if user_id != entry["user_id"]:
        raise HTTPException(403, "Preview belongs to a different user")
    return _pdf_response(entry["bytes"], entry["filename"], dl=False,
                         pipeline="inline_stash")


@router.post("/files/inline-pdf")
async def stash_inline_endpoint(request: Request, user: dict = Depends(get_current_user)):
    """Generic stash endpoint: POST a PDF body (binary or multipart) and get
    back `{stash_id, token, expires_in}`. The iframe then loads
    `/api/files/inline/{stash_id}?t={token}` to render the PDF as a normal
    HTTPS document — sidestepping the `blob:` URL ad-blocker block.

    Designed for callers that already have the PDF bytes locally (e.g.
    they did a normal `axios.post(..., {responseType:'blob'})` and want to
    show the result in PdfPreviewModal without a `blob:` URL)."""
    body = await request.body()
    if not body:
        raise HTTPException(400, "Empty body")
    if not _is_pdf(body):
        raise HTTPException(400, "Body is not a PDF")
    # Filename is optional, comes from X-Filename header if provided.
    filename = request.headers.get("x-filename") or "document.pdf"
    sid = stash_inline_pdf(body, user["id"], user.get("org_id") or "",
                           filename=filename)
    token = _mint_preview_token(sid, user["id"])
    return {"stash_id": sid, "token": token, "expires_in": 300}


class BundleIn(BaseModel):
    file_ids: list[str]


@router.post("/files/pdf-bundle")
async def file_pdf_bundle(body: BundleIn, user: dict = Depends(get_current_user)):
    if user.get("role") not in {"admin", "manager", "hseq_lead"}:
        raise HTTPException(403, "Bulk PDF bundles require admin/manager role")
    if not body.file_ids:
        raise HTTPException(400, "file_ids cannot be empty")
    if len(body.file_ids) > 25:
        raise HTTPException(400, "Max 25 files per bundle")
    from PyPDF2 import PdfMerger
    merger = PdfMerger()
    converted = 0; skipped: list[dict] = []
    for fid in body.file_ids:
        try:
            doc, path = await _resolve_file(fid, user)
            pdf, _ = await _convert(doc, path)
            merger.append(io.BytesIO(pdf))
            converted += 1
        except HTTPException as e:
            skipped.append({"file_id": fid, "reason": e.detail})
    if converted == 0:
        raise HTTPException(415, f"No files could be converted: {skipped}")
    buf = io.BytesIO()
    merger.write(buf); merger.close()
    headers = {
        "Content-Disposition": f'attachment; filename="paneltec-bundle-{converted}.pdf"',
        "X-Bundle-Converted": str(converted),
        "X-Bundle-Skipped": str(len(skipped)),
    }
    return Response(content=buf.getvalue(), media_type="application/pdf", headers=headers)


# ────────────────── admin install hook (v146 — background job) ──────────────────
#
# Prior to v146 this endpoint blocked the HTTP request for the full duration
# of `apt-get install`, which on a first-time install fetches ~650 MB of
# LibreOffice deps and takes 5–10 minutes. The public URL runs behind
# Cloudflare + Kubernetes ingress, both of which enforce a ~100 s idle-read
# timeout; the browser saw a 5xx while the backend kept installing to
# completion server-side. Users experienced "install goes part-way then
# stops" but the packages actually finished landing.
#
# v146 fix: POST returns 202 immediately with a job_id, and the subprocess
# runs as a detached asyncio task. `_INSTALL_STATE` tracks live progress
# (rolling last-50-lines log tail) so `GET /admin/server-tools/health`
# can surface it. The frontend polls health every 5 s until
# `install_running=false`.
_INSTALL_STATE: Dict[str, Any] = {
    "install_running": False,
    "job_id": None,
    "started_at": None,
    "finished_at": None,
    "exit_code": None,
    "packages": None,
    "log_tail": deque(maxlen=50),
}
_INSTALL_WALL_CLOCK_S = 20 * 60   # 20-minute hard ceiling


def _install_log_tail_str() -> str:
    return "\n".join(_INSTALL_STATE["log_tail"])


async def _run_apt_install(job_id: str, pkgs: list) -> None:
    """Background task — runs `apt-get install` and streams stdout/stderr
    line-by-line into the module-level `_INSTALL_STATE["log_tail"]` deque.
    Enforces the 20-min wall-clock cap by SIGKILLing the subprocess."""
    _INSTALL_STATE["log_tail"].append(
        f"[paneltec] Job {job_id} — installing: {', '.join(pkgs)}"
    )
    cmd = [
        "bash", "-lc",
        # v151.1 — `dpkg --configure -a` first cleans up any interrupted
        # prior install (e.g. an earlier apt run that was killed by a
        # container refresh mid-transaction). Idempotent no-op on a clean
        # system. Without it, apt-get refuses to proceed with "dpkg was
        # interrupted, you must manually run 'dpkg --configure -a'".
        f"dpkg --configure -a 2>&1 || true; "
        f"apt-get update -qq && apt-get install -y --no-install-recommends "
        f"{' '.join(pkgs)} 2>&1; echo '---'; "
        f"which libreoffice || which soffice || true; "
        f"echo '---'; which tesseract || true; "
        f"echo '---'; which pdftotext || true",
    ]
    proc = None
    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )

        async def _reader():
            assert proc is not None and proc.stdout is not None
            while True:
                line_b = await proc.stdout.readline()
                if not line_b:
                    return
                line = line_b.decode("utf-8", errors="replace").rstrip()
                if line:
                    _INSTALL_STATE["log_tail"].append(line)

        try:
            await asyncio.wait_for(_reader(), timeout=_INSTALL_WALL_CLOCK_S)
            rc = await proc.wait()
        except asyncio.TimeoutError:
            _INSTALL_STATE["log_tail"].append(
                f"[paneltec] wall-clock timeout after {_INSTALL_WALL_CLOCK_S}s — killing subprocess"
            )
            try:
                proc.kill()
                await proc.wait()
            except Exception as ke:
                _INSTALL_STATE["log_tail"].append(f"[paneltec] kill failed: {ke}")
            rc = -1

        _INSTALL_STATE["exit_code"] = rc
        _INSTALL_STATE["log_tail"].append(f"[paneltec] apt-get exited rc={rc}")
    except Exception as e:
        log.exception("install_libreoffice background task crashed")
        _INSTALL_STATE["exit_code"] = -1
        _INSTALL_STATE["log_tail"].append(f"[paneltec] background task crashed: {e}")
    finally:
        _INSTALL_STATE["install_running"] = False
        _INSTALL_STATE["finished_at"] = datetime.now(timezone.utc).isoformat()


# ────────────────── v151.1 — auto-install on backend boot ──────────────────
#
# WHY: Our pod runs the Emergent-managed base image
# `mono_fullstack_base_image_cloud_arm:release-15062026-2` which does NOT
# include LibreOffice / Tesseract / Poppler. The image's writable overlay
# is tied to the container lifecycle — supervisor service restarts survive,
# but any container refresh (memory pressure, forced Emergent update, node
# reschedule) resets `/usr`, wiping the apt-installed packages we needed
# for DOCX→PDF conversion and OCR. Users saw the Server Tools admin pill
# regress to red after each container refresh (v146 first spotted it).
#
# Emergent doesn't expose a Dockerfile / apt-packages hook to us from
# inside the pod (only `/app/.emergent/emergent.yml` which is an opaque
# image reference). So the durable fix here is to detect the missing
# tools every time the backend boots and dispatch the same apt-get task
# that `POST /admin/install-libreoffice` uses. Fire-and-forget:
# `on_startup` returns immediately and apt runs in the background. The
# health endpoint's `install_running=true` + `log_tail` show progress
# to any admin watching, without them touching a button.
def ensure_server_tools_or_install_bg() -> Dict[str, Any]:
    """Detect libreoffice/tesseract/poppler and kick off an async apt-get
    if any are missing. Idempotent — a no-op if all tools are present or
    if an install is already in flight. Fires-and-forgets; caller must
    never await. Returns a small status dict for logging."""
    import shutil
    tools = {
        "libreoffice": shutil.which("libreoffice") or shutil.which("soffice"),
        "tesseract":   shutil.which("tesseract"),
        "poppler":     shutil.which("pdftotext"),
    }
    missing = [name for name, path in tools.items() if not path]

    if not missing:
        return {"missing": [], "action": "noop", "reason": "all tools present"}

    if _INSTALL_STATE.get("install_running"):
        return {
            "missing": missing,
            "action": "skip",
            "reason": "install already running",
            "job_id": _INSTALL_STATE.get("job_id"),
        }

    # Always install the full toolchain so a partial wipe doesn't leave us
    # with mismatched versions. Same package list as the manual endpoint.
    pkgs = [
        "libreoffice-core", "libreoffice-writer",
        "libreoffice-calc", "libreoffice-impress",
        "tesseract-ocr", "poppler-utils",
    ]
    job_id = str(uuid.uuid4())
    started = datetime.now(timezone.utc).isoformat()
    _INSTALL_STATE.update({
        "install_running": True,
        "job_id": job_id,
        "started_at": started,
        "finished_at": None,
        "exit_code": None,
        "packages": list(pkgs),
    })
    _INSTALL_STATE["log_tail"].clear()
    _INSTALL_STATE["log_tail"].append(
        f"[paneltec] auto-install on boot — missing: {', '.join(missing)}"
    )
    # Fire-and-forget. `_run_apt_install` handles its own exceptions and
    # always flips `install_running` back to False in its finally block.
    asyncio.create_task(_run_apt_install(job_id, pkgs))
    return {"missing": missing, "action": "queued", "job_id": job_id}



@router.post("/admin/install-libreoffice", status_code=202)
async def install_libreoffice(
    include_ocr: bool = Query(True, description="Also install Tesseract + Poppler for OCR"),
    user: dict = Depends(get_current_user),
):
    """Admin-only one-click toolchain installer. Kicks off apt-get in a
    background asyncio task and returns 202 immediately with a job_id so
    the caller can poll `/admin/server-tools/health` for progress —
    apt-get runs for 5–10 min on a cold cache, longer than any edge
    HTTP proxy will hold a connection open."""
    if user.get("role") != "admin":
        raise HTTPException(403, "Only admin can install system packages")
    if _INSTALL_STATE["install_running"]:
        raise HTTPException(
            409,
            {
                "detail": "Install already running",
                "job_id": _INSTALL_STATE["job_id"],
                "started_at": _INSTALL_STATE["started_at"],
            },
        )
    pkgs = [
        "libreoffice-core", "libreoffice-writer",
        "libreoffice-calc", "libreoffice-impress",
    ]
    if include_ocr:
        pkgs += ["tesseract-ocr", "poppler-utils"]

    job_id = str(uuid.uuid4())
    started = datetime.now(timezone.utc).isoformat()
    _INSTALL_STATE.update({
        "install_running": True,
        "job_id": job_id,
        "started_at": started,
        "finished_at": None,
        "exit_code": None,
        "packages": list(pkgs),
    })
    _INSTALL_STATE["log_tail"].clear()

    asyncio.create_task(_run_apt_install(job_id, pkgs))

    return {
        "job_id": job_id,
        "started_at": started,
        "install_running": True,
        "packages": pkgs,
    }


async def _tool_status() -> dict:
    """which-style status of optional server toolchains."""
    async def _which(name: str) -> str | None:
        p = await asyncio.create_subprocess_exec(
            "bash", "-lc", f"which {name} 2>/dev/null || true",
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.DEVNULL,
        )
        out, _ = await p.communicate()
        path = out.decode().strip().splitlines()[0] if out else ""
        return path or None

    async def _version(bin_path: str, flag: str = "--version") -> str | None:
        if not bin_path:
            return None
        p = await asyncio.create_subprocess_exec(
            "bash", "-lc", f"{bin_path} {flag} 2>&1 | head -1",
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.DEVNULL,
        )
        out, _ = await p.communicate()
        return (out.decode().strip() or None)

    lo = await _which("libreoffice") or await _which("soffice")
    ts = await _which("tesseract")
    pp = await _which("pdftotext")
    return {
        "libreoffice": {"installed": bool(lo), "path": lo, "version": await _version(lo) if lo else None},
        "tesseract":   {"installed": bool(ts), "path": ts, "version": await _version(ts) if ts else None},
        "poppler":     {"installed": bool(pp), "path": pp, "version": await _version(pp, "-v") if pp else None},
    }


@router.get("/admin/system-tools")
async def system_tools(user: dict = Depends(get_current_user)):
    """Status of optional server toolchains (admin-only). Drives the
    Settings → System page in the UI."""
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    return {"tools": await _tool_status()}


@router.get("/admin/server-tools/health")
async def server_tools_health(user: dict = Depends(get_current_user)):
    """Phase 3.13 — health-check shape requested by the Settings page.
    Returns `{libreoffice:{ok,version,path}, tesseract:{...}, poppler:{...}}`.
    Same data as `/admin/system-tools` but normalised to the `ok` key the
    UI uses to colour the chip green/red without a key-mapping helper.

    v146 — also surfaces the background install job progress so the UI
    can poll a single endpoint and show a live log tail while apt-get
    is still running (`install_running`, `install_job_id`,
    `install_log_tail`, `install_exit_code`, `install_started_at`,
    `install_finished_at`)."""
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    s = await _tool_status()
    def _norm(t: dict) -> dict:
        return {"ok": bool(t.get("installed")), "version": t.get("version"), "path": t.get("path")}
    return {
        "libreoffice": _norm(s["libreoffice"]),
        "tesseract":   _norm(s["tesseract"]),
        "poppler":     _norm(s["poppler"]),
        "install_running":     bool(_INSTALL_STATE["install_running"]),
        "install_job_id":      _INSTALL_STATE["job_id"],
        "install_started_at":  _INSTALL_STATE["started_at"],
        "install_finished_at": _INSTALL_STATE["finished_at"],
        "install_exit_code":   _INSTALL_STATE["exit_code"],
        "install_log_tail":    _install_log_tail_str() or None,
    }


# ────────────── Phase 3.14 — Auto-OCR-to-SmartSearch on upload ──────────────
# Fire-and-forget extractor: after a file is converted to PDF (cached on disk
# at <UPLOAD_DIR>/cache/<file_id>.pdf), we spawn a background task that runs
# the existing `ocr_pdf_to_text()` util and persists the result onto
# `doc_files.search_text`. Triggered from the /pdf endpoint so it benefits
# from the LibreOffice cache without re-converting.

# 50 MB cap — anything larger blows past the tesseract timeout and the
# search-relevance per byte falls off a cliff.
OCR_INDEX_MAX_BYTES = 50 * 1024 * 1024


async def _ocr_index_file(file_id: str, pdf_path: Path) -> None:
    """Background task: extract text and persist to doc_files.search_text.
    Cheap when the PDF has a text layer (pdftotext fast path); slow only
    when tesseract has to OCR rasterised pages."""
    try:
        existing = await db.doc_files.find_one({"id": file_id}, {"_id": 0, "search_text": 1, "size": 1})
        if not existing:
            return
        if existing.get("search_text"):
            log.info("ocr skipped file=%s reason=already_indexed", file_id)
            return
        if int(existing.get("size") or 0) > OCR_INDEX_MAX_BYTES:
            log.info("ocr skipped file=%s reason=size", file_id)
            await db.doc_files.update_one({"id": file_id},
                {"$set": {"search_text_status": "skipped_size", "search_text_at": now_iso()}})
            return
        text = ocr_pdf_to_text(pdf_path)
        await db.doc_files.update_one({"id": file_id},
            {"$set": {"search_text": text or "", "search_text_chars": len(text or ""),
                      "search_text_status": "indexed", "search_text_at": now_iso()}})
        log.info("ocr indexed file=%s chars=%d", file_id, len(text or ""))
    except Exception as e:
        log.warning("ocr failed file=%s err=%s", file_id, e)
        try:
            await db.doc_files.update_one({"id": file_id},
                {"$set": {"search_text_status": f"error: {str(e)[:120]}",
                          "search_text_at": now_iso()}})
        except Exception:
            pass


@router.get("/admin/files/{file_id}/search-text")
async def admin_file_search_text(file_id: str, user: dict = Depends(get_current_user)):
    """Debug-only — returns the OCR'd text persisted on the file doc."""
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    doc = await db.doc_files.find_one(
        {"id": file_id, "org_id": user["org_id"]},
        {"_id": 0, "search_text": 1, "search_text_chars": 1,
         "search_text_status": 1, "search_text_at": 1, "filename": 1, "size": 1},
    )
    if not doc:
        raise HTTPException(404, "File not found")
    return doc
