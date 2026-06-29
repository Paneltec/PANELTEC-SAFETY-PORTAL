"""Phase 4.x — SWMS extras.

Extends the legacy AI-draft SWMS module with:
  • A startup seed of the SWMS-06 Concrete/Asphalt Cutting V12.0 record so
    customers always see at least one structured SWMS in their library.
  • POST /api/swms/import-docx — admin endpoint that accepts a .docx URL or
    raw upload, parses it via python-docx, and returns the inferred
    structured payload for review (does NOT auto-save).
"""
from __future__ import annotations

import io
import logging
import re
import uuid
from datetime import datetime, timezone

import httpx
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel

from db import db
from auth import get_current_user


log = logging.getLogger("paneltec.swms.extras")

router = APIRouter(prefix="/swms", tags=["swms"])


# ────────────────── SWMS-06 seed ──────────────────

SWMS_06_PAYLOAD = {
    "code": "SWMS-06",
    "version": "V12.0",
    "slug": "swms-06-concrete-or-asphalt-cutting",
    "title": "Concrete or Asphalt Cutting",
    "scope": "High-risk construction work involving concrete or asphalt cutting.",
    "high_risk_construction_work": "CONCRETE OR ASPHALT CUTTING",
    "job_description": "High-risk construction work involving concrete or asphalt cutting.",
    "prepared_by": {
        "name": "Patrick Monaghan", "role": "Operations Manager",
        "organisation": "Paneltec Group", "date_prepared": "2025-08-31",
    },
    "approved_by": {
        "name": "John Guy", "position": "Director - Paneltec Pty Ltd",
        "contact": "0408 449 023", "date_approved": "2025-08-31",
    },
    "review_date": "2026-08-31",
    "activity_analysis": [
        {"step": "Training of operator", "potential_hazards": ["An untrained operator is a danger to himself and others"], "risk_class_before": 1, "controls": ["Make sure the operator is familiar with the type of equipment they are going to use and are trained in the supplier's safe use manual and is able to operate the equipment to the supplier's recommendations.", "Make sure the operator understands what the hazards associated with the cutting operation are, why the risks need to be controlled and how the risks are to be controlled.", "Make sure the operator understands what protective equipment is needed, why it is needed, and how it is to be fitted and used."], "responsible": ["Management", "Supervisor"], "risk_class_after": 2},
        {"step": "Protection of other workers and the public", "potential_hazards": ["Flying debris"], "risk_class_before": 1, "controls": ["Barricades and signs may be needed to warn people who may be at risk from the cutting operation."], "responsible": ["Supervisor"], "risk_class_after": 2},
        {"step": "Type and condition of the concrete", "potential_hazards": ["Obstructions or resistance in the material being cut – can cause sudden kick-back, pushback or pull-in movements of the saw"], "risk_class_before": 1, "controls": ["Check whether the material contains reinforcing steel, electrical cable conduits, gas pipes or other services.", "Service plans should be checked to ensure that live services will not be cut by the saw."], "responsible": ["Supervisor"], "risk_class_after": 2},
        {"step": "Confined spaces", "potential_hazards": ["Carbon monoxide"], "risk_class_before": 1, "controls": ["Petrol driven equipment should not be used in confined spaces because of the potential for carbon monoxide poisoning."], "responsible": ["Supervisor"], "risk_class_after": 2},
        {"step": "Electrical equipment in wet conditions", "potential_hazards": ["Electrocution"], "risk_class_before": 1, "controls": ["If electrical equipment is used in wet conditions, the potential for electrocution needs to be considered and be guarded against."], "responsible": ["Supervisor"], "risk_class_after": 2},
        {"step": "Personal Protective Equipment", "potential_hazards": ["Cuts", "Noise (hearing loss)", "Dust", "Flying debris"], "risk_class_before": 3, "controls": ["Safety helmet; Safety footwear; Safety goggles; Face shield; Hearing protection; Sun and weather protection; Gloves to improve grip and reduce force and vibration; Respiratory protection where hazardous dusts or fumes cannot be eliminated."], "responsible": ["Operator"], "risk_class_after": 3},
        {"step": "Using the equipment — Setting Up", "potential_hazards": ["Manual handling sprains and strains", "Noise (hearing loss)", "Vibration (circulatory damage)", "Electricity (electric shock)", "Surfaces slippery from cutting residue (falls, sprains)", "Dust"], "risk_class_before": 1, "controls": ["Exact location of the cut or penetration is clearly marked on the work area.", "A trolley supports the cutting machine for horizontal work at low level so operators do not bend forward or work on their knees.", "Cutting blade is the right size and type for the machine.", "Blade is in good working condition (no cracks, gaps, warping, deterioration).", "Electrically powered machines protected at power outlet with RCD.", "Appropriate barricading and warning signs erected.", "Adequate lighting provided.", "Method of collecting residue in place to prevent slippery surfaces."], "responsible": ["Operator"], "risk_class_after": 2},
        {"step": "Before cutting operation", "potential_hazards": ["Cuts", "Manual handling sprains and strains", "Noise (hearing loss)", "Vibration (circulatory damage)", "Electricity (electric shock)", "Surfaces slippery from cutting residue (falls, sprains)", "Dust"], "risk_class_before": 1, "controls": ["General condition of equipment checked by operator before each job (cutting tool, guards, leads, hydraulic hoses). Tag out defective equipment.", "Cutting speed matches drive speed per manufacturer's specification.", "Shaft and flanges clean and undamaged.", "Blade fits securely over the shaft.", "Shaft nut securely tightened against outside flange.", "Blade guard fitted and in good working order.", "Drive belt tensioned correctly.", "Adequate coolant or water readily available for wet cutting."], "responsible": ["Operator"], "risk_class_after": 2},
        {"step": "During the cutting operation", "potential_hazards": ["Cuts", "Manual handling sprains and strains", "Noise (hearing loss)", "Vibration (circulatory damage)", "Electricity (electric shock)", "Surfaces slippery from cutting residue (falls, sprains)", "Dust"], "risk_class_before": 1, "controls": ["Blade guard in lowered position.", "When starting, operator and others stand outside path of blade.", "If machine stalls, raise blade and check outside flange/nut for tightness.", "When resuming, blade aligned with previous cut.", "Wall cuts performed with operator's back close to vertical, hands not above shoulder height.", "Cut from a standing posture with feet braced and body balanced; kneel on one knee with knee protection when needed.", "Cutting horizontally across a wall, hands at waist height.", "Minimise time in fixed postures.", "Plenty of water/coolant to suppress dust at point of generation (critical for silica risk).", "Throttle lock only used when starting the equipment.", "Stop equipment when changing grip between horizontal and vertical cuts.", "Use the provided handles to support equipment (do not support by belt guard).", "Electrical leads not cut during operation; mind water/electrocution risks.", "Assistants located clear of saw movement, ejected material, dropped machine, falling offcuts.", "Saw used only with blade rotating opposite to cut direction; not used for inverted cutting.", "Hand saw never used above shoulder height.", "Hand saw work practices minimise kick backs."], "responsible": ["Operator"], "risk_class_after": 2},
        {"step": "Working alone", "potential_hazards": ["Inability to respond to incident/emergency without assistance"], "risk_class_before": 1, "controls": ["Do not work alone as this can be hazardous because of the potential need to provide assistance in the event of an unsafe incident or emergency."], "responsible": ["Supervisor", "Operator"], "risk_class_after": 2},
        {"step": "Maintenance", "potential_hazards": ["Equipment failure / various"], "risk_class_before": 1, "controls": ["Equipment inspected and maintained regularly by a competent person."], "responsible": ["Supervisor"], "risk_class_after": 2},
    ],
    "environmental_risks": [
        {"work_activity": "Emergencies", "risk": "Accident/Incident, Fire, Spill", "risk_class_before": 1, "controls": ["Personnel advised at toolbox meeting and prepared.", "Emergency numbers available on site.", "Fire extinguishers / First aid kits / Spill kits on site."], "responsible": ["Supervisor", "All"], "risk_class_after": 2},
        {"work_activity": "Air Quality", "risk": "Exhaust Fumes, Dust", "risk_class_before": 2, "controls": ["Control exhaust emissions (regular maintenance, replace old machinery, fit emission control).", "Dampen milled surface with water cart.", "Use water cart where required."], "responsible": ["Supervisor", "All"], "risk_class_after": 3},
        {"work_activity": "Community Relations", "risk": "Inconvenience from works/traffic management", "risk_class_before": 2, "controls": ["Identify potentially impacted parties.", "Letter drops to notify residents/businesses."], "responsible": ["Supervisor", "All"], "risk_class_after": 3},
        {"work_activity": "Flora & Fauna", "risk": "Vegetation removal; native fauna risk from excavation", "risk_class_before": 2, "controls": ["Assess work area for significant vegetation / faunal habitat — define work and exclusion areas using fencing and signage.", "Machinery thoroughly washed down prior to leaving site where noxious weeds present."], "responsible": ["Supervisor", "All"], "risk_class_after": 3},
        {"work_activity": "Heritage & Archaeology", "risk": "Loss/destruction of artifacts; disturbance of historical sites", "risk_class_before": 1, "controls": ["Undertake site survey for significant areas (client rep / government bodies).", "Develop project-specific procedure.", "Install exclusion fencing and signage if required."], "responsible": ["Supervisor", "All"], "risk_class_after": 2},
        {"work_activity": "Noise & Vibration", "risk": "Noise sensitivity to public and fauna", "risk_class_before": 1, "controls": ["Work during normal hours where possible.", "Affected residents notified.", "Plant regularly serviced with efficient mufflers.", "Use smaller plant to reduce vibration."], "responsible": ["Supervisor", "All"], "risk_class_after": 2},
        {"work_activity": "Soil Management & Water Quality", "risk": "Off-site soil impacts; stockpile damage; mud on roadways", "risk_class_before": 2, "controls": ["Wash down machinery only in nominated areas (>20m from waterways).", "Use only approved stockpile areas.", "Build bund walls from available site material.", "Drivers check vehicle tyres/bodies for mud build-up."], "responsible": ["Supervisor", "All"], "risk_class_after": 3},
        {"work_activity": "Fuels and Chemicals", "risk": "Waterway contamination from spillages", "risk_class_before": 1, "controls": ["Store per standards; minimise on-site storage away from drains.", "All containers labelled; SDS available on site.", "Drums stored securely in bunded area.", "Spill kits available.", "Regular maintenance of plant hoses.", "Correct refuelling equipment."], "responsible": ["Supervisor", "All"], "risk_class_after": 2},
        {"work_activity": "Visual Impact / Waste Management", "risk": "Litter / dust / mud / lighting impact on residents and waterways", "risk_class_before": 2, "controls": ["Use site induction to communicate tidy-site requirement.", "Inspect site frequently; insist on tidiness.", "Keep public roads free of dirt/mud.", "Repair broken-out paved surfaces ASAP.", "Remove all rubbish from site daily.", "Waste recycled where possible."], "responsible": ["Supervisor", "All"], "risk_class_after": 3},
    ],
    "ppe": [
        "Safety Hi-vis Vests/Shirts", "Safety Boots", "Sunscreen",
        "Broad Brimmed Sun Hat", "Sunglasses",
        "Gloves specific to the task", "White reflective overalls/night",
        "Hard Hat where required", "Safety Glasses",
        "Hearing protection where required", "Long clothing", "Dust masks",
        "Respiratory equipment if required",
    ],
    "training_requirements": [
        "Task Specific Training depending on activity",
        "Current Construction Induction Card",
        "Client Site Induction (where required)",
        "Activity Induction (SWMS & SWPs)",
        "Vehicle Licences (car or truck depending on class of vehicle or mobile plant driven)",
        "Plant Operator licences or competency assessments where required",
        "Competencies for Work Zone Traffic Management",
        "First Aid certificate training", "Manual Handling training",
        "Fire Extinguisher training", "Fatigue Management Training",
    ],
    "equipment_list": [
        "Cutting machine (concrete saw)",
        "Trolley to support cutting machine",
        "Cutting blade", "Blade guard",
        "Electrically powered machines",
        "Residual current device (earth leakage circuit breaker)",
        "Water (for wet cutting / dust suppression)",
        "Spill kits", "Fire extinguishers", "First aid kits", "Water cart",
    ],
    "emergency_procedures": {
        "general": "Personnel advised at toolbox meeting and prepared in the event of an incident occurring. Emergency numbers available on site. Fire extinguishers / First aid kits / Spill kits on site.",
        "accident_incident": "Provide first aid; call 000 if required; notify Supervisor immediately; secure the area; preserve evidence; complete Incident Report in Paneltec app.",
        "fire": "Raise alarm; evacuate to muster point; use appropriate fire extinguisher if safe; call 000; notify Supervisor.",
        "spill": "Stop source if safe; contain using spill kit; protect drains/waterways; notify Supervisor; arrange disposal per SDS.",
    },
    "status": "approved",
    "attendance_sheet_template": True,
    "source_file": {
        "url": "https://www.dropbox.com/scl/fi/ba4mpdsukoyirlur3ekhz/2025_SWMS-06_Concrete_or_Asphalt_Cutting-V12.0.docx?rlkey=rb3y1q774rkvrp9skox64qtws&dl=1",
        "filename": "2025_SWMS-06_Concrete_or_Asphalt_Cutting-V12.0.docx",
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
    "applies_to": {
        "asset_kinds": ["plant"],
        "asset_types": ["concrete_saw", "slab_cutter"],
        "worker_ids": [],
        "roles": [{"role": "operator"}, {"role": "foreman"}],
        "companies": [],
    },
}


async def seed_swms_06() -> dict:
    """Insert SWMS-06 into every org workspace that doesn't yet have it.
    Idempotent: matches on (org_id, code, version)."""
    inserted = 0
    skipped = 0
    async for ws in db.workspaces.find({"deleted_at": None}, {"_id": 0, "id": 1, "org_id": 1, "name": 1}):
        existing = await db.swms.find_one({
            "org_id": ws["org_id"], "code": "SWMS-06",
            "version": "V12.0", "deleted_at": None,
        }, {"_id": 0, "id": 1})
        if existing:
            skipped += 1
            continue
        now = datetime.now(timezone.utc).isoformat()
        doc = {
            **SWMS_06_PAYLOAD,
            "id": str(uuid.uuid4()),
            "org_id": ws["org_id"],
            "workspace_id": ws["id"],
            "created_at": now,
            "updated_at": now,
            "deleted_at": None,
        }
        await db.swms.insert_one(doc)
        inserted += 1
        # Only seed once per org (first workspace) to avoid duplicate copies.
        await db.swms.create_index([("org_id", 1), ("code", 1), ("version", 1)])
    return {"inserted": inserted, "skipped": skipped}


# ────────────────── Import .docx ──────────────────

class ImportDocxIn(BaseModel):
    url: str
    workspace_id: str | None = None


@router.post("/import-docx")
async def import_swms_docx(body: ImportDocxIn, user: dict = Depends(get_current_user)):
    """Fetch a .docx, parse via python-docx, return a best-effort structured
    payload for the admin to review. **Does NOT auto-save** — the front-end
    must POST the returned payload to /api/swms after the admin clicks
    confirm so missing/incorrect fields can be edited first.
    """
    if user.get("role") not in {"admin", "manager", "hseq_lead"}:
        raise HTTPException(403, "Only admin/manager/HSEQ can import SWMS")

    async with httpx.AsyncClient(timeout=20, follow_redirects=True) as c:
        r = await c.get(body.url)
        if r.status_code >= 400:
            raise HTTPException(400, f"Could not fetch .docx ({r.status_code})")
        content = r.content

    try:
        from docx import Document  # python-docx
    except Exception:
        raise HTTPException(500, "python-docx not installed on server")

    try:
        d = Document(io.BytesIO(content))
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"Invalid .docx: {e}")

    # Heuristic extraction — pull headings + bullets. The richer table
    # extraction is deferred to the admin's edit pass.
    title_guess = (d.core_properties.title or "").strip()
    paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]

    # First non-empty H1/H2 paragraph is usually the title.
    code = None; version = None
    for txt in paragraphs[:20]:
        m = re.match(r"^(SWMS[-\s]?\d+)\b.*?(V\d+(?:\.\d+)?)?", txt, re.I)
        if m:
            code = m.group(1).upper().replace(" ", "-")
            version = m.group(2)
            break

    inferred = {
        "title": title_guess or (paragraphs[0] if paragraphs else "Untitled SWMS"),
        "code": code or "SWMS-XX",
        "version": version or "V1.0",
        "scope": "",
        "high_risk_construction_work": "",
        "prepared_by": {"name": "", "role": "", "organisation": "", "date_prepared": ""},
        "approved_by": {"name": "", "position": "", "contact": "", "date_approved": ""},
        "ppe": [], "training_requirements": [], "equipment_list": [],
        "activity_analysis": [],
        "environmental_risks": [],
        "emergency_procedures": {"general": "", "accident_incident": "", "fire": "", "spill": ""},
        "status": "draft",
        "attendance_sheet_template": True,
        "source_file": {
            "url": body.url,
            "filename": body.url.split("/")[-1].split("?")[0],
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        "_diagnostics": {
            "paragraph_count": len(paragraphs),
            "table_count": len(d.tables),
            "note": "Heuristic parse — admin should verify before saving.",
        },
    }
    return inferred


# ────────────────── Phase 4.1 — Version chain history + backfill ──────────────────

@router.get("/{swms_id}/history")
async def swms_history(swms_id: str, user: dict = Depends(get_current_user)):
    """Walk the chain in BOTH directions from `swms_id` so an admin opening
    history on either the latest or any ancestor sees the full lineage.

    Algorithm:
      1. Walk backward via `supersedes` (ancestors).
      2. Walk forward via `superseded_by` (descendants).
      3. Concatenate ancestors + start + descendants.
      4. Cap at 20 nodes across the combined walk (loop defence).
      5. Final sort tie-breaker on created_at (oldest first).
    """
    cur = await db.swms.find_one({"id": swms_id, "org_id": user["org_id"]}, {"_id": 0})
    if not cur:
        raise HTTPException(404, "SWMS not found")

    seen: set[str] = {cur["id"]}
    ancestors: list[dict] = []
    descendants: list[dict] = []

    # ── Ancestors (backward).
    node = cur
    for _ in range(20):
        prev_id = node.get("supersedes")
        if not prev_id or prev_id in seen:
            break
        node = await db.swms.find_one({"id": prev_id, "org_id": user["org_id"]}, {"_id": 0})
        if not node:
            break
        ancestors.append(node); seen.add(node["id"])

    # ── Descendants (forward).
    node = cur
    for _ in range(20):
        if len(seen) >= 20:
            break
        nxt_id = node.get("superseded_by")
        if not nxt_id or nxt_id in seen:
            break
        node = await db.swms.find_one({"id": nxt_id, "org_id": user["org_id"]}, {"_id": 0})
        if not node:
            break
        descendants.append(node); seen.add(node["id"])

    chain = list(reversed(ancestors)) + [cur] + descendants
    # Final tie-breaker — keep oldest-first regardless of walk direction so the
    # UI timeline always reads V1 → V2 → V3.
    chain.sort(key=lambda d: (d.get("created_at") or "", d.get("version") or ""))
    return {"swms_id": swms_id, "chain": chain, "depth": len(chain)}


@router.get("/{swms_id}/diff/{previous_id}")
async def swms_diff(swms_id: str, previous_id: str,
                    user: dict = Depends(get_current_user)):
    """Phase 4.1c — structured delta between two SWMS revisions.

    Computes set-diffs on the four list-shaped fields auditors care about
    most: hazards, controls, ppe, activity_analysis. Each field returns
    {added:[...], removed:[...], unchanged:[...]} so the UI can render
    green/strikethrough/grey pills side-by-side."""
    cur = await db.swms.find_one({"id": swms_id, "org_id": user["org_id"]}, {"_id": 0})
    prev = await db.swms.find_one({"id": previous_id, "org_id": user["org_id"]}, {"_id": 0})
    if not cur or not prev:
        raise HTTPException(404, "SWMS not found")

    def _items(d: dict, key: str) -> list[str]:
        v = d.get(key)
        if isinstance(v, list):
            return [str(x).strip() for x in v if x is not None and str(x).strip()]
        if isinstance(v, str):
            return [s.strip() for s in v.splitlines() if s.strip()]
        return []

    out: dict = {}
    for k in ("hazards", "controls", "ppe", "activity_analysis"):
        a = set(_items(prev, k)); b = set(_items(cur, k))
        out[k] = {
            "added": sorted(b - a),
            "removed": sorted(a - b),
            "unchanged": sorted(a & b),
        }
    return {
        "swms_id": swms_id,
        "previous_id": previous_id,
        "header": {
            "from_version": prev.get("version"),
            "to_version": cur.get("version"),
            "from_created_at": prev.get("created_at"),
            "to_created_at": cur.get("created_at"),
        },
        "diff": out,
    }


admin_router = APIRouter(prefix="/admin/swms", tags=["admin-swms"])


@admin_router.post("/backfill-version-chain")
async def backfill_version_chain(user: dict = Depends(get_current_user)):
    """One-shot: link duplicates with the same `title` in import_date order.
    Idempotent — rows already in a chain are skipped."""
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    by_title: dict[str, list[dict]] = {}
    async for r in db.swms.find({"org_id": user["org_id"], "deleted_at": None}, {"_id": 0}):
        t = (r.get("title") or "").strip()
        if not t:
            continue
        by_title.setdefault(t, []).append(r)
    linked = 0
    skipped = 0
    for title, rows in by_title.items():
        if len(rows) < 2:
            continue
        rows.sort(key=lambda d: d.get("created_at") or "")
        for i in range(1, len(rows)):
            prev, curr = rows[i - 1], rows[i]
            if curr.get("supersedes") or prev.get("superseded_by"):
                skipped += 1
                continue
            await db.swms.update_one({"id": prev["id"]},
                {"$set": {"superseded_by": curr["id"], "status": "superseded",
                          "updated_at": now_iso()}})
            await db.swms.update_one({"id": curr["id"]},
                {"$set": {"supersedes": prev["id"], "updated_at": now_iso()}})
            log.info("backfill linked: %s -> %s (title=%r)", prev["id"], curr["id"], title)
            linked += 1
    return {"linked": linked, "skipped": skipped}


# ────────────────── Phase 4.1 — SWMS Assignments ──────────────────

class AssignmentsIn(BaseModel):
    applies_to: dict  # {roles:[], worker_ids:[], company_ids:[], asset_types:[]}


class BulkAssignmentsIn(BaseModel):
    swms_ids: list[str]
    applies_to: dict


def _clean_applies_to(raw: dict) -> dict:
    raw = raw or {}
    return {
        "roles":        [str(x) for x in (raw.get("roles") or [])],
        "worker_ids":   [str(x) for x in (raw.get("worker_ids") or [])],
        "company_ids":  [str(x) for x in (raw.get("company_ids") or [])],
        "asset_types":  [str(x) for x in (raw.get("asset_types") or [])],
    }


@router.get("/assignments")
async def list_assignments(user: dict = Depends(get_current_user)):
    out = {}
    async for s in db.swms.find(
        {"org_id": user["org_id"], "deleted_at": None,
         "status": {"$ne": "superseded"}},
        {"_id": 0, "id": 1, "applies_to": 1},
    ):
        out[s["id"]] = _clean_applies_to(s.get("applies_to") or {})
    return out


def _require_swms_edit(user: dict):
    if user.get("role") not in {"admin", "manager", "hseq_lead"}:
        raise HTTPException(403, "Permission denied: swms.edit")


@router.put("/assignments/bulk")
async def put_assignment_bulk(body: BulkAssignmentsIn,
                              user: dict = Depends(get_current_user)):
    _require_swms_edit(user)
    if not body.swms_ids:
        raise HTTPException(400, "swms_ids cannot be empty")
    cleaned = _clean_applies_to(body.applies_to)
    r = await db.swms.update_many(
        {"id": {"$in": body.swms_ids}, "org_id": user["org_id"], "deleted_at": None},
        {"$set": {"applies_to": cleaned, "updated_at": now_iso(),
                  "updated_by": user["id"]}},
    )
    return {"matched": r.matched_count, "modified": r.modified_count,
            "applies_to": cleaned}


@router.put("/assignments/{swms_id}")
async def put_assignment(swms_id: str, body: AssignmentsIn,
                         user: dict = Depends(get_current_user)):
    _require_swms_edit(user)
    cleaned = _clean_applies_to(body.applies_to)
    result = await db.swms.find_one_and_update(
        {"id": swms_id, "org_id": user["org_id"], "deleted_at": None},
        {"$set": {"applies_to": cleaned, "updated_at": now_iso(),
                  "updated_by": user["id"]}},
        return_document=True, projection={"_id": 0},
    )
    if not result:
        raise HTTPException(404, "SWMS not found")
    return {"swms_id": swms_id, "applies_to": cleaned}


def now_iso():  # local re-export so this module is self-sufficient
    return datetime.now(timezone.utc).isoformat()
